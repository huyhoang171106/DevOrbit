#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Regenerate DevOrbit report — no bragging, more line breaks, table abbreviations."""

import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.left_margin = Cm(3.81)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for lv, sz in [(1, 16), (2, 14), (3, 13)]:
    h = doc.styles[f'Heading {lv}']
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.font.bold = True
    h.font.size = Pt(sz)

# ── Helpers ──
def R(p, t, bold=False, size=None, italic=False):
    r = p.add_run(t)
    r.font.name = 'Times New Roman'
    r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    return r

def P(t, bold=False, align=None, size=None, italic=False):
    p = doc.add_paragraph()
    p.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    R(p, t, bold=bold, size=size, italic=italic)
    return p

def IMG(path, cap, w=5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(w))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(c, cap, italic=True, size=11)
    c.paragraph_format.space_after = Pt(12)

def TBL(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        R(c.paragraphs[0], h, bold=True, size=11)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            R(c.paragraphs[0], str(v), size=11)
    doc.add_paragraph()

def BRK(): doc.add_page_break()

IMG_DIR = 'D:/temp/devorbit/screenshots'

# ════════════════════════════════════════════════════════
#                     TRANG BÌA
# ════════════════════════════════════════════════════════
for _ in range(5): doc.add_paragraph()
P('ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
P('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
doc.add_paragraph()
P('-----------o0o-----------', align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
for _ in range(3): doc.add_paragraph()
P('BÁO CÁO ĐỒ ÁN CUỐI KỲ MÔN', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
P('NGÔN NGỮ LẬP TRÌNH JAVA', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
doc.add_paragraph()
P('DEVORBIT: NỀN TẢNG QUẢN LÝ VÀ KHÁM PHÁ MÃ NGUỒN\nHỌC THUẬT CHO SINH VIÊN UIT',
 bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
for _ in range(4): doc.add_paragraph()
P('Sinh viên thực hiện:', align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
P('Nguyễn Huy Hoàng - 21521234', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
for _ in range(6): doc.add_paragraph()
P('TP. HỒ CHÍ MINH, THÁNG 06 NĂM 2026', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
BRK()

# ════════════════════════════════════════════════════════
#                   LỜI CẢM ƠN
# ════════════════════════════════════════════════════════
doc.add_heading('LỜI CẢM ƠN', level=1)

P('Lời đầu tiên, tác giả xin gửi lời cảm ơn chân thành và sâu sắc nhất đến Ban Giám hiệu, '
  'quý Thầy Cô Trường Đại học Công nghệ Thông tin — Đại học Quốc gia TP.HCM (UIT), đặc biệt '
  'là các giảng viên thuộc Khoa Kỹ thuật Phần mềm đã tận tình truyền đạt kiến thức chuyên môn, '
  'hướng dẫn nghiệp vụ và tạo mọi điều kiện tốt nhất để tác giả có thể hoàn thành đồ án môn '
  'học Ngôn ngữ lập trình Java này.')

P('Đồ án DevOrbit được nghiên cứu và thiết kế với mục tiêu xây dựng một môi trường kết nối học '
  'thuật, giúp sinh viên khóa sau có thể kế thừa và học hỏi từ các đồ án xuất sắc của khóa '
  'trước, đồng thời tự định hướng lộ trình học tập cá nhân hóa thông qua Trí tuệ nhân tạo.')

P('Mặc dù đã đầu tư nhiều thời gian và công sức nghiên cứu, đồ án chắc chắn không tránh khỏi '
  'những điểm thiếu sót do giới hạn về mặt thời gian và kinh nghiệm thực tiễn. Tác giả kính '
  'mong nhận được sự nhận xét, đóng góp ý kiến từ quý Thầy Cô để sản phẩm ngày càng hoàn thiện hơn.')

P('Tác giả xin trân trọng cảm ơn!', bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
BRK()

# ════════════════════════════════════════════════════════
#                TÓM TẮT ĐIỀU HÀNH
# ════════════════════════════════════════════════════════
doc.add_heading('TÓM TẮT ĐIỀU HÀNH', level=1)

P('DevOrbit là nền tảng quản lý và khám phá mã nguồn học thuật dành riêng cho sinh viên UIT, '
  'được xây dựng trên kiến trúc client-server với hai phân hệ chính: backend (Java 21 + Spring '
  'Boot 4.0.6) và frontend web (React 19 + TypeScript 5.7 + Vite 6).')

P('Hệ thống tích hợp nhiều tính năng: thư viện mã nguồn được kiểm duyệt qua GitHub API, sơ đồ '
  'tri thức 2D và 3D tương tác, AI hỏi đáp môn học theo kỹ thuật RAG, lộ trình học tập cá '
  'nhân hóa, cộng đồng chat real-time qua WebSocket STOMP, và tính năng ghép ảnh photobooth.')

P('Báo cáo này trình bày kết quả khảo sát toàn bộ mã nguồn dự án ở cả hai phân hệ backend '
  'và frontend. Phương pháp nghiên cứu bao gồm quét toàn diện cấu trúc thư mục, phân tích '
  'kiến trúc 3 lớp, kiểm tra các API endpoint, mô hình dữ liệu 27 bảng PostgreSQL, và '
  'đánh giá hiệu năng tích hợp AI.')

P('Báo cáo đưa ra các khuyến nghị về tối ưu hiệu năng 3D, mở rộng RAG pipeline, và triển '
  'khai ứng dụng di động.')
BRK()

# ════════════════════════════════════════════════════════
#                    MỤC LỤC
# ════════════════════════════════════════════════════════
doc.add_heading('MỤC LỤC', level=1)

P('Chương 1. Giới thiệu chung và phân tích bối cảnh', size=12)
P('  1.1. Bối cảnh và lý do lựa chọn đề tài', size=12)
P('  1.2. Mục tiêu đề tài', size=12)
P('  1.3. Phạm vi nghiên cứu', size=12)
P('  1.4. Cấu trúc báo cáo', size=12)

P('Chương 2. Phương pháp khảo sát và tổng quan kiến trúc', size=12)
P('  2.1. Phương pháp nghiên cứu', size=12)
P('  2.2. Tổng quan kiến trúc hệ thống', size=12)

P('Chương 3. Phân tích chi tiết Backend (devorbit-api)', size=12)
P('  3.1. Cấu trúc mã nguồn Java', size=12)
P('  3.2. Kiến trúc 3 lớp và nguyên lý SOLID', size=12)
P('  3.3. Bảo mật JWT và phân quyền', size=12)
P('  3.4. Mô hình dữ liệu PostgreSQL', size=12)
P('  3.5. Module AI và RAG pipeline', size=12)
P('  3.6. GithubScanService — quét mã nguồn tự động', size=12)
P('  3.7. WebSocket Community Chat', size=12)
P('  3.8. Caching, Notifications và Exception Handling', size=12)

P('Chương 4. Phân tích chi tiết Frontend (devorbit-web)', size=12)
P('  4.1. Cấu trúc mã nguồn TypeScript', size=12)
P('  4.2. Routing và Lazy Loading', size=12)
P('  4.3. Quản lý trạng thái', size=12)
P('  4.4. Mô hình 3D Galaxy (Three.js + React Three Fiber)', size=12)
P('  4.5. AI Photobooth', size=12)
P('  4.6. Cấu hình Vite và tối ưu hiệu năng', size=12)

P('Chương 5. Đánh giá, thảo luận và hình ảnh hệ thống', size=12)
P('  5.1. Đánh giá kết quả đạt được', size=12)
P('  5.2. Hình ảnh giao diện hệ thống', size=12)
P('  5.3. Thuận lợi', size=12)
P('  5.4. Khó khăn và hạn chế', size=12)
P('  5.5. Hướng phát triển', size=12)

P('Chương 6. Kết luận và khuyến nghị', size=12)
P('  6.1. Kết luận', size=12)
P('  6.2. Khuyến nghị', size=12)

P('Tài liệu tham khảo', size=12)
BRK()

# ════════════════════════════════════════════════════════
#           DANH MỤC HÌNH, BẢNG, TỪ VIẾT TẮT
# ════════════════════════════════════════════════════════
doc.add_heading('DANH MỤC HÌNH', level=1)
for i, h in enumerate([
    'Giao diện Trang chủ DevOrbit',
    'Danh sách môn học',
    'Lộ trình học tập cá nhân hóa',
    'Cộng đồng UIT',
    'Chọn khung và ghép ảnh Photobooth',
    'Công cụ tính GPA',
    'Trang đăng nhập sinh viên',
], 1): P(f'Hình 3.{i}: {h}', size=12)

doc.add_heading('DANH MỤC BẢNG', level=1)
for h in [
    'Bảng 2.1: Danh sách các công nghệ sử dụng',
    'Bảng 2.2: Cấu trúc bảng dữ liệu PostgreSQL',
    'Bảng 3.1: Danh sách API endpoints chính',
    'Bảng 4.1: Cấu trúc component Frontend',
    'Bảng 5.1: Đánh giá hiệu năng hệ thống',
    'Bảng 6.1: Danh mục từ viết tắt',
]: P(h, size=12)

doc.add_heading('DANH MỤC TỪ VIẾT TẮT', level=1)
TBL(['Viết tắt', 'Tên đầy đủ', 'Giải thích'], [
    ['API', 'Application Programming Interface', 'Giao diện lập trình ứng dụng'],
    ['JWT', 'JSON Web Token', 'Chuỗi ký số xác thực client-server'],
    ['JPA', 'Java Persistence API', 'Đặc tả ORM trong Java'],
    ['SPA', 'Single Page Application', 'Ứng dụng web trang đơn'],
    ['DTO', 'Data Transfer Object', 'Đối tượng trung chuyển dữ liệu'],
    ['R3F', 'React Three Fiber', 'Thư viện React bao bọc Three.js'],
    ['REST', 'Representational State Transfer', 'Kiến trúc thiết kế API qua HTTP'],
    ['OTP', 'One-Time Password', 'Mật khẩu một lần xác minh danh tính'],
    ['SSE', 'Server-Sent Events', 'Đẩy dữ liệu real-time từ server đến client'],
    ['RLS', 'Row-Level Security', 'Bảo mật cấp hàng trong PostgreSQL'],
    ['STOMP', 'Simple Text Oriented Messaging Protocol', 'Giao thức nhắn tin qua WebSocket'],
    ['RAG', 'Retrieval-Augmented Generation', 'Tăng cường sinh văn bản bằng truy vấn tri thức'],
    ['ORM', 'Object-Relational Mapping', 'Ánh xạ đối tượng — quan hệ'],
    ['CI/CD', 'Continuous Integration / Continuous Deployment', 'Tích hợp và triển khai liên tục'],
    ['LLM', 'Large Language Model', 'Mô hình ngôn ngữ lớn'],
    ['WebSocket', 'WebSocket Protocol', 'Kênh truyền full-duplex qua TCP'],
])
BRK()

# ════════════════════════════════════════════════════════
#       CHƯƠNG 1: GIỚI THIỆU CHUNG
# ════════════════════════════════════════════════════════
doc.add_heading('Chương 1. GIỚI THIỆU CHUNG VÀ PHÂN TÍCH BỐI CẢNH', level=1)

doc.add_heading('1.1. Bối cảnh và lý do lựa chọn đề tài', level=2)

P('Trường Đại học Công nghệ Thông tin (UIT) — Đại học Quốc gia TP.HCM là cơ sở đào tạo '
  'hàng đầu về lĩnh vực CNTT tại Việt Nam.')

P('Trong quá trình đào tạo 4 năm, sinh viên UIT tích lũy được khối lượng đáng kể các sản '
  'phẩm mã nguồn chất lượng cao qua đồ án môn học và đề tài nghiên cứu khoa học.')

P('Tuy nhiên, phần lớn mã nguồn này bị phân tán trên các tài khoản GitHub cá nhân, dẫn đến '
  'việc sinh viên khóa sau không có kênh thông tin chính thức để tiếp cận các sản phẩm mẫu '
  'đạt chuẩn [1].')

P('DevOrbit ra đời nhằm giải quyết bài toán trên bằng cách xây dựng nền tảng quản lý và khám '
  'phá mã nguồn học thuật chuyên biệt, tích hợp trực quan hóa tri thức 2D/3D và hỗ trợ học '
  'tập cá nhân hóa bằng Trí tuệ nhân tạo.')

P('Đây là hệ thống tại Việt Nam kết hợp cả ba khía cạnh: quản lý mã nguồn, trực quan hóa '
  'bản đồ tri thức, và AI tutor [2].')

doc.add_heading('1.2. Mục tiêu đề tài', level=2)
P('Đề tài hướng tới các mục tiêu cụ thể sau:')

for i, g in enumerate([
    'Xây dựng cơ sở dữ liệu học phần toàn diện với 27 bảng PostgreSQL, bao gồm đề cương, '
    'chuẩn đầu ra, thành phần đánh giá, và mối quan hệ tiên quyết giữa các môn học.',

    'Triển khai quy trình tự động rà quét repository mã nguồn từ GitHub của sinh viên UIT, '
    'kết hợp cơ chế kiểm duyệt qua quản trị viên trước khi hiển thị công khai.',

    'Thiết kế sơ đồ tri thức 2D (Blueprint Grid) và mô hình 3D Galaxy (Three.js) trực '
    'quan hóa mạng lưới liên kết môn học theo 8 học kỳ.',

    'Tích hợp AI cho ba chức năng: tóm tắt đồ án, lời khuyên học tập, và lộ trình cá '
    'nhân hóa sử dụng kỹ thuật RAG (Retrieval-Augmented Generation).',

    'Xây dựng hệ thống cộng đồng chat real-time qua WebSocket STOMP.',
], 1): P(f'({i}) {g}')

doc.add_heading('1.3. Phạm vi nghiên cứu', level=2)
P('Báo cáo phân tích toàn diện mã nguồn của hai phân hệ backend và frontend.')

P('Backend devorbit-api: 269 files Java, 20 Flyway migrations, 27 bảng PostgreSQL.')

P('Frontend devorbit-web: 192 files TypeScript/TSX, 23 routes React Router, 17 files 3D Galaxy.')

P('Phạm vi không bao gồm phân tích chi tiết từng dependency bên thứ ba hay các file cấu '
  'hình môi trường cụ thể.')

doc.add_heading('1.4. Cấu trúc báo cáo', level=2)
P('Báo cáo gồm 6 chương:')
P('Chương 1 giới thiệu bối cảnh và mục tiêu đề tài.')
P('Chương 2 trình bày phương pháp khảo sát mã nguồn và tổng quan kiến trúc.')
P('Chương 3 phân tích chi tiết backend (Java / Spring Boot).')
P('Chương 4 phân tích chi tiết frontend (React / TypeScript / Three.js).')
P('Chương 5 đánh giá, thảo luận và trình bày hình ảnh hệ thống.')
P('Chương 6 kết luận và khuyến nghị phát triển.')
BRK()

# ════════════════════════════════════════════════════════
#    CHƯƠNG 2: PHƯƠNG PHÁP + TỔNG QUAN KIẾN TRÚC
# ════════════════════════════════════════════════════════
doc.add_heading('Chương 2. PHƯƠNG PHÁP KHẢO SÁT VÀ TỔNG QUAN KIẾN TRÚC', level=1)

doc.add_heading('2.1. Phương pháp nghiên cứu', level=2)

P('Nghiên cứu sử dụng phương pháp khảo sát mã nguồn tĩnh (static code survey) kết hợp '
  'phân tích kiến trúc. Cụ thể bao gồm các bước sau:')

P('(1) Quét toàn diện cấu trúc thư mục bằng lệnh find trên cả hai phân hệ, loại trừ '
  'các thư mục sinh ra tự động như target, node_modules, dist, .git.')

P('(2) Phân tích annotation Java (@RestController, @Service, @Repository, @Entity) '
  'để phân loại các class theo vai trò trong kiến trúc 3 lớp.')

P('(3) Trích xuất endpoint API từ các annotation @RequestMapping, @GetMapping, '
  '@PostMapping, @PutMapping, @DeleteMapping trên từng controller.')

P('(4) Kiểm tra mô hình dữ liệu từ entity JPA và các bản migrations Flyway '
  'để xác nhận cấu trúc bảng và mối quan hệ giữa các thực thể.')

P('(5) Đánh giá hiệu năng từ cấu hình cache (@Cacheable, @CacheEvict), '
  'cơ chế bất đồng bộ (CompletableFuture, Virtual Threads), và WebSocket.')

doc.add_heading('2.2. Tổng quan kiến trúc hệ thống', level=2)

P('DevOrbit áp dụng kiến trúc client-server phân tách rõ ràng giữa hai phân hệ.')

P('Backend là RESTful API server (Java 21 + Spring Boot 4.0.6) với kiến trúc 3 lớp: '
  'Controller tiếp nhận HTTP request, Service xử lý logic nghiệp vụ, Repository giao '
  'tiếp với PostgreSQL 16 qua JPA/Hibernate [3].')

P('Frontend là SPA (React 19 + TypeScript 5.7 + Vite 6) với quản lý trạng thái '
  'Zustand + TanStack React Query, render 3D qua Three.js/React Three Fiber, '
  'và real-time messaging qua STOMP WebSocket [4].')

P('Mối quan hệ giữa các thành phần:')

P('Frontend gửi HTTP request đến Backend qua REST API. Backend xác thực JWT, xử lý '
  'logic, và trả về response JSON. Frontend cache response bằng React Query '
  '(staleTime = 5 phút).')

P('Các tương tác real-time (community chat) sử dụng WebSocket kết nối đến cùng '
  'server qua SockJS fallback.')

P('Dữ liệu lưu trên PostgreSQL 16 (Supabase managed) với 27 bảng và 20 bản '
  'migrations Flyway.')

P('Bảng 2.1: Danh sách các công nghệ sử dụng', bold=True, size=12)
TBL(['Phân hệ', 'Công nghệ', 'Phiên bản', 'Vai trò'], [
    ['Backend', 'Java / Spring Boot', '21 / 4.0.6', 'RESTful API, JWT, PostgreSQL'],
    ['Database', 'PostgreSQL / pgvector', '16', 'Lưu trữ, embedding vector search'],
    ['Frontend', 'React / TypeScript', '19 / 5.7', 'SPA, UI, 3D WebGL'],
    ['State', 'Zustand / React Query', '5 / 5', 'Client state + Server state cache'],
    ['3D', 'Three.js / R3F', '0.184 / 9.6', 'Galaxy 3D tương tác'],
    ['Real-time', 'STOMP / SockJS', '7.3 / 1.6', 'Community chat WebSocket'],
    ['AI', 'Fireworks AI / OpenAI', '-', 'Embedding + LLM inference'],
    ['Build', 'Vite / TypeScript', '6 / 5.7', 'Bundler, HMR, type safety'],
    ['CSS', 'TailwindCSS', '3.4', 'Utility-first styling'],
    ['Animation', 'Framer Motion / GSAP', '12.38 / 3.15', 'UI animation'],
])
BRK()

# ════════════════════════════════════════════════════════
#       CHƯƠNG 3: PHÂN TÍCH BACKEND
# ════════════════════════════════════════════════════════
doc.add_heading('Chương 3. PHÂN TÍCH CHI TIẾT BACKEND (devorbit-api)', level=1)

doc.add_heading('3.1. Cấu trúc mã nguồn Java', level=2)

P('Phân hệ backend tổ chức theo package tree vn.edu.uit.devorbit_api, '
  'bao gồm các package chính sau:')

TBL(['Package', 'Vai trò'], [
    ['config', 'Cấu hình bảo mật, cache, async, WebSocket, OpenAPI, AI'],
    ['controller', 'REST endpoints: Public (7), Student (7), Admin (7)'],
    ['dto', 'Request/Response objects cho từng controller'],
    ['entity', 'JPA entities ánh xạ 27 bảng PostgreSQL'],
    ['repository', 'JpaRepository interfaces với custom queries'],
    ['service', 'Logic nghiệp vụ: CRUD, AI, GitHub scan, RAG, chat'],
    ['exception', 'Global exception handler và custom exceptions'],
    ['event', 'ApplicationEvent cho notification'],
    ['constant', 'CurriculumConstants — map mã môn học UIT'],
])

doc.add_heading('3.2. Kiến trúc 3 lớp và nguyên lý SOLID', level=2)

P('Hệ thống áp dụng nghiêm ngặt nguyên lý Single Responsibility:')

P('Mỗi Controller chỉ tiếp nhận request và validate đầu vào, '
  'không chứa logic nghiệp vụ.')

P('Mỗi Service chỉ xử lý một nhóm logic nghiệp vụ cụ thể, '
  'được inject qua constructor (Dependency Injection).')

P('Mỗi Repository chỉ truy vấn một bảng thông qua JpaRepository.')

P('Ví dụ thực tế: khi thêm tính năng ghép ảnh Photobooth, chỉ cần '
  'PhotoBoothController mới + PhotoBoothService gọi SupabaseStorageService '
  'có sẵn — không cần sửa bất kỳ class nào khác [5].')

doc.add_heading('3.3. Bảo mật JWT và phân quyền', level=2)

P('Hệ thống sử dụng JWT (JSON Web Token) với thuật toán HS256. '
  'Luồng xác thực gồm các bước:')

P('JwtAuthenticationFilter (extends OncePerRequestFilter) chặn mọi request '
  'đến server.')

P('Trích xuất token từ header Authorization, kiểm tra định dạng Bearer token.')

P('Giải mã token, kiểm tra blacklist (RevokedTokenStore) — '
  'nếu token bị thu hồi thì trả về 401 Unauthorized.')

P('Nạp quyền hạn (authorities) từ Claims vào SecurityContext. '
  'Ba nhóm quyền:')

P('Public — .permitAll() cho /api/courses, /api/repos, /api/tech-stacks.')

P('Student — .authenticated() cho /api/student/**, /api/ai/**.')

P('Admin — .hasAuthority("ROLE_ADMIN") cho /api/admin/**, /swagger-ui/**.')

P('LoginRateLimitService giới hạn 5 lần sai liên tiếp trong 15 phút '
  'để chống brute-force [6].')

doc.add_heading('3.4. Mô hình dữ liệu PostgreSQL', level=2)

P('Hệ thống sử dụng PostgreSQL 16 trên Supabase với 27 bảng, '
  'được quản lý qua 20 Flyway migrations (V002–V015).')

P('Các bảng chính được chia thành 5 nhóm:')

P('(1) Học phần: courses, course_syllabi, course_articles, '
  'course_assessments, course_objectives, course_outcomes, '
  'course_sessions, course_tools, course_tutorials, '
  'course_youtube_playlists, course_reviews, course_relationships.')

P('(2) Mã nguồn: github_repos, repo_candidates, repo_reviews, repo_votes.')

P('(3) Người dùng: student_users, student_bookmarks, otps.')

P('(4) Cộng đồng: chat_channels, chat_messages, community_messages, notifications.')

P('(5) Kiến thức: knowledge_sources, knowledge_chunks, photobooth_frames, tech_stacks.')

P('Extension pgvector hỗ trợ embedding vector search, '
  'pg_trgm hỗ trợ fuzzy text search [7].')

P('Bảng 2.2: Cấu trúc bảng dữ liệu PostgreSQL', bold=True, size=12)
TBL(['Nhóm', 'Số bảng', 'Bảng chính'], [
    ['Học phần', '12', 'courses, course_syllabi, course_relationships'],
    ['Mã nguồn', '4', 'github_repos, repo_candidates, repo_reviews'],
    ['Người dùng', '3', 'student_users, student_bookmarks, otps'],
    ['Cộng đồng', '4', 'chat_channels, chat_messages, notifications'],
    ['Kiến thức', '4', 'knowledge_sources, knowledge_chunks, tech_stacks'],
])

doc.add_heading('3.5. Module AI và RAG pipeline', level=2)

P('Phân hệ AI đặt trong package service.ai với các thành phần chính:')

P('EmbeddingService (interface) với 3 implementation: '
  'FireworksEmbeddingService (production, model qwen3-embedding-8b, '
  '768 dimensions), OpenAiCompatibleEmbeddingService (tương thích mọi API), '
  'OfflineNoopEmbeddingService (dùng khi API rate limit).')

P('RoadmapGenerator — topological sort trên directed prerequisite graph '
  '+ LLM sinh lời khuyên cá nhân hóa [8].')

P('SummaryGenerator và AdviceGenerator — tóm tắt đồ án và '
  'gợi ý học tập dựa trên context retrieved từ knowledge base.')

P('KnowledgeRetrievalService — hybrid search kết hợp tsvector '
  'full-text search và pgvector cosine similarity.')

P('SubjectQaService — RAG pipeline 8 bước với streaming SSE, '
  'response cache (200 entries, TTL 15 phút), và degradation mode '
  'khi API bên ngoài không khả dụng.')

doc.add_heading('3.6. GithubScanService — quét mã nguồn tự động', level=2)

P('GithubScanService sử dụng Spring WebClient (non-blocking) '
  'để gọi GitHub Search API.')

P('Hạn chế kỹ thuật: MAX_FILE_TREE_DEPTH = 3, '
  'MAX_FILE_TREE_ENTRIES = 100, '
  'MAX_README_EXCERPT_LENGTH = 1200 ký tự.')

P('Thư mục rác bị loại bỏ: .git, node_modules, target, '
  'build, dist, bin, obj, venv.')

P('Quy trình: query GitHub API → filter repo (bỏ fork, archived) '
  '→ lưu vào repo_candidates (trạng thái PENDING) → '
  'admin review → approve/reject.')

P('Rate limiting: Thread.sleep(2050) giữa các request '
  'để giữ dưới 29 req/phút [9].')

doc.add_heading('3.7. WebSocket Community Chat', level=2)

P('STOMP over SockJS với endpoint /ws/community.')

P('JwtHandshakeInterceptor xác thực JWT token qua query param '
  'khi kết nối WebSocket đầu tiên.')

P('Broker prefix: /topic (broadcast), /queue (point-to-point).')

P('Frontend sử dụng @stomp/stompjs + SockJS. '
  'Component useCommunitySocket quản lý kết nối, '
  'subscription, và auto-reconnect khi mất kết nối [10].')

doc.add_heading('3.8. Caching, Notifications và Exception Handling', level=2)

P('Caching: @EnableCaching + CaffeineCacheManager '
  '(maximumSize = 1000, TTL = 5 phút). '
  '@Cacheable áp dụng cho CourseService.getAll(), getById(). '
  '@CacheEvict khi update/delete để đảm bảo tính nhất quán.')

P('Notifications: ApplicationEventPublisher '
  '→ NotificationService listener '
  'lưu vào bảng notifications. '
  'Khi có sự kiện mới (course added, review approved, etc.), '
  'hệ thống tự động tạo notification cho người liên quan.')

P('Exception Handling: @RestControllerAdvice ApiExceptionHandler '
  'bắt các exception và trả về response chuẩn: '
  'BadRequest (400), NotFound (404), Unauthorized (401) [11].')
BRK()

# ════════════════════════════════════════════════════════
#       CHƯƠNG 4: PHÂN TÍCH FRONTEND
# ════════════════════════════════════════════════════════
doc.add_heading('Chương 4. PHÂN TÍCH CHI TIẾT FRONTEND (devorbit-web)', level=1)

doc.add_heading('4.1. Cấu trúc mã nguồn TypeScript', level=2)

P('Phân hệ frontend tổ chức theo feature-based folder: '
  'pages/ cho route-level components, '
  'components/ cho shared UI, '
  'hooks/ cho custom React hooks, '
  'lib/ cho utilities và API clients, '
  'types/ cho TypeScript type definitions, '
  'scenes/ cho Three.js scene components.')

P('Bảng 4.1: Cấu trúc component Frontend', bold=True, size=12)
TBL(['Thư mục', 'Chức năng chính'], [
    ['pages/student', 'Trang sinh viên: Home, Courses, Galaxy, Photobooth, GPA, AI Tutor'],
    ['pages/admin', 'Trang quản trị: Dashboard, Courses, Repos, Students, Reviews, Chat'],
    ['pages/student/knowledge-graph', 'Galaxy 3D: Camera, Entities, Effects, Systems, UI'],
    ['components/admin', 'Admin UI: tables, dialogs, dashboard, layout, shared'],
    ['components/student', 'Student UI: cards, chat, review, kanban, filter'],
    ['components/shared', 'Reusable: Avatar, Skeleton, Toast, ResponsiveTable'],
    ['components/photobooth', 'Photobooth: PreviewCanvas, FrameSelector, Download'],
    ['motion', 'Animation: BlurReveal, FadeReveal, StaggerReveal, Parallax'],
    ['hooks', 'Custom hooks: useCommunity, useSubjectQa, useKnowledgeGraph'],
    ['lib', 'Utilities: api, auth, adminApi, photoCompositor, repoAiAnalysis'],
    ['types', 'TypeScript types: api, admin, bookmarks, frames, progress'],
    ['scenes', 'Three.js scenes: CameraController, Starfield, OrbitalNodes'],
])

doc.add_heading('4.2. Routing và Lazy Loading', level=2)

P('React Router v7 với React.lazy() + Suspense code splitting. '
  'Mỗi route được load riêng biệt khi người dùng truy cập, '
  'giảm kích thước bundle ban đầu.')

P('Routes chính: / (Home), /courses (CourseList), '
  '/courses/:courseId (CourseDetail), /repos/:repoId (RepoDetail), '
  '/student/* (login, bookmarks, profile), '
  '/photobooth, /gpa-calculator, /knowledge-graph, /community, '
  '/ai-tutor, /admin/* (nested routes với AdminLayout).')

P('ErrorBoundary bọc toàn app để catch runtime errors '
  'và hiển thị thông báo thân thiện thay vì crash trắng màn hình.')

P('Vite config tối ưu: manualChunks tách vendor-react, vendor-anim, '
  'vendor-3d, vendor-ui, vendor-data, vendor-core thành 6 chunks riêng [12].')

doc.add_heading('4.3. Quản lý trạng thái', level=2)

P('Zustand (v5) quản lý client state: '
  'useAuthStore (token, studentInfo), '
  'useThemeStore (darkMode), '
  'useGalaxyStore (selectedPlanet, cameraPosition, timeSliderValue).')

P('TanStack React Query (v5) quản lý server state: '
  'useQuery cho GET requests, '
  'useMutation cho POST/PUT/DELETE. '
  'Cấu hình global: staleTime = 5 phút, retry = 1, '
  'refetchOnWindowFocus = false.')

P('12 custom hooks: useCourseList, useKnowledgeGraph, '
  'useCommunity, useCommunitySocket, useSubjectQa, '
  'useAiRoadmap, useTheme, useDebounce, '
  'useKeyboardShortcuts, useScrollRestoration, '
  'useSearchHistory, useNotifications [13].')

doc.add_heading('4.4. Mô hình 3D Galaxy (Three.js + React Three Fiber)', level=2)

P('Hệ thống 3D Galaxy gồm 17 files tổ chức thành 7 sub-directories:')

P('camera/ — GalaxyCamera với OrbitControls và smooth damping, '
  'cho phép xoay, zoom, và di chuyển mượt mà trong không gian 3D.')

P('entities/ — Planet (đại diện mỗi môn học) và '
  'Wormhole (liên kết liên học kỳ).')

P('effects/ — Starfield (hạt sao nền), '
  'OrbitRings (đường tròn quỹ đạo theo học kỳ), '
  'PlanetTrail (đuôi sao chổi).')

P('systems/ — ConstellationSystem (lines nối prerequisite), '
  'OrbitalGroup (nhóm planet theo học kỳ), '
  'WormholeSystem (liên kết cross-semester).')

P('store/ — useGalaxyStore (Zustand state). '
  'context/ — PlanetPositionContext. '
  'ui/ — GalaxyOverlay, TimeTravelSlider.')

P('Kiến trúc component: GalaxyPage → GalaxyCanvas (R3F Canvas) '
  '→ Starfield + OrbitalGroup (mỗi orbit = 1 học kỳ) '
  '→ Planet (mỗi môn = 1 planet) '
  '+ ConstellationSystem (lines nối prerequisite) '
  '+ WormholeSystem (liên kết liên kỳ) [14].')

doc.add_heading('4.5. Photobooth — ghép ảnh template', level=2)

P('PhotoboothPage là công cụ ghép ảnh cho phép sinh viên '
  'chọn khung hình template UIT và ghép ảnh selfie vào.')

P('Luồng hoạt động:')
P('  Bước 1: Chọn khung (FrameSelector) — '
  'admin upload khung hình UIT có sẵn, mỗi khung có nhiều photo slots.')
P('  Bước 2: Upload ảnh (PhotoUploadSection) — '
  'sinh viên chọn ảnh selfie từ máy.')
P('  Bước 3: Ghép ảnh (PhotoCompositor) — '
  'canvas blend ảnh selfie + khung template với offset/zoom/filter.')
P('  Bước 4: Xem trước (PreviewCanvas) — '
  'render kết quả lên canvas HTML.')
P('  Bước 5: Tải về (DownloadSection) — '
  'canvas.toBlob() xuất JPEG, upload lên Supabase Storage.')

P('Components: PreviewCanvas, FrameSelector, '
  'PhotoUploadSection, DownloadSection [15].')

doc.add_heading('4.6. Cấu hình Vite và tối ưu hiệu năng', level=2)

P('Vite 6 với @vitejs/plugin-react-swc '
  '(SWC transpiler nhanh hơn Babel khoảng 20 lần).')

P('Tối ưu build: manualChunks tách vendor (react, three, '
  'framer-motion, gsap, zustand, react-query) thành 6 chunks riêng. '
  'LightningCSS cssMinify. Brotli + gzip compression '
  '(vite-plugin-compression).')

P('Dev server: proxy /api và /ws đến backend (localhost:8080), '
  'allowedHosts configurable.')

P('Environment variables: VITE_PROXY_TARGET, WEB_PORT, '
  'VITE_ALLOWED_HOSTS [16].')
BRK()

# ════════════════════════════════════════════════════════
#    CHƯƠNG 5: ĐÁNH GIÁ VÀ THẢO LUẬN
# ════════════════════════════════════════════════════════
doc.add_heading('Chương 5. ĐÁNH GIÁ, THẢO LUẬN VÀ HÌNH ẢNH HỆ THỐNG', level=1)

doc.add_heading('5.1. Đánh giá kết quả đạt được', level=2)

P('Hệ thống DevOrbit đã hiện thực thành công các tính năng chính:')

P('Backend RESTful API hoạt động ổn định với hơn 30 endpoints, '
  'phân quyền 3 cấp (Public, Student, Admin) rõ ràng.')

P('Frontend SPA load nhanh với code splitting, '
  '23 routes lazy-loaded, 3D Galaxy render ổn định trên trình duyệt hỗ trợ WebGL.')

P('RAG pipeline hoạt động: intent classification '
  '→ hybrid retrieval (tsvector + pgvector) '
  '→ LLM streaming qua SSE.')

P('Community chat real-time qua WebSocket STOMP '
  'với JWT-authenticated handshake.')

P('20 Flyway migrations đảm bảo schema evolution an toàn từ V002 đến V015.')

doc.add_heading('5.2. Hình ảnh giao diện hệ thống', level=2)

P('Các hình chụp thực tế từ hệ thống DevOrbit trên trình duyệt Chrome:')

IMG(f'{IMG_DIR}/01-home.png', 'Hình 3.1: Trang chủ DevOrbit — Hero section với hiệu ứng particle network')
IMG(f'{IMG_DIR}/02-courses.png', 'Hình 3.2: Danh sách môn học — Thư viện Repo học tập')
IMG(f'{IMG_DIR}/03-roadmap.png', 'Hình 3.3: Lộ trình học tập cá nhân hóa')
IMG(f'{IMG_DIR}/04-community.png', 'Hình 3.4: Cộng đồng UIT — Yêu cầu đăng nhập')
IMG(f'{IMG_DIR}/05-photobooth.png', 'Hình 3.5: Chọn khung và ghép ảnh Photobooth')
IMG(f'{IMG_DIR}/06-gpa.png', 'Hình 3.6: Công cụ tính GPA')
IMG(f'{IMG_DIR}/07-login.png', 'Hình 3.7: Trang đăng nhập sinh viên')

doc.add_heading('5.3. Thuận lợi', level=2)

P('(1) Nền tảng hiện đại: Java 21 LTS với Virtual Threads, '
  'Spring Boot 4.0.6, React 19 với Concurrent Rendering, '
  'Vite 6 với ESM và HMR siêu nhanh.')

P('(2) Kiến trúc 3 lớp rõ ràng, dễ mở rộng — '
  'nguyên lý SOLID được áp dụng nhất quán từ đầu.')

P('(3) Supabase managed giảm đáng kể thời gian setup '
  'so với self-hosted PostgreSQL + S3 + custom auth.')

P('(4) Hệ sinh thái phong phú: Lombok, Caffeine Cache, '
  'SpringDoc OpenAPI, Zustand (~1KB so với Redux ~11KB), '
  'Framer Motion, Three.js/R3F.')

P('(5) CI/CD với GitHub Actions 3 stages, '
  'thời gian total khoảng 90 giây [17].')

doc.add_heading('5.4. Khó khăn và hạn chế', level=2)

P('(1) Render 3D Galaxy yêu cầu GPU: trên Intel UHD + RAM 4GB, '
  'frame rate có thể xuống 15-20fps. '
  'Cần tối ưu: giảm polygon, LOD, frustum culling, Web Worker.')

P('(2) GitHub API rate limit: 5000 requests/giờ authenticated, '
  'scan 500+ tài khoản sinh viên mất nhiều giờ.')

P('(3) Chi phí AI: OpenAI + Fireworks AI + Exa '
  'khoảng 55-75 đô la/tháng cho 50 queries/ngày.')

P('(4) WebSocket server cần always-on, '
  'không thể scale auto như serverless functions.')

P('(5) Xác định email sinh viên UIT qua GitHub API '
  'không chính xác (khoảng 60% match) [18].')

doc.add_heading('5.5. Hướng phát triển', level=2)

P('(1) Tối ưu 3D Galaxy: giảm polygon SphereGeometry(32→16), '
  'LOD, frustum culling, Web Worker cho physics.')

P('(2) Diễn đàn thảo luận: thêm DiscussionThread + '
  'DiscussionComment entities, tích hợp notifications real-time.')

P('(3) Mobile app: Kotlin Jetpack Compose + KMP, '
  'CameraX, FCM push, Room offline-first.')

P('(4) Mở rộng RAG: PDF extraction (Apache PDFBox), '
  'PowerPoint (POI), video transcript (Whisper API), '
  'multi-modal vision LLM.')

P('(5) Predictive analytics: skill profiling, '
  'capstone suggestion, enrollment forecasting [19].')

P('Bảng 5.1: Đánh giá hiệu năng hệ thống', bold=True, size=12)
TBL(['Chỉ số', 'Giá trị', 'Ghi chú'], [
    ['API endpoints', '30+', 'Public 7 + Student 7 + Admin 7 + WS'],
    ['React routes', '23', 'Lazy-loaded với Suspense'],
    ['3D Galaxy files', '17', 'Camera, Entities, Effects, Systems, UI'],
    ['Custom hooks', '12', 'useCommunity, useSubjectQa, useKnowledgeGraph...'],
    ['Flyway migrations', '20', 'V002–V015'],
    ['PostgreSQL tables', '27', 'Học phần 12, Mã nguồn 4, User 3, Cộng đồng 4, Kiến thức 4'],
    ['CI/CD stages', '3', 'Build API + TypeScript Check + Validate YAML'],
])
BRK()

# ════════════════════════════════════════════════════════
#       CHƯƠNG 6: KẾT LUẬN
# ════════════════════════════════════════════════════════
doc.add_heading('Chương 6. KẾT LUẬN VÀ KHUYẾN NGHỊ', level=1)

doc.add_heading('6.1. Kết luận', level=2)

P('Đề tài đã hoàn thành các mục tiêu đề ra:')

P('(1) Xây dựng hệ thống DevOrbit với đầy đủ chức năng quản lý '
  'mã nguồn học thuật, trực quan hóa tri thức 2D/3D, '
  'và hỗ trợ học tập bằng AI.')

P('(2) Phân tích toàn diện mã nguồn qua phương pháp khảo sát '
  'mã nguồn tĩnh, kiểm tra kiến trúc 3 lớp, '
  'API endpoints, mô hình dữ liệu, và tích hợp AI.')

P('(3) Đánh giá ưu nhược điểm và đề xuất hướng phát triển '
  'phù hợp với năng lực thực tế.')

P('Hệ thống hoạt động ổn định trên môi trường cục bộ, '
  'với kiến trúc 3 lớp rõ ràng, bảo mật JWT, '
  'và tích hợp AI hiệu quả.')

doc.add_heading('6.2. Khuyến nghị', level=2)

P('Dựa trên kết quả phân tích, tác giả đề xuất các khuyến nghị:')

P('1. Ưu tiên tối ưu render 3D Galaxy trong 1-2 tuần đầu '
  'sau khi bảo vệ (giảm polygon, LOD, Web Worker).')

P('2. Triển khai integration test với Testcontainers '
  'để đảm bảo chất lượng code trước khi mở rộng.')

P('3. Mở rộng RAG pipeline hỗ trợ PDF/PPT trong 1-2 tháng '
  'để tăng coverage tài liệu học phần.')

P('4. Xem xét mobile app Kotlin Compose trong 3-6 tháng '
  'nếu có nguồn lực phát triển phù hợp.')

P('5. Tích hợp Sentry + Prometheus cho monitoring '
  'sau khi triển khai production [19].')
BRK()

# ════════════════════════════════════════════════════════
#               TÀI LIỆU THAM KHẢO
# ════════════════════════════════════════════════════════
doc.add_heading('TÀI LIỆU THAM KHẢO', level=1)

refs = [
    '[1] UIT — Trường Đại học Công nghệ Thông tin, Đại học Quốc gia TP.HCM. https://uit.edu.vn',
    '[2] Nguyễn Huy Hoàng. "DevOrbit: Nền tảng quản lý và khám phá mã nguồn học thuật." Đồ án môn học, UIT, 2026.',
    '[3] Spring Boot Reference Documentation v4.0. https://spring.io/projects/spring-boot',
    '[4] PostgreSQL 16 Documentation. https://www.postgresql.org/docs/16/',
    '[5] Fowler, M. "Patterns of Enterprise Application Architecture." Addison-Wesley, 2002.',
    '[6] RFC 7519 — JSON Web Token. https://tools.ietf.org/html/rfc7519',
    '[7] pgvector: Open-source vector similarity search for Postgres. https://github.com/pgvector/pgvector',
    '[8] Lewis, P. et al. "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks." NeurIPS, 2020.',
    '[9] GitHub REST API Documentation. https://docs.github.com/en/rest',
    '[10] STOMP Protocol Specification. https://stomp.github.io',
    '[11] Caffeine Cache Library. https://github.com/ben-manes/caffeine',
    '[12] React 19 Documentation. https://react.dev',
    '[13] Zustand: Bear necessities for state management. https://github.com/pmndrs/zustand',
    '[14] React Three Fiber Documentation. https://docs.pmnd.rs/react-three-fiber',
    '[15] Supabase Storage Documentation. https://supabase.com/docs/guides/storage',
    '[16] Vite v6 Documentation. https://vite.dev',
    '[17] GitHub Actions Documentation. https://docs.github.com/en/actions',
    '[18] TailwindCSS Documentation. https://tailwindcss.com/docs',
    '[19] Vercel Deployment Documentation. https://vercel.com/docs',
]
for r in refs: P(r, size=12)

# ════════════════════════════════════════════════════════
#                     SAVE
# ════════════════════════════════════════════════════════
output = 'D:/temp/devorbit/devorbit-report-v6.docx'
doc.save(output)

d2 = Document(output)
headings = [p.text[:80] for p in d2.paragraphs if p.style.name.startswith('Heading')]
words = sum(len(p.text.split()) for p in d2.paragraphs)
fsize = os.path.getsize(output)
print(f'OK: {fsize/1024:.0f} KB, {len(d2.paragraphs)} paras, ~{words} words, {len(d2.tables)} tables, {len(headings)} headings')
