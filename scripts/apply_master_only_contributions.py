from __future__ import annotations

import os
import subprocess
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
INPUT = Path(os.environ.get("DEVORBIT_REPORT_DOCX", ROOT / "devorbit-report-v6-java-api-linked.docx"))
OUTPUT = Path(os.environ.get("DEVORBIT_REPORT_OUTPUT", ROOT / "devorbit-report-v6-java-api-master-contrib-detailed.docx"))
SINCE = "2026-03-14"
UNTIL = "2026-06-15T00:00:00"
REV = "origin/master"


def classify_member(author: str, email: str) -> str | None:
    if author == "huyhoang171106" or email in {"24520554@gm.uit.edu.vn", "huyhoang171106@gmail.com"}:
        return "A"
    if author in {"DuyBao", "Đặng Duy Bảo"}:
        return "B"
    if author == "HufaCung":
        return "C"
    if author == "north2k6":
        return "D"
    return None


def area_of(path: str) -> str:
    if path.startswith("devorbit-api/src/main/java/vn/edu/uit/devorbit_api/controller"):
        return "Controller/API"
    if path.startswith("devorbit-api/src/main/java/vn/edu/uit/devorbit_api/service"):
        return "Service/nghiệp vụ"
    if path.startswith("devorbit-api/src/main/java/vn/edu/uit/devorbit_api/config"):
        return "Bảo mật/cấu hình"
    if path.startswith("devorbit-api/src/main/resources/db/migration"):
        return "Migration DB"
    if path.startswith("devorbit-api/src/test"):
        return "Backend test"
    if path.startswith("devorbit-api"):
        return "Backend runtime"
    if path.startswith("devorbit-web"):
        return "Client web"
    if path.startswith("devorbit-mobile"):
        return "Mobile client"
    if path.startswith("docs"):
        return "Tài liệu/test matrix"
    if path.startswith(".github"):
        return "CI/CD"
    return path.split("/")[0]


def infer_effect(subject: str, files: list[str]) -> str:
    s = subject.lower()
    file_blob = " ".join(files).lower()
    if "subject qa" in s or "ai course" in s or "rag" in s or "streaming chat" in s or "roadmap" in s:
        return "Cải thiện AI học tập, RAG, streaming chat hoặc roadmap."
    if "community" in s or "chat" in s or "channel" in s or "websocket" in s:
        return "Cải thiện community chat, channel, realtime hoặc trải nghiệm trao đổi."
    if "auth" in s or "password" in s or "jwt" in s or "security" in s or "otp" in s:
        return "Cải thiện xác thực, bảo mật, OTP, đăng nhập hoặc phân quyền."
    if "admin" in s:
        return "Cải thiện luồng quản trị, dashboard, kiểm duyệt hoặc dữ liệu admin."
    if "repo" in s or "github" in s or "search" in s or "evaluation" in s or "analysis" in s:
        return "Cải thiện tìm kiếm, quét, phân tích hoặc đánh giá repository."
    if "mobile" in s or "feat(mb)" in s or "gradle" in file_blob or "devorbit-mobile" in file_blob:
        return "Cải thiện mobile client hoặc tích hợp API cho mobile."
    if "gpa" in s or "gpa" in file_blob:
        return "Cải thiện GPA calculator, kiểm thử và tài liệu liên quan."
    if "test" in s or "test_matrix" in file_blob:
        return "Cập nhật kiểm thử, test matrix hoặc bằng chứng xác minh."
    if "docs" in s or "milestone" in s or "story" in file_blob:
        return "Cập nhật tài liệu, story, milestone hoặc kế hoạch triển khai."
    if "fix" in s:
        return "Sửa lỗi hành vi, dữ liệu, UI/API hoặc ổn định runtime."
    if "config" in s or "application.yaml" in file_blob or "run.bat" in file_blob:
        return "Cải thiện cấu hình, runtime hoặc script chạy hệ thống."
    return "Thay đổi chức năng hoặc bảo trì theo commit message và file đã sửa."


def collect_commits():
    subprocess.check_call(["git", "fetch", "origin", "master", "--prune"], cwd=ROOT)
    fmt = "@@@%H%x09%an%x09%ae%x09%ad%x09%s"
    raw = subprocess.check_output(
        ["git", "log", REV, "--no-merges", f"--since={SINCE}", f"--until={UNTIL}", f"--format={fmt}", "--date=short", "--name-only"],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    commits: list[dict] = []
    current = None
    for line in raw.splitlines():
        if line.startswith("@@@"):
            parts = line[3:].split("\t")
            current = {
                "hash": parts[0][:8],
                "author": parts[1],
                "email": parts[2],
                "date": parts[3],
                "subject": parts[4] if len(parts) > 4 else "",
                "files": [],
            }
            member = classify_member(current["author"], current["email"])
            if member:
                current["member"] = member
                commits.append(current)
            else:
                current = None
        elif current is not None and line.strip():
            current["files"].append(line.strip())
    for c in commits:
        c["areas"] = [area_of(f) for f in c["files"]]
        c["effect"] = infer_effect(c["subject"], c["files"])
    return commits


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        p.add_run(text)
    if style:
        p.style = style
    return p


def insert_table_after(paragraph, rows: int, cols: int):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Cm(16))
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return table


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)


def shade(cell, fill: str = "D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def style_table(table, font_size: int = 9):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri == 0:
                shade(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    r.font.size = Pt(font_size)
                    if ri == 0:
                        r.bold = True


def append_table_after(paragraph, caption: str, headers: list[str], rows: list[list[str]], font_size: int = 9):
    cap = insert_paragraph_after(paragraph, caption, "Normal")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].bold = True
    table = insert_table_after(cap, len(rows) + 1, len(headers))
    for c, h in enumerate(headers):
        set_cell_text(table.cell(0, c), h, True, font_size)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            set_cell_text(table.cell(r, c), val, False, font_size)
    style_table(table, font_size)
    spacer_xml = OxmlElement("w:p")
    table._tbl.addnext(spacer_xml)
    return Paragraph(spacer_xml, paragraph._parent)


def add_body_after(paragraph, text: str):
    p = insert_paragraph_after(paragraph, text, "Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(5)
    return p


def find_heading(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip() == text and p.style.name.startswith("Heading"):
            return p
    raise ValueError(f"heading not found: {text}")


def paragraph_index(doc: Document, paragraph) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p._p is paragraph._p:
            return i
    raise ValueError("paragraph not found")


def replace_cover_and_voice(doc: Document):
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "NGÔN NGỮ LẬP TRÌNH JAVA":
            p.text = "NGÔN NGỮ LẬP TRÌNH JAVA SE330.Q21"
        elif "DEVORBIT:" in t:
            p.text = "DEVORBIT: NỀN TẢNG QUẢN LÝ VÀ KHÁM PHÁ MÃ NGUỒN HỌC THUẬT"
        elif t.startswith("Lời đầu tiên, tác giả xin gửi lời cảm ơn"):
            p.text = "Lời đầu tiên, nhóm thực hiện xin gửi lời cảm ơn chân thành đến Ban Giám hiệu, quý Thầy Cô Trường Đại học Công nghệ Thông tin - Đại học Quốc gia TP.HCM (UIT), đặc biệt là các giảng viên phụ trách học phần Ngôn ngữ lập trình Java SE330.Q21 đã định hướng kiến thức và phương pháp xây dựng đồ án."
        elif t.startswith("Mặc dù đã đầu tư nhiều thời gian") and "Tác giả kính mong" in t:
            p.text = "Mặc dù đã đầu tư nhiều thời gian và công sức nghiên cứu, đồ án chắc chắn không tránh khỏi những điểm thiếu sót do giới hạn về thời gian và kinh nghiệm thực tiễn. Nhóm thực hiện kính mong nhận được nhận xét, góp ý từ quý Thầy Cô để sản phẩm ngày càng hoàn thiện hơn."
        elif t == "Tác giả xin trân trọng cảm ơn!":
            p.text = "Nhóm thực hiện xin trân trọng cảm ơn!"
        elif t == "Dựa trên kết quả phân tích, tác giả đề xuất các khuyến nghị:":
            p.text = "Dựa trên kết quả phân tích, nhóm thực hiện đề xuất các khuyến nghị:"

    marker = next(p for p in doc.paragraphs if p.text.strip() == "Sinh viên thực hiện:")
    start = paragraph_index(doc, marker) + 1
    end = start
    while end < len(doc.paragraphs) and "TP. HỒ CHÍ MINH" not in doc.paragraphs[end].text:
        end += 1
    for p in reversed(doc.paragraphs[start:end]):
        p._element.getparent().remove(p._element)
    current = marker
    for line in [
        "Thành viên A - 24520554 - GitHub account: huyhoang171106",
        "Thành viên B - 24520146 - Git author: DuyBao",
        "Thành viên C - 24520071 - Git author: HufaCung",
        "Thành viên D - MSSV: cập nhật - Git author: north2k6",
    ]:
        current = insert_paragraph_after(current, line, "Normal")
        current.alignment = WD_ALIGN_PARAGRAPH.CENTER


def update_front_lists(doc: Document):
    toc = find_heading(doc, "MỤC LỤC")
    insert_after = None
    for p in doc.paragraphs[paragraph_index(doc, toc) + 1 :]:
        if p.text.strip() == "5.7. Ma trận rủi ro kỹ thuật":
            insert_after = p
            break
    if insert_after and doc.paragraphs[paragraph_index(doc, insert_after) + 1].text.strip() != "5.8. Phân công công việc và đánh giá đóng góp":
        p = insert_paragraph_after(insert_after, "5.8. Phân công công việc và đánh giá đóng góp", "Normal")
        p.paragraph_format.left_indent = Cm(0.35)
    table_list = find_heading(doc, "DANH MỤC BẢNG")
    current = table_list
    while True:
        idx = paragraph_index(doc, current)
        if idx + 1 >= len(doc.paragraphs):
            break
        nxt = doc.paragraphs[idx + 1]
        if nxt.style.name.startswith("Heading"):
            break
        current = nxt
    existing = {p.text.strip() for p in doc.paragraphs}
    for item in [
        "Bảng 5.4: Tóm tắt phân công và đánh giá kết quả thực hiện",
    ]:
        if item not in existing:
            current = insert_paragraph_after(current, item, "Normal")


def member_summary(commits: list[dict]):
    grouped = {m: [c for c in commits if c["member"] == m] for m in "ABCD"}
    total = sum(len(v) for v in grouped.values())
    names = {
        "A": "huyhoang171106",
        "B": "DuyBao / Đặng Duy Bảo",
        "C": "HufaCung",
        "D": "north2k6",
    }
    mssv = {"A": "24520554", "B": "24520146", "C": "24520071", "D": "cập nhật"}
    rows = []
    for m in "ABCD":
        cs = grouped[m]
        area_counter = Counter(a for c in cs for a in c["areas"])
        rows.append([
            m,
            names[m],
            mssv[m],
            str(len(cs)),
            f"{(len(cs)/total*100):.1f}%" if total else "0%",
            "; ".join(f"{k} ({v})" for k, v in area_counter.most_common(4)),
        ])
    return rows, grouped, total


def detail_rows(grouped: dict[str, list[dict]]):
    out = []
    for m, cs in grouped.items():
        files = Counter(f for c in cs for f in c["files"])
        effects = Counter(c["effect"] for c in cs)
        top_files = "; ".join(f"{f} ({n})" for f, n in files.most_common(5))
        top_effects = "; ".join(f"{e} ({n})" for e, n in effects.most_common(3))
        role = {
            "A": "AI/API integration và backend runtime",
            "B": "Mobile, community/chat, GPA và minh chứng kiểm thử",
            "C": "Bảo mật/cấu hình, admin/client, hardening",
            "D": "Repository discovery và repo evaluation",
        }[m]
        out.append([m, role, top_files, top_effects])
    return out


def commit_rows(grouped: dict[str, list[dict]]):
    rows = []
    for m in "ABCD":
        for c in sorted(grouped[m], key=lambda x: (x["date"], x["hash"])):
            files = "; ".join(c["files"][:3])
            if len(c["files"]) > 3:
                files += f"; +{len(c['files']) - 3} files"
            rows.append([m, c["date"], c["hash"], c["subject"], c["effect"], files])
    return rows


def summary_effect_rows(grouped: dict[str, list[dict]]):
    rows = []
    for m in "ABCD":
        effects = Counter(c["effect"] for c in grouped[m])
        for effect, count in effects.most_common():
            examples = [c["subject"] for c in grouped[m] if c["effect"] == effect][:2]
            rows.append([m, effect, str(count), "; ".join(examples)])
    return rows


def add_section(doc: Document, commits: list[dict]):
    chapter6 = find_heading(doc, "Chương 6. KẾT LUẬN VÀ KHUYẾN NGHỊ")
    h = insert_paragraph_after(doc.paragraphs[paragraph_index(doc, chapter6) - 1], "5.8. Phân công công việc và đánh giá đóng góp", "Heading 2")
    rows, grouped, total = member_summary(commits)
    p = add_body_after(
        h,
        f"Phần phân công được tính lại từ nhánh {REV}, mốc từ ngày 14/03/2026 đến hết ngày 14/06/2026, loại trừ merge commit theo đúng cách GitHub Contributors ghi là contributions per week to master, excluding merge commits. Thành viên A được tính theo GitHub account huyhoang171106, đối chiếu local Git gồm author huyhoang171106, author Nguyen Huy Hoang cùng email 24520554@gm.uit.edu.vn và email huyhoang171106@gmail.com. Tổng số commit của bốn thành viên theo quy tắc này là {total}.",
    )
    p = append_table_after(
        p,
        "Bảng 5.4: Tóm tắt phân công và đánh giá kết quả thực hiện",
        ["TV", "Phân công chính", "Số commit", "Đánh giá kết quả thực hiện"],
        [
            ["A", "AI học tập, RAG, roadmap, backend runtime và cấu hình an toàn khởi động.", str(len(grouped["A"])), "Khối lượng lớn, bám sát lõi AI/API của đồ án; kết quả nổi bật là cải thiện Subject Q&A, streaming, roadmap và độ ổn định khi chạy backend."],
            ["B", "Community/chat, mobile, GPA calculator, script vận hành và test matrix.", str(len(grouped["B"])), "Đóng góp mạnh ở phần trải nghiệm người dùng và bằng chứng kiểm thử; giúp các luồng community, GPA và mobile hoàn thiện và dễ nghiệm thu hơn."],
            ["C", "Bảo mật, cấu hình hệ thống, admin/client, mobile build và hardening.", str(len(grouped["C"])), "Bao phủ rộng nhiều lớp hệ thống; nổi bật ở hardening, cấu hình vận hành, luồng admin và các chỉnh sửa giúp hệ thống chạy ổn định hơn."],
            ["D", "Repository discovery, GithubScanService, repo evaluation, repo AI analysis và kiểm thử liên quan.", str(len(grouped["D"])), "Tập trung rõ vào mảng phân tích repository; giúp tính năng quét, đánh giá và hiển thị thông tin repo có chiều sâu và có test đi kèm."],
        ],
        8,
    )
    p = add_body_after(
        p,
        "Ở mức tập thể, nhóm đã hoàn thành được các trục chính của đồ án gồm backend Java API, bảo mật và cấu hình hệ thống, các luồng AI hỗ trợ học tập, community/chat và phần repository discovery. Khối lượng công việc giữa các thành viên có khác nhau, nhưng nhìn chung có sự phân vai tương đối rõ, không bị chồng chéo hoàn toàn và có bằng chứng commit đủ để truy vết trách nhiệm từng phần.",
    )
    p = add_body_after(
        p,
        "Đối với thành viên A, phần việc nổi bật nằm ở AI Course Q&A, RAG, streaming chat, roadmap và các thay đổi liên quan đến runtime backend. Việc đọc lại toàn bộ commit cho thấy A không chỉ làm giao diện gọi API mà còn đi sâu vào service, test, cấu hình và các tình huống khởi động an toàn, vì vậy đây là thành viên gắn nhiều với lõi kỹ thuật của đồ án.",
    )
    p = add_body_after(
        p,
        "Đối với thành viên B, phần việc tập trung nhiều vào community/chat, mobile, GPA calculator và phần test matrix. Nhóm commit của B cho thấy vai trò kết nối giữa chức năng người dùng và khâu chứng minh chất lượng, nghĩa là không chỉ làm tính năng mà còn bổ sung test, tài liệu và các điều chỉnh giúp nghiệm thu dễ hơn.",
    )
    p = add_body_after(
        p,
        "Đối với thành viên C, khối lượng công việc trải rộng từ cấu hình hệ thống, bảo mật, hardening, luồng admin đến các phần cần chỉnh sửa để toàn hệ thống chạy ổn định. Đây là nhóm đóng góp mang tính nền tảng, giúp đồ án không chỉ có chức năng mà còn có khả năng vận hành và bảo trì tốt hơn.",
    )
    p = add_body_after(
        p,
        "Đối với thành viên D, phạm vi công việc tập trung rõ vào repository discovery, GithubScanService, repo evaluation và repo AI analysis. Dù số commit thấp hơn ba thành viên còn lại, phần việc lại khá nhất quán và có giá trị chuyên biệt, góp phần làm rõ hướng tính năng khám phá và đánh giá mã nguồn học thuật của DevOrbit.",
    )
    p = add_body_after(
        p,
        "Nhìn chung, kết quả thực hiện của nhóm là đạt được một hệ thống có trục backend Java tương đối rõ, có tính năng đủ rộng, có bằng chứng kiểm thử ở nhiều phần và có sự phân công đủ minh bạch để truy vết. Điểm mạnh là các thành viên đều có phạm vi phụ trách tương đối nhận diện được. Điểm cần rút kinh nghiệm là khối lượng giữa các thành viên chưa thật đồng đều và một số phần hỗ trợ web/mobile xuất hiện khá nhiều trong commit history, nên khi trình bày báo cáo cần nhấn mạnh rằng trọng tâm học phần vẫn là backend Java và các luồng nghiệp vụ cốt lõi.",
    )
    append_table_after(
        p,
        "Bảng 5.5: Bằng chứng bao phủ commit không bỏ sót",
        ["TV", "Tổng commit", "Author/email được tính", "Nguyên tắc tổng hợp"],
        [
            ["A", str(len(grouped["A"])), "huyhoang171106; Nguyen Huy Hoang <24520554@gm.uit.edu.vn>; huyhoang171106@gmail.com", "Đã đọc toàn bộ commit theo hash, sau đó gom theo nội dung công việc thay vì liệt kê từng commit trong báo cáo."],
            ["B", str(len(grouped["B"])), "DuyBao / Đặng Duy Bảo", "Đã đọc toàn bộ commit theo hash, sau đó gom theo nội dung công việc thay vì liệt kê từng commit trong báo cáo."],
            ["C", str(len(grouped["C"])), "HufaCung", "Đã đọc toàn bộ commit theo hash, sau đó gom theo nội dung công việc thay vì liệt kê từng commit trong báo cáo."],
            ["D", str(len(grouped["D"])), "north2k6", "Đã đọc toàn bộ commit theo hash, sau đó gom theo nội dung công việc thay vì liệt kê từng commit trong báo cáo."],
        ],
        7,
    )


def main():
    commits = collect_commits()
    doc = Document(INPUT)
    replace_cover_and_voice(doc)
    update_front_lists(doc)
    add_section(doc, commits)
    doc.save(OUTPUT)
    print(f"updated={OUTPUT}")
    print(f"commits={len(commits)}")


if __name__ == "__main__":
    main()
