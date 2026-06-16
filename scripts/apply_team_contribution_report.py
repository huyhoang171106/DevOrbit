from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
INPUT = Path(os.environ.get("DEVORBIT_REPORT_DOCX", ROOT / "devorbit-report-v6-java-api-linked.docx"))
OUTPUT = Path(os.environ.get("DEVORBIT_REPORT_OUTPUT", ROOT / "devorbit-report-v6-java-api-team.docx"))


LIGHT_BLUE = "D9EAF7"


TEAM_ROWS = [
    [
        "A",
        "Nguyen Huy Hoang / huyhoang171106 / Ubuntu",
        "24520554",
        "161",
        "39.0%",
        "Backend Java, Subject Q&A/RAG, cấu hình runtime, API tích hợp, đồng bộ client.",
        "Chủ lực tích hợp hệ thống, cải thiện Subject Q&A/RAG, cấu hình API, xử lý vận hành và đồng bộ dữ liệu giữa backend Java với client.",
    ],
    [
        "B",
        "DuyBao / Đặng Duy Bảo",
        "24520146",
        "99",
        "24.0%",
        "Mobile client, community/chat, GPA calculator, docs/test matrix, một số cấu hình backend.",
        "Phụ trách mobile app, community/chat UI/API liên quan, GPA calculator, cập nhật tài liệu kiểm thử và hỗ trợ ổn định luồng chạy.",
    ],
    [
        "C",
        "HufaCung",
        "24520071",
        "105",
        "25.4%",
        "SecurityConfig, application.yaml, route/layout client, mobile build, hardening docs.",
        "Phụ trách bảo mật/cấu hình, UI admin/student, tài liệu hardening và hỗ trợ ổn định kiến trúc.",
    ],
    [
        "D",
        "north2k6",
        "cập nhật",
        "48",
        "11.6%",
        "GithubScanService, repo evaluation, repo analysis UI, test matrix.",
        "Phụ trách cải thiện tìm kiếm repository, phân tích repo, kiểm thử repo evaluation và tài liệu minh chứng.",
    ],
]


DETAILED_ROWS = [
    [
        "A",
        "Tích hợp hệ thống và backend API trọng tâm",
        "SubjectQaService.java; SubjectQaServiceTest.java; application.yaml; SecurityConfig.java; AiChatWidget.tsx; useSubjectQa.ts.",
        "Cải thiện Subject Q&A/RAG, streaming chat, cấu hình runtime, bảo mật và đồng bộ API với client.",
        "Đóng góp lớn nhất sau khi gom alias trên master. Điểm mạnh là bao quát toàn hệ thống và kết nối backend Java với các client. Cần chuẩn hóa Git identity vì có nhiều alias trong lịch sử commit.",
    ],
    [
        "B",
        "Mobile, community/chat, GPA calculator, tài liệu kiểm thử",
        "AcademicRepository.kt; CourseViewModel.kt; ApiService.kt; GpaCalculatorPage.tsx; router.test.tsx; docs/TEST_MATRIX.md.",
        "Hoàn thiện nhiều luồng mobile, community/chat, bookmark/course, GPA calculator và cập nhật test matrix/story cho minh chứng.",
        "Đóng góp ổn định ở các phần client sử dụng API và tài liệu kiểm thử. Phần việc giúp sản phẩm có thêm kênh sử dụng ngoài web và có bằng chứng kiểm thử rõ hơn.",
    ],
    [
        "C",
        "Bảo mật/cấu hình, client routing, hardening docs",
        "SecurityConfig.java; application.yaml; router.tsx; Layout.tsx; schema.sql; package.json; app/build.gradle.kts; docs hardening.",
        "Cập nhật bảo mật, cấu hình ứng dụng, route/layout, build mobile, sửa test WebFlux và bổ sung tài liệu hardening/technical debt.",
        "Đóng góp rộng ở lớp cấu hình và vận hành. Có nhiều dấu vết ở backend Java, client và tài liệu kỹ thuật, phù hợp vai trò ổn định kiến trúc.",
    ],
    [
        "D",
        "Repository discovery và repo evaluation",
        "GithubScanService.java; GithubRepoService.java; RepoAiAnalysisSection.tsx; repoEvaluation.ts; repoAiAnalysis.ts; RepoDetailPage.tsx; repoEvaluation.test.ts.",
        "Cải thiện tìm kiếm repository, alias map, matching theo môn học, phân tích repo, X-ray summary, test repo evaluation và cập nhật test matrix.",
        "Đóng góp tập trung, rõ module. Phần việc làm mạnh hơn năng lực khám phá và đánh giá repository, là một nghiệp vụ quan trọng của DevOrbit.",
    ],
]


def find_paragraph(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f"paragraph not found: {text}")


def find_heading(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip() == text and p.style.name.startswith("Heading"):
            return p
    raise ValueError(f"heading not found: {text}")


def paragraph_index(doc: Document, paragraph) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p._p is paragraph._p:
            return i
    raise ValueError("paragraph not found in document")


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        p.add_run(text)
    if style:
        p.style = style
    return p


def insert_table_after(paragraph, rows: int, cols: int):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Cm(16))
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return table


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(11)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_i == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(3)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10 if len(cell.text) > 80 else 11)
                    if row_i == 0:
                        run.bold = True


def add_body_after(paragraph, text: str):
    p = insert_paragraph_after(paragraph, text, "Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    return p


def append_table_after(paragraph, caption: str, headers: list[str], rows: list[list[str]]):
    cap = insert_paragraph_after(paragraph, caption, "Normal")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].bold = True
    table = insert_table_after(cap, len(rows) + 1, len(headers))
    for c, h in enumerate(headers):
        set_cell_text(table.cell(0, c), h, True)
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            set_cell_text(table.cell(r, c), value)
    style_table(table)
    spacer_xml = OxmlElement("w:p")
    table._tbl.addnext(spacer_xml)
    return Paragraph(spacer_xml, paragraph._parent)


def replace_cover(doc: Document):
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "NGÔN NGỮ LẬP TRÌNH JAVA":
            p.text = "NGÔN NGỮ LẬP TRÌNH JAVA SE330.Q21"
        elif "DEVORBIT:" in t:
            p.text = "DEVORBIT: NỀN TẢNG QUẢN LÝ VÀ KHÁM PHÁ MÃ NGUỒN HỌC THUẬT"

    marker = find_paragraph(doc, "Sinh viên thực hiện:")
    start = paragraph_index(doc, marker) + 1
    end = start
    while end < len(doc.paragraphs) and "TP. HỒ CHÍ MINH" not in doc.paragraphs[end].text:
        end += 1
    for p in reversed(doc.paragraphs[start:end]):
        p._element.getparent().remove(p._element)

    current = marker
    cover_lines = [
        "Thành viên A - 24520554 - Git: Nguyen Huy Hoang / huyhoang171106",
        "Thành viên B - 24520146 - Git: DuyBao / Đặng Duy Bảo",
        "Thành viên C - 24520071 - Git: HufaCung",
        "Thành viên D - MSSV: cập nhật - Git: north2k6",
    ]
    for line in cover_lines:
        current = insert_paragraph_after(current, line, "Normal")
        current.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in current.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(13)

    replacements = {
        "Lời đầu tiên, tác giả xin gửi lời cảm ơn chân thành và sâu sắc nhất đến Ban Giám hiệu, quý Thầy Cô Trường Đại học Công nghệ Thông tin - Đại học Quốc gia TP.HCM (UIT), đặc biệt là các giảng viên thuộc Khoa Kỹ thuật Phần mềm đã tận tình truyền đạt kiến thức chuyên môn trong suốt quá trình học tập.": "Lời đầu tiên, nhóm thực hiện xin gửi lời cảm ơn chân thành đến Ban Giám hiệu, quý Thầy Cô Trường Đại học Công nghệ Thông tin - Đại học Quốc gia TP.HCM (UIT), đặc biệt là các giảng viên phụ trách học phần Ngôn ngữ lập trình Java SE330.Q21 đã định hướng kiến thức và phương pháp xây dựng đồ án.",
        "Mặc dù đã đầu tư nhiều thời gian và công sức nghiên cứu, đồ án chắc chắn không tránh khỏi những điểm thiếu sót do giới hạn về mặt thời gian và kinh nghiệm thực tiễn. Tác giả kính mong nhận được sự nhận xét, đóng góp ý kiến từ quý Thầy Cô để sản phẩm ngày càng hoàn thiện hơn.": "Mặc dù đã đầu tư nhiều thời gian và công sức nghiên cứu, đồ án chắc chắn không tránh khỏi những điểm thiếu sót do giới hạn về thời gian và kinh nghiệm thực tiễn. Nhóm thực hiện kính mong nhận được nhận xét, góp ý từ quý Thầy Cô để sản phẩm ngày càng hoàn thiện hơn.",
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in replacements:
            p.text = replacements[t]
        elif t.startswith("Lời đầu tiên, tác giả xin gửi lời cảm ơn"):
            p.text = "Lời đầu tiên, nhóm thực hiện xin gửi lời cảm ơn chân thành đến Ban Giám hiệu, quý Thầy Cô Trường Đại học Công nghệ Thông tin - Đại học Quốc gia TP.HCM (UIT), đặc biệt là các giảng viên phụ trách học phần Ngôn ngữ lập trình Java SE330.Q21 đã định hướng kiến thức và phương pháp xây dựng đồ án."
        elif t.startswith("Mặc dù đã đầu tư nhiều thời gian") and "Tác giả kính mong" in t:
            p.text = "Mặc dù đã đầu tư nhiều thời gian và công sức nghiên cứu, đồ án chắc chắn không tránh khỏi những điểm thiếu sót do giới hạn về thời gian và kinh nghiệm thực tiễn. Nhóm thực hiện kính mong nhận được nhận xét, góp ý từ quý Thầy Cô để sản phẩm ngày càng hoàn thiện hơn."
        elif t == "Tác giả xin trân trọng cảm ơn!":
            p.text = "Nhóm thực hiện xin trân trọng cảm ơn!"
        elif t == "Dựa trên kết quả phân tích, tác giả đề xuất các khuyến nghị:":
            p.text = "Dựa trên kết quả phân tích, nhóm thực hiện đề xuất các khuyến nghị:"


def update_lists(doc: Document):
    toc = find_heading(doc, "MỤC LỤC")
    insert_after = None
    for p in doc.paragraphs[paragraph_index(doc, toc) + 1 :]:
        if p.text.strip() == "5.7. Ma trận rủi ro kỹ thuật":
            insert_after = p
            break
    if insert_after is not None:
        next_p = doc.paragraphs[paragraph_index(doc, insert_after) + 1]
        if next_p.text.strip() != "5.8. Phân công công việc và đánh giá đóng góp":
            p = insert_paragraph_after(insert_after, "5.8. Phân công công việc và đánh giá đóng góp", "Normal")
            p.paragraph_format.left_indent = Cm(0.35)

    table_list = find_heading(doc, "DANH MỤC BẢNG")
    current = table_list
    while True:
        idx = paragraph_index(doc, current)
        if idx + 1 >= len(doc.paragraphs):
            break
        nxt = doc.paragraphs[idx + 1]
        if nxt.style.name.startswith("Heading"):
            break
        current = nxt
    if "Bảng 5.4: Thống kê đóng góp theo GitHub/Git" not in [p.text.strip() for p in doc.paragraphs]:
        current = insert_paragraph_after(current, "Bảng 5.4: Thống kê đóng góp theo GitHub/Git trên master", "Normal")
        current = insert_paragraph_after(current, "Bảng 5.5: Phân công module và đánh giá cá nhân", "Normal")
        current = insert_paragraph_after(current, "Bảng 5.6: Chi tiết phân công theo module và file tiêu biểu", "Normal")


def add_contribution_section(doc: Document):
    chapter6 = find_heading(doc, "Chương 6. KẾT LUẬN VÀ KHUYẾN NGHỊ")
    h = insert_paragraph_after(doc.paragraphs[paragraph_index(doc, chapter6) - 1], "5.8. Phân công công việc và đánh giá đóng góp", "Heading 2")
    p = add_body_after(
        h,
        "Phần phân công dưới đây được tổng hợp từ lịch sử Git của repository. Do một số thành viên sử dụng nhiều alias Git khác nhau, báo cáo gom theo email/username chính và quy về bốn thành viên A, B, C, D. Số commit chỉ phản ánh dấu vết Git, không thay thế hoàn toàn đánh giá chất lượng, nhưng là cơ sở khách quan để mô tả module đã tham gia.",
    )
    p = append_table_after(
        p,
        "Bảng 5.4: Thống kê đóng góp theo GitHub/Git trên master",
        ["TV", "Git identity", "MSSV", "Commit", "Tỉ lệ", "Khu vực có dấu vết chính"],
        [[row[0], row[1], row[2], row[3], row[4], row[5]] for row in TEAM_ROWS],
    )
    p = add_body_after(
        p,
        "Số liệu trong bảng được tính từ nhánh origin/master, mốc từ ngày 14/03/2026, loại trừ merge commit để bám sát cách GitHub Contributors hiển thị contribution trên nhánh master. Tổng số commit không merge trong phạm vi này là 413. Sau khi gom các alias Git của cùng thành viên A, thành viên A chiếm khoảng 39.0% số commit; thành viên B chiếm khoảng 24.0%; thành viên C chiếm khoảng 25.4%; thành viên D chiếm khoảng 11.6%.",
    )
    p = add_body_after(
        p,
        "Nhìn từ lịch sử commit master-only, nhóm có phân bố công việc theo cả chiều ngang và chiều dọc. Thành viên A tham gia nhiều luồng tích hợp hệ thống và AI/API. Thành viên B nổi bật ở mobile, community/chat, GPA calculator và tài liệu kiểm thử. Thành viên C tham gia mạnh vào cấu hình bảo mật, client, mobile build và tài liệu hardening. Thành viên D tập trung vào phân tích repository, cải thiện tìm kiếm và kiểm thử repo evaluation.",
    )
    p = append_table_after(
        p,
        "Bảng 5.5: Phân công module và đánh giá cá nhân",
        ["TV", "Module/chức năng phụ trách", "Đánh giá đóng góp"],
        [[row[0], row[6], "Hoàn thành tốt phần việc có dấu vết Git; cần tiếp tục chuẩn hóa tên/MSSV trong commit để báo cáo nhóm rõ hơn."] for row in TEAM_ROWS],
    )
    p = add_body_after(
        p,
        "Để mô tả rõ hơn trách nhiệm từng người, bảng sau liệt kê các file hoặc khu vực có dấu vết Git nổi bật. Các file này không có nghĩa là thành viên chỉ làm duy nhất phần đó, nhưng là bằng chứng tốt nhất để nối giữa phân công và lịch sử phát triển thực tế.",
    )
    p = append_table_after(
        p,
        "Bảng 5.6: Chi tiết phân công theo module và file tiêu biểu",
        ["TV", "Vai trò chính", "File/khu vực tiêu biểu", "Kết quả bàn giao", "Nhận xét"],
        DETAILED_ROWS,
    )
    p = add_body_after(
        p,
        "Thành viên A giữ vai trò tích hợp chính, nhất là ở các luồng backend Java có kết nối AI và dữ liệu. Trên origin/master, các alias Nguyen Huy Hoang, huyhoang171106 và Ubuntu được gom lại vì cùng đại diện cho thành viên A trong lịch sử repository. Thành viên này thường xử lý các thay đổi có phạm vi rộng, bao gồm cấu hình runtime, Subject Q&A, API repo/course và đồng bộ kiểu dữ liệu với client.",
    )
    p = add_body_after(
        p,
        "Thành viên B có đóng góp rõ ở mobile client và phần chứng cứ kiểm thử. Các commit liên quan đến AcademicRepository, AuthViewModel, CourseViewModel, ApiService và test matrix cho thấy thành viên này tập trung vào việc đưa dữ liệu từ backend Java ra trải nghiệm sử dụng thực tế. GPA calculator và các story/tài liệu đi kèm cũng giúp báo cáo có thêm phần đánh giá chức năng cụ thể.",
    )
    p = add_body_after(
        p,
        "Thành viên C tham gia nhiều vào phần cấu hình, bảo mật và ổn định dự án. Các thay đổi ở SecurityConfig, application.yaml, schema.sql, router, layout và tài liệu hardening cho thấy vai trò của thành viên này thiên về làm cho hệ thống có cấu trúc rõ, chạy ổn và có tài liệu kỹ thuật đi kèm. Đây là phần quan trọng với đồ án Java vì backend không chỉ cần có API mà còn cần cơ chế bảo vệ và vận hành.",
    )
    p = add_body_after(
        p,
        "Thành viên D tập trung vào nghiệp vụ repository discovery và repo evaluation. Các dấu vết ở GithubScanService, GithubRepoService, repoEvaluation, repoAiAnalysis và RepoDetailPage cho thấy thành viên này xử lý phần tìm kiếm, phân tích và trình bày thông tin repository. Đây là nhóm chức năng sát với mục tiêu ban đầu của DevOrbit: giúp sinh viên tìm được mã nguồn học thuật có liên quan đến môn học.",
    )
    add_body_after(
        p,
        "Về đánh giá tập thể, nhóm hoàn thành được một hệ thống có backend Java làm lõi, có dữ liệu học phần, repository, xác thực, phân quyền, AI service và các client sử dụng API. Bài học kinh nghiệm quan trọng là cần thống nhất Git identity từ đầu, tách nhánh theo module rõ hơn và ghi commit message có liên hệ trực tiếp tới chức năng hoặc story để việc truy vết cuối kỳ chính xác hơn.",
    )


def main():
    doc = Document(INPUT)
    replace_cover(doc)
    update_lists(doc)
    add_contribution_section(doc)
    doc.save(OUTPUT)
    print(f"updated={OUTPUT}")


if __name__ == "__main__":
    main()
