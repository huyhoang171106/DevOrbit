from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX = Path(os.environ.get("DEVORBIT_REPORT_DOCX", ROOT / "devorbit-report-v6-clean.docx"))
OUTPUT = Path(os.environ.get("DEVORBIT_REPORT_OUTPUT", DOCX))


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def run_style(run, *, hyperlink: bool = False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if hyperlink:
        run.font.color.rgb = RGBColor(5, 99, 193)
        run.font.underline = True


def add_internal_hyperlink(paragraph, text: str, anchor: str):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    txt = OxmlElement("w:t")
    txt.text = text
    run.append(r_pr)
    run.append(txt)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_external_hyperlink(paragraph, text: str, url: str):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    txt = OxmlElement("w:t")
    txt.text = text
    run.append(r_pr)
    run.append(txt)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, name: str, bookmark_id: int):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def rewrite_with_internal_citations(paragraph, valid_refs: set[int]):
    text = paragraph.text
    matches = list(re.finditer(r"\[(\d+)\]", text))
    if not matches:
        return 0
    clear_paragraph(paragraph)
    pos = 0
    linked = 0
    for match in matches:
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run_style(run)
        ref_num = int(match.group(1))
        cite_text = match.group(0)
        if ref_num in valid_refs:
            add_internal_hyperlink(paragraph, cite_text, f"ref_{ref_num}")
            linked += 1
        else:
            run = paragraph.add_run(cite_text)
            run_style(run)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run_style(run)
    return linked


def rewrite_reference_url(paragraph):
    text = paragraph.text
    match = re.search(r"https?://\S+", text)
    if not match:
        return 0
    before, url, after = text[: match.start()], match.group(0), text[match.end() :]
    clear_paragraph(paragraph)
    if before:
        run = paragraph.add_run(before)
        run_style(run)
    add_external_hyperlink(paragraph, url, url)
    if after:
        run = paragraph.add_run(after)
        run_style(run)
    return 1


def main():
    doc = Document(DOCX)
    refs: dict[int, int] = {}
    in_refs = False
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "TÀI LIỆU THAM KHẢO":
            in_refs = True
            continue
        if in_refs:
            match = re.match(r"^\[(\d+)\]", text)
            if match:
                refs[int(match.group(1))] = idx

    linked_citations = 0
    linked_urls = 0
    in_refs = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "TÀI LIỆU THAM KHẢO":
            in_refs = True
            continue
        if in_refs:
            linked_urls += rewrite_reference_url(paragraph)
        else:
            linked_citations += rewrite_with_internal_citations(paragraph, set(refs))

    # Add bookmarks after paragraph rewrites so clear_paragraph() cannot remove them.
    for bookmark_id, (ref_num, paragraph_idx) in enumerate(sorted(refs.items()), start=1):
        add_bookmark(doc.paragraphs[paragraph_idx], f"ref_{ref_num}", bookmark_id)

    doc.save(OUTPUT)
    print(f"updated={OUTPUT}")
    print(f"bookmarks={len(refs)} linked_citations={linked_citations} linked_urls={linked_urls}")


if __name__ == "__main__":
    main()
