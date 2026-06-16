from __future__ import annotations

import shutil
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX = Path(os.environ.get("DEVORBIT_REPORT_DOCX", ROOT / "devorbit-report-v6.docx"))
BACKUP = Path(os.environ.get("DEVORBIT_REPORT_BACKUP", ROOT / "devorbit-report-v6.before-polish.docx"))


BLUE = RGBColor(31, 78, 121)
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"


def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p


def find_paragraph(doc: Document, text: str):
    for p in iter_paragraphs(doc):
        if p.text.strip() == text:
            return p
    raise ValueError(f"paragraph not found: {text}")


def find_heading(doc: Document, text: str):
    for p in iter_paragraphs(doc):
        if p.text.strip() == text and p.style and p.style.name.startswith("Heading"):
            return p
    raise ValueError(f"heading not found: {text}")


def paragraph_index(doc: Document, paragraph) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p._p is paragraph._p:
            return i
    raise ValueError("paragraph not in document")


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        p.add_run(text)
    if style:
        p.style = style
    return p


def insert_table_after(paragraph, rows: int, cols: int):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Cm(16))
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return table


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)
    return cell


def style_table(table, header_fill: str = LIGHT_BLUE):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(3)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(11)
            if row_i == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)


def add_heading_after(paragraph, text: str, level: int = 2):
    return insert_paragraph_after(paragraph, text, f"Heading {level}")


def add_body_after(paragraph, text: str):
    p = insert_paragraph_after(paragraph, text, "Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets_after(paragraph, items: list[str]):
    current = paragraph
    for item in items:
        current = insert_paragraph_after(current, item, "Normal")
        current.paragraph_format.left_indent = Cm(0.75)
        current.paragraph_format.first_line_indent = Cm(0)
        current.paragraph_format.space_after = Pt(3)
    return current


def replace_section_body(doc: Document, heading_text: str, paragraphs: list[str]):
    heading = find_heading(doc, heading_text)
    start = paragraph_index(doc, heading) + 1
    end = len(doc.paragraphs)
    for i in range(start, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if p.text.strip() and p.style and p.style.name.startswith("Heading"):
            end = i
            break
    for p in reversed(doc.paragraphs[start:end]):
        p._element.getparent().remove(p._element)
    current = heading
    for text in paragraphs:
        current = add_body_after(current, text)
    return current


def append_table_after(paragraph, caption: str, headers: list[str], rows: list[list[str]]):
    cap = insert_paragraph_after(paragraph, caption, "Normal")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].bold = True
    table = insert_table_after(cap, len(rows) + 1, len(headers))
    for col, header in enumerate(headers):
        set_cell_text(table.cell(0, col), header, True)
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            set_cell_text(table.cell(r, c), value)
    style_table(table)
    spacer_xml = OxmlElement("w:p")
    table._tbl.addnext(spacer_xml)
    return Paragraph(spacer_xml, paragraph._parent)


def set_document_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    for name, size in [("Heading 1", 15), ("Heading 2", 13)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.keep_with_next = True
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        footer = section.footer.paragraphs[0]
        footer.text = "DevOrbit - Báo cáo đồ án cuối kỳ"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)


def update_lists(doc: Document):
    toc = find_heading(doc, "MỤC LỤC")
    next_h = None
    for p in doc.paragraphs[paragraph_index(doc, toc) + 1 :]:
        if p.text.strip() and p.style.name.startswith("Heading"):
            next_h = p
            break
    if next_h is None:
        return
    start = paragraph_index(doc, toc) + 1
    end = paragraph_index(doc, next_h)
    for p in reversed(doc.paragraphs[start:end]):
        p._element.getparent().remove(p._element)
    entries = [
        "Lời cảm ơn",
        "Tóm tắt điều hành",
        "Mục lục",
        "Danh mục hình",
        "Danh mục bảng",
        "Danh mục từ viết tắt",
        "Chương 1. Giới thiệu chung và phân tích bối cảnh",
        "1.1. Bối cảnh và lý do lựa chọn đề tài",
        "1.2. Mục tiêu đề tài",
        "1.3. Phạm vi nghiên cứu",
        "1.4. Tiêu chí thành công và phương pháp đánh giá",
        "1.5. Cấu trúc báo cáo",
        "Chương 2. Phương pháp khảo sát và tổng quan kiến trúc",
        "2.1. Phương pháp nghiên cứu",
        "2.2. Tổng quan kiến trúc hệ thống",
        "2.3. Luồng nghiệp vụ tổng thể",
        "2.4. Ma trận truy vết yêu cầu",
        "Chương 3. Phân tích chi tiết Backend (devorbit-api)",
        "3.1-3.8. Kiến trúc, dữ liệu, AI, scan GitHub, chat real-time",
        "3.9. Bảo mật, vận hành và kiểm soát lỗi backend",
        "Chương 4. Lớp giao diện minh họa và tích hợp API (devorbit-web)",
        "4.1-4.6. Cấu trúc giao diện, routing, state, Photobooth, hiệu năng",
        "4.7. Vai trò của devorbit-web trong đồ án Java",
        "Chương 5. Đánh giá, thảo luận và hình ảnh hệ thống",
        "5.1-5.5. Kết quả, giao diện, thuận lợi, hạn chế, hướng phát triển",
        "5.6. Kế hoạch kiểm thử và minh chứng vận hành",
        "5.7. Ma trận rủi ro kỹ thuật",
        "Chương 6. Kết luận và khuyến nghị",
        "6.1. Kết luận",
        "6.2. Khuyến nghị",
        "6.3. Đóng góp chính của đồ án",
        "Tài liệu tham khảo",
    ]
    current = toc
    for entry in entries:
        current = insert_paragraph_after(current, entry, "Normal")
        current.paragraph_format.left_indent = Cm(0.35 if entry[:1].isdigit() else 0)

    fig = find_heading(doc, "DANH MỤC HÌNH")
    tbl = find_heading(doc, "DANH MỤC BẢNG")
    fig_idx = paragraph_index(doc, fig)
    tbl_idx = paragraph_index(doc, tbl)
    for p in reversed(doc.paragraphs[fig_idx + 1 : tbl_idx]):
        p._element.getparent().remove(p._element)
    current = fig
    for entry in [
        "Hình 3.1: Giao diện Trang chủ DevOrbit",
        "Hình 3.2: Danh sách môn học",
        "Hình 3.3: Lộ trình học tập cá nhân hóa",
        "Hình 3.4: Giao diện bản đồ học phần",
        "Hình 3.5: Photobooth UIT",
        "Hình 3.6: Quản trị repository",
        "Hình 3.7: Giao diện quản trị hệ thống",
    ]:
        current = insert_paragraph_after(current, entry, "Normal")

    next_h = None
    for p in doc.paragraphs[paragraph_index(doc, tbl) + 1 :]:
        if p.text.strip() and p.style.name.startswith("Heading"):
            next_h = p
            break
    start = paragraph_index(doc, tbl) + 1
    end = paragraph_index(doc, next_h)
    for p in reversed(doc.paragraphs[start:end]):
        p._element.getparent().remove(p._element)
    current = tbl
    for entry in [
        "Bảng 1.1: Tiêu chí thành công của đồ án",
        "Bảng 2.1: Danh sách các công nghệ sử dụng",
        "Bảng 2.2: Cấu trúc bảng dữ liệu PostgreSQL",
        "Bảng 2.3: Luồng nghiệp vụ và tác nhân sử dụng",
        "Bảng 2.4: Ma trận truy vết yêu cầu - triển khai",
        "Bảng 3.1: Danh sách API endpoints chính",
        "Bảng 3.2: Nhóm bảng dữ liệu PostgreSQL",
        "Bảng 3.3: Kiểm soát bảo mật và vận hành backend",
        "Bảng 4.1: Cấu trúc component giao diện",
        "Bảng 4.2: Vai trò minh họa của devorbit-web",
        "Bảng 5.1: Chỉ số hoàn thành hệ thống",
        "Bảng 5.2: Kế hoạch kiểm thử và minh chứng",
        "Bảng 5.3: Ma trận rủi ro kỹ thuật",
    ]:
        current = insert_paragraph_after(current, entry, "Normal")


def enhance_content(doc: Document):
    replace_section_body(
        doc,
        "TÓM TẮT ĐIỀU HÀNH",
        [
            "DevOrbit là nền tảng quản lý và khám phá mã nguồn học thuật dành cho sinh viên UIT, trong đó trọng tâm kỹ thuật của đồ án là backend Java 21/Spring Boot. Phân hệ devorbit-api chịu trách nhiệm chuẩn hóa dữ liệu học phần, quản lý repository GitHub, xác thực JWT, phân quyền, xử lý nghiệp vụ sinh viên/quản trị viên, tích hợp PostgreSQL/Supabase và cung cấp API cho các lớp giao diện.",
            "Điểm nổi bật của đồ án nằm ở cách backend Java đóng vai trò lõi vận hành: dữ liệu học phần được quản lý bằng migration, repository được thu thập và kiểm duyệt qua pipeline admin, sinh viên nhận gợi ý lộ trình học và câu trả lời AI theo ngữ cảnh môn học. Lớp web chỉ được trình bày như một client minh họa cho khả năng tích hợp API, không phải trọng tâm đánh giá của đồ án.",
            "Báo cáo này được hoàn thiện theo hướng báo cáo kỹ thuật có thể bảo vệ: nêu rõ bối cảnh, mục tiêu, phạm vi, phương pháp khảo sát mã nguồn, kiến trúc backend, mô hình dữ liệu, API, bảo mật, kiểm thử, vận hành, hạn chế, rủi ro và lộ trình phát triển. Các nhận định trong báo cáo ưu tiên bám vào devorbit-api, tài liệu kiến trúc, test matrix và trạng thái triển khai hiện có.",
            "Kết quả cho thấy DevOrbit đã đạt mức sản phẩm học thuật hoàn chỉnh ở các luồng cốt lõi phía backend: public browsing, admin pipeline, student auth/bookmark, AI roadmap, AI subject Q&A, photobooth frame API và hạ tầng Docker/Supabase. Các hạn chế còn lại tập trung vào mở rộng kiểm thử backend, kiểm soát chi phí AI, hardening vận hành và hoàn thiện client sử dụng API.",
        ],
    )

    replace_section_body(
        doc,
        "1.4. Cấu trúc báo cáo",
        [
            "Báo cáo gồm 6 chương. Chương 1 trình bày bối cảnh, mục tiêu, phạm vi, tiêu chí đánh giá và cấu trúc báo cáo. Chương 2 mô tả phương pháp khảo sát mã nguồn, kiến trúc tổng thể, luồng nghiệp vụ và ma trận truy vết yêu cầu. Chương 3 phân tích chuyên sâu backend devorbit-api. Chương 4 trình bày ngắn gọn devorbit-web như client sử dụng API Java. Chương 5 đánh giá kết quả, kiểm thử, hạn chế và rủi ro kỹ thuật. Chương 6 tổng kết đóng góp, khuyến nghị và hướng phát triển.",
        ],
    )
    h = find_heading(doc, "1.4. Cấu trúc báo cáo")
    h.clear()
    h.add_run("1.5. Cấu trúc báo cáo")
    previous = doc.paragraphs[paragraph_index(doc, h) - 1]
    new_h = add_heading_after(previous, "1.4. Tiêu chí thành công và phương pháp đánh giá", 2)
    p = add_body_after(
        new_h,
        "Để tránh đánh giá cảm tính, báo cáo sử dụng nhóm tiêu chí thành công theo bốn lớp: giá trị người dùng, độ đầy đủ chức năng, chất lượng kiến trúc và khả năng vận hành. Mỗi tiêu chí được liên kết với bằng chứng quan sát được trong mã nguồn, tài liệu hoặc luồng chạy hệ thống.",
    )
    p = append_table_after(
        p,
        "Bảng 1.1: Tiêu chí thành công của đồ án",
        ["Nhóm tiêu chí", "Yêu cầu", "Bằng chứng đánh giá"],
        [
            ["Giá trị người dùng", "Sinh viên tìm được môn học, repo mẫu, lộ trình và hỗ trợ AI theo ngữ cảnh.", "Public API, Course Detail, AI Tutor, AI Roadmap."],
            ["Độ đầy đủ chức năng", "Có luồng public, student và admin tách bạch; dữ liệu repository đi qua kiểm duyệt.", "README, route structure, controller/service packages."],
            ["Chất lượng kiến trúc", "Backend 3 lớp, DTO/entity tách biệt, service xử lý nghiệp vụ rõ ràng.", "docs/ARCHITECTURE.md, package map, controller/service/repository."],
            ["Bảo mật và vận hành", "JWT, phân quyền, migration, cache, exception handling, WebSocket auth.", "Security filter, Flyway migrations, ApiExceptionHandler, Test Matrix."],
            ["Khả năng mở rộng", "Có hướng mở rộng AI, mobile client, integration test và kiểm soát chi phí vận hành.", "Chương 5.5-5.7 và docs/TEST_MATRIX.md."],
        ],
    )

    after_arch = find_heading(doc, "2.2. Tổng quan kiến trúc hệ thống")
    # Find the last paragraph before chapter 3.
    idx = paragraph_index(doc, after_arch)
    end = idx + 1
    for i in range(idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip() == "Chương 3. PHÂN TÍCH CHI TIẾT BACKEND (devorbit-api)":
            end = i
            break
    anchor = doc.paragraphs[end - 1]
    h23 = add_heading_after(anchor, "2.3. Luồng nghiệp vụ tổng thể", 2)
    p = add_body_after(
        h23,
        "DevOrbit được thiết kế quanh ba tác nhân chính: khách truy cập, sinh viên đã đăng nhập và quản trị viên. Cách phân vai này giúp hệ thống vừa mở đủ để sinh viên khám phá tài nguyên học tập, vừa kiểm soát chất lượng dữ liệu repository trước khi công khai.",
    )
    p = append_table_after(
        p,
        "Bảng 2.3: Luồng nghiệp vụ và tác nhân sử dụng",
        ["Tác nhân", "Luồng chính", "Giá trị tạo ra"],
        [
            ["Khách truy cập", "Xem danh sách môn học, chi tiết môn và repository công khai.", "Tìm tài nguyên học tập mà không cần tài khoản."],
            ["Sinh viên", "Đăng ký OTP, đăng nhập JWT, bookmark môn/repo, hỏi AI, tạo roadmap.", "Cá nhân hóa việc học và lưu lại tài nguyên quan trọng."],
            ["Quản trị viên", "CRUD môn học, scan GitHub, duyệt repo candidate, quản lý tài nguyên và frame photobooth.", "Đảm bảo dữ liệu công khai có kiểm duyệt và có thể vận hành dài hạn."],
            ["Hệ thống nền", "Migration, cache, notification event, WebSocket broker, AI provider fallback.", "Giữ dữ liệu nhất quán, phản hồi nhanh và giảm rủi ro phụ thuộc provider."],
        ],
    )
    h24 = add_heading_after(p, "2.4. Ma trận truy vết yêu cầu", 2)
    p = add_body_after(
        h24,
        "Ma trận truy vết giúp hội đồng nhìn thấy mối liên hệ giữa mục tiêu đề tài, thành phần triển khai và minh chứng đánh giá. Đây là lớp kiểm soát quan trọng để báo cáo không chỉ mô tả công nghệ mà còn chứng minh mục tiêu đã được hiện thực.",
    )
    append_table_after(
        p,
        "Bảng 2.4: Ma trận truy vết yêu cầu - triển khai",
        ["Mục tiêu", "Thành phần triển khai", "Minh chứng trong báo cáo"],
        [
            ["Chuẩn hóa dữ liệu học phần", "PostgreSQL, Flyway migrations, Course service/controller.", "Chương 3.4, Bảng 3.2."],
            ["Khám phá repository học thuật", "GitHub scan, repo candidate, admin review, public repo detail.", "Chương 3.6, Chương 4.2."],
            ["Hỗ trợ học tập bằng AI", "RAG subject Q&A, roadmap generator, embedding service fallback.", "Chương 3.5, Chương 5.1."],
            ["Tra cứu tri thức học phần", "API course relationship và dữ liệu tiên quyết/bổ trợ.", "Chương 3.4, Chương 4.4."],
            ["Vận hành an toàn", "JWT, role-based access, exception handling, cache eviction.", "Chương 3.3, Chương 3.9."],
        ],
    )

    anchor = find_heading(doc, "3.8. Caching, Notifications và Exception Handling")
    idx = paragraph_index(doc, anchor)
    end = idx + 1
    for i in range(idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip() == "Chương 4. PHÂN TÍCH CHI TIẾT FRONTEND (devorbit-web)":
            end = i
            break
    h39 = add_heading_after(doc.paragraphs[end - 1], "3.9. Bảo mật, vận hành và kiểm soát lỗi backend", 2)
    p = add_body_after(
        h39,
        "Backend của DevOrbit có nhiều điểm kiểm soát phù hợp với một hệ thống học thuật có dữ liệu người dùng và quyền quản trị. Trọng tâm bảo mật nằm ở JWT, phân quyền route, tách DTO khỏi entity, giới hạn đầu vào ở service và cơ chế xử lý exception thống nhất. Trọng tâm vận hành nằm ở Flyway migration, cache có TTL, event notification và fallback cho các provider AI.",
    )
    append_table_after(
        p,
        "Bảng 3.3: Kiểm soát bảo mật và vận hành backend",
        ["Khu vực", "Cách triển khai", "Ý nghĩa kỹ thuật"],
        [
            ["Xác thực", "JwtAuthenticationFilter kiểm tra Bearer token và nạp SecurityContext.", "Tách xác thực khỏi controller, giảm lặp logic bảo mật."],
            ["Phân quyền", "Public/Student/Admin endpoint được phân lớp theo security config.", "Giới hạn quyền ghi dữ liệu quan trọng cho admin."],
            ["Dữ liệu", "JPA entity, repository và migration tách khỏi DTO/API contract.", "Giảm rủi ro lộ cấu trúc nội bộ và hỗ trợ tiến hóa schema."],
            ["AI provider", "EmbeddingService có implementation production và fallback.", "Hạn chế sự cố khi API ngoài rate limit hoặc tạm ngừng."],
            ["Lỗi hệ thống", "ApiExceptionHandler chuẩn hóa 400/401/404/500.", "Client nhận được response dự đoán được thay vì stack trace."],
            ["Hiệu năng", "Caffeine cache với TTL và eviction khi dữ liệu thay đổi.", "Tăng tốc truy vấn đọc nhiều nhưng vẫn giữ nhất quán cơ bản."],
        ],
    )

    anchor = find_heading(doc, "4.6. Cấu hình Vite và tối ưu hiệu năng")
    idx = paragraph_index(doc, anchor)
    end = idx + 1
    for i in range(idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip() == "Chương 5. ĐÁNH GIÁ, THẢO LUẬN VÀ HÌNH ẢNH HỆ THỐNG":
            end = i
            break
    h47 = add_heading_after(doc.paragraphs[end - 1], "4.7. Vai trò của devorbit-web trong đồ án Java", 2)
    p = add_body_after(
        h47,
        "Trong phạm vi báo cáo đồ án Java, devorbit-web được xem là client minh họa cho API thay vì đối tượng phân tích chính. Phần này chỉ làm rõ cách giao diện tiêu thụ endpoint từ devorbit-api, cách quản lý token, cách gọi các luồng course, repo, AI và photobooth. Trọng tâm đánh giá vẫn là thiết kế backend, mô hình dữ liệu, service nghiệp vụ, bảo mật và kiểm thử phía Java.",
    )
    append_table_after(
        p,
        "Bảng 4.2: Vai trò minh họa của devorbit-web",
        ["Khu vực giao diện", "API Java liên quan", "Ý nghĩa trong báo cáo"],
        [
            ["Danh sách môn học", "CourseController, CourseService, CourseRepository.", "Chứng minh API public và mapping dữ liệu học phần."],
            ["Chi tiết repository", "RepoController, GithubScanService, RepoCandidateService.", "Chứng minh pipeline thu thập và kiểm duyệt mã nguồn."],
            ["AI Tutor/Roadmap", "SubjectQaController, RoadmapGenerator, EmbeddingService.", "Chứng minh backend tích hợp AI theo nghiệp vụ học tập."],
            ["Đăng nhập sinh viên", "StudentAuthController, JwtService, SecurityConfig.", "Chứng minh xác thực JWT và phân quyền client-server."],
            ["Admin review", "Admin controllers, service nghiệp vụ, repository.", "Chứng minh luồng quản trị dữ liệu và kiểm duyệt."],
        ],
    )

    anchor = find_heading(doc, "5.5. Hướng phát triển")
    idx = paragraph_index(doc, anchor)
    end = idx + 1
    for i in range(idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip() == "Chương 6. KẾT LUẬN VÀ KHUYẾN NGHỊ":
            end = i
            break
    h56 = add_heading_after(doc.paragraphs[end - 1], "5.6. Kế hoạch kiểm thử và minh chứng vận hành", 2)
    p = add_body_after(
        h56,
        "Báo cáo phân biệt ba mức minh chứng: kiểm tra tĩnh từ mã nguồn, kiểm thử tự động và smoke test vận hành. Với phạm vi đồ án Java, ưu tiên kiểm thử nằm ở backend: controller, service, repository, security filter, migration, AI service fallback và các endpoint quan trọng. Kiểm thử UI chỉ đóng vai trò phụ để xác nhận client gọi API đúng.",
    )
    p = append_table_after(
        p,
        "Bảng 5.2: Kế hoạch kiểm thử và minh chứng",
        ["Hạng mục", "Minh chứng hiện có", "Bổ sung nên thực hiện"],
        [
            ["Backend API", "Maven tests, service/controller focused tests, schema initializer tests trong Test Matrix.", "Thêm Testcontainers cho luồng dữ liệu xuyên tầng."],
            ["Client Web", "TypeScript/Vitest đã được dùng cho một số hook/component theo Test Matrix.", "Bổ sung smoke test client ở các luồng gọi API quan trọng."],
            ["Mobile", "Gradle unit test cho ViewModel và màn hình course/bookmark.", "Thêm instrumentation test trên thiết bị thật."],
            ["AI/RAG", "Focused tests và runtime smoke cho subject Q&A/roadmap.", "Đo latency, chi phí/token, fallback provider và chất lượng câu trả lời."],
            ["Tích hợp API", "Smoke test thủ công qua trình duyệt.", "Tự động hóa luồng login, course, repo review, AI chat bằng Playwright khi cần."],
        ],
    )
    h57 = add_heading_after(p, "5.7. Ma trận rủi ro kỹ thuật", 2)
    p = add_body_after(
        h57,
        "Các rủi ro dưới đây không phủ nhận kết quả đạt được; chúng giúp xác định đúng phần cần kiểm soát khi DevOrbit được triển khai rộng hơn trong môi trường thật.",
    )
    append_table_after(
        p,
        "Bảng 5.3: Ma trận rủi ro kỹ thuật",
        ["Rủi ro", "Mức ảnh hưởng", "Biện pháp giảm thiểu"],
        [
            ["GitHub API rate limit khi quét nhiều tài khoản", "Cao", "Queue scan, cache kết quả, backoff, token rotation hợp lệ."],
            ["Chi phí AI tăng theo lượng truy vấn", "Cao", "Cache câu hỏi phổ biến, giới hạn quota, fallback model rẻ hơn."],
            ["Luồng client phụ thuộc API backend", "Trung bình", "Chuẩn hóa DTO, thông báo lỗi rõ ràng, contract test cho endpoint quan trọng."],
            ["Dữ liệu học phần thay đổi theo năm học", "Trung bình", "Versioning syllabus, migration có idempotency, admin import workflow."],
            ["Thiếu E2E test cho luồng dài", "Trung bình", "Playwright smoke theo critical path trước mỗi release."],
            ["Phụ thuộc Supabase/provider ngoài", "Trung bình", "Health check, retry có giới hạn, backup/export định kỳ."],
        ],
    )

    concl = find_heading(doc, "6.2. Khuyến nghị")
    idx = paragraph_index(doc, concl)
    end = idx + 1
    for i in range(idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip() == "TÀI LIỆU THAM KHẢO":
            end = i
            break
    h63 = add_heading_after(doc.paragraphs[end - 1], "6.3. Đóng góp chính của đồ án", 2)
    p = add_body_after(
        h63,
        "Đồ án đóng góp ở ba lớp. Ở lớp sản phẩm, DevOrbit cung cấp một nền tảng tập trung để sinh viên UIT khám phá học phần, repository mẫu, quan hệ tiên quyết và hỗ trợ AI. Ở lớp kỹ thuật, đồ án chứng minh khả năng xây dựng một backend Java/Spring Boot có phân quyền, xử lý dữ liệu PostgreSQL, tích hợp GitHub API, WebSocket và AI/RAG trong một kiến trúc rõ ràng. Ở lớp học thuật, báo cáo cung cấp một cách tiếp cận có thể tái sử dụng để đánh giá sản phẩm phần mềm: bắt đầu từ mục tiêu, truy vết tới triển khai, rồi kết thúc bằng minh chứng, hạn chế và rủi ro.",
    )
    add_body_after(
        p,
        "Với phạm vi môn học Java, giá trị lớn nhất của DevOrbit là cho thấy một backend Java có thể trở thành lõi vận hành của hệ sinh thái học tập đa nền tảng. Các phần AI, web client và mobile client được trình bày như các kênh tiêu thụ API, qua đó làm rõ năng lực tích hợp, bảo mật, xử lý dữ liệu và cung cấp API ổn định của devorbit-api.",
    )


def polish_existing_paragraphs(doc: Document):
    replacements = {
        "Báo cáo đưa ra các khuyến nghị về tối ưu hiệu năng 3D, mở rộng RAG pipeline, và triển khai ứng dụng di động.": "Báo cáo đồng thời chỉ ra các ưu tiên cải tiến: mở rộng RAG pipeline, tăng độ phủ kiểm thử backend, kiểm soát chi phí AI và hoàn thiện các client sử dụng API.",
        "(4) Kết nối cộng đồng học tập qua WebSocket real-time chat, giúp sinh viên trao đổi kinh nghiệm học tập và repository hữu ích.": "(4) Kết nối cộng đồng học tập qua WebSocket real-time chat, giúp sinh viên trao đổi kinh nghiệm học tập, chia sẻ repository hữu ích và nhận thông báo đúng ngữ cảnh.",
        "Báo cáo phân tích toàn diện mã nguồn của hai phân hệ backend và frontend.": "Báo cáo phân tích chuyên sâu mã nguồn backend Java và trình bày ngắn gọn lớp client sử dụng API.",
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in replacements:
            p.text = replacements[t]
        elif t == "Chương 4. PHÂN TÍCH CHI TIẾT FRONTEND (devorbit-web)":
            p.text = "Chương 4. LỚP GIAO DIỆN MINH HỌA VÀ TÍCH HỢP API"
        elif t == "4.4. Mô hình 3D Galaxy (Three.js + React Three Fiber)":
            p.text = "4.4. Client minh họa luồng học phần và repository"

    replace_section_body(
        doc,
        "4.4. Client minh họa luồng học phần và repository",
        [
            "Phần giao diện liên quan đến học phần và repository được sử dụng để chứng minh khả năng cung cấp API ổn định từ devorbit-api. Client lấy dữ liệu môn học, quan hệ tiên quyết, repository đã kiểm duyệt và trạng thái người dùng thông qua các endpoint REST. Vì đây là báo cáo đồ án Java, phần phân tích tập trung vào contract API, DTO, service xử lý nghiệp vụ và tính nhất quán dữ liệu phía backend.",
            "Các màn hình client giúp kiểm tra nhanh kết quả tích hợp, nhưng không được xem là đóng góp chính của đồ án. Đóng góp chính vẫn nằm ở backend Java: controller nhận request, service áp dụng nghiệp vụ, repository truy vấn PostgreSQL và security layer bảo vệ các thao tác cần quyền sinh viên hoặc quản trị viên.",
        ],
    )

    text_replacements = {
        "—": "-",
        "3D Galaxy": "bản đồ học phần",
        "Galaxy 3D": "bản đồ học phần",
        "2D/3D Knowledge Graph": "bản đồ quan hệ học phần",
        "Knowledge Graph 2D/3D": "bản đồ quan hệ học phần",
        "knowledge graph": "bản đồ quan hệ học phần",
        "Knowledge Graph": "bản đồ quan hệ học phần",
        "Three.js + React Three Fiber": "client trực quan",
        "Three.js/React Three Fiber": "client trực quan",
        "Three.js/R3F": "client web",
        "Three.js": "client web",
        "React Three Fiber": "client trực quan",
        "2D/3D": "quan hệ học phần",
        "render 3D": "hiển thị dữ liệu",
        "trực quan hóa tri thức quan hệ học phần": "tra cứu quan hệ học phần",
        "tích hợp trực quan hóa tri thức quan hệ học phần": "tích hợp tra cứu quan hệ học phần",
        "Thiết kế sơ đồ tri thức 2D (Blueprint Grid) và mô hình bản đồ học phần (client web) trực quan hóa mạng lưới liên kết môn học theo 8 học kỳ.": "Thiết kế API quan hệ học phần và dữ liệu tiên quyết/bổ trợ để client có thể hiển thị mạng lưới liên kết môn học theo 8 học kỳ.",
        "useGalaxyStore (selectedPlanet, cameraPosition, timeSliderValue).": "useCourseList và các hook liên quan đến dữ liệu học phần.",
        "scenes/ cho client web scene components.": "lib/ cho API clients và utilities.",
        "WebGL": "giao diện web",
        "frontend 3D": "web client",
        "Frontend": "Client web",
        "frontend": "client web",
    }
    for p in doc.paragraphs:
        if not p.text:
            continue
        updated = p.text
        for old, new in text_replacements.items():
            updated = updated.replace(old, new)
        if updated != p.text:
            p.text = updated
    for p in doc.paragraphs:
        if p.text.strip() and p.style.name == "Normal":
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.space_after = Pt(5)
            for run in p.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(13)


def add_page_breaks_before_chapters(doc: Document):
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("Chương ") or text in {"LỜI CẢM ƠN", "TÓM TẮT ĐIỀU HÀNH", "MỤC LỤC", "TÀI LIỆU THAM KHẢO"}:
            if p.runs:
                p.runs[0].add_break(WD_BREAK.PAGE)


def main():
    if not DOCX.exists():
        raise SystemExit(f"Missing {DOCX}")
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
    doc = Document(DOCX)
    set_document_styles(doc)
    enhance_content(doc)
    update_lists(doc)
    polish_existing_paragraphs(doc)
    for table in doc.tables:
        style_table(table, LIGHT_BLUE)
    doc.save(DOCX)
    print(f"updated={DOCX}")
    print(f"backup={BACKUP}")


if __name__ == "__main__":
    main()
