from __future__ import annotations

import os
import subprocess
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = Path(os.environ.get("DEVORBIT_COURSE_REPORT_OUTPUT", ROOT / "devorbit-report-v9-final.docx"))
REV = "origin/master"
SINCE = "2026-03-14T00:00:00"
UNTIL = "2026-06-15T00:00:00"


def set_run_font(run, size: int = 13, bold: bool = False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc: Document, text: str = "", *, first_line: bool = True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.75)
    if text:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 14 if level == 1 else 13, True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_center(doc: Document, text: str, size: int = 13, bold: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size, bold)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], font_size: int = 11):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, header in enumerate(headers):
        cell = table.cell(0, c)
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, font_size, True)
        shade(cell)
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(value))
            set_run_font(run, font_size)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    return table


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, 12, True)
    return p


def add_image(doc: Document, image_path: str, caption: str, width_cm: float = 14.0):
    """Add an image with caption centered below it."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    add_caption(doc, caption)
    return p


def add_mermaid(doc: Document, code: str):
    for line in ["```mermaid", *code.strip().splitlines(), "```"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(10)
    doc.add_paragraph()


def shade(cell, fill: str = "D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def classify_member(author: str, email: str) -> str | None:
    if author == "huyhoang171106" or email in {"24520554@gm.uit.edu.vn", "huyhoang171106@gmail.com"}:
        return "A"
    if author in {"DuyBao", "Đặng Duy Bảo"}:
        return "B"
    if author == "HufaCung":
        return "C"
    if author == "north2k6":
        return "D"
    return None


def collect_commit_counts():
    subprocess.check_call(["git", "fetch", "origin", "master", "--prune"], cwd=ROOT)
    fmt = "%H%x09%an%x09%ae"
    raw = subprocess.check_output(
        ["git", "log", REV, "--no-merges", f"--since={SINCE}", f"--until={UNTIL}", f"--format={fmt}"],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    counts = Counter()
    aliases: dict[str, Counter[str]] = defaultdict(Counter)
    for line in raw.splitlines():
        _, author, email = line.split("\t")
        member = classify_member(author, email)
        if not member:
            continue
        counts[member] += 1
        aliases[member][f"{author} <{email}>"] += 1
    return counts, aliases


def build_report():
    counts, aliases = collect_commit_counts()
    total = sum(counts.values())

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    add_center(doc, "ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH", 13, True)
    add_center(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN", 13, True)
    for _ in range(5):
        doc.add_paragraph()
    add_center(doc, "BÁO CÁO ĐỒ ÁN CUỐI KỲ MÔN", 15, True)
    add_center(doc, "NGÔN NGỮ LẬP TRÌNH JAVA SE330.Q21", 15, True)
    add_center(doc, "DEVORBIT: NỀN TẢNG QUẢN LÝ VÀ KHÁM PHÁ MÃ NGUỒN HỌC THUẬT", 14, True)
    for _ in range(3):
        doc.add_paragraph()
    add_center(doc, "Sinh viên thực hiện:", 13, True)
    add_center(doc, "Thành viên A - 24520554")
    add_center(doc, "Thành viên B - 24520146")
    add_center(doc, "Thành viên C - 24520071")
    add_center(doc, "Thành viên D - MSSV: cập nhật")
    for _ in range(5):
        doc.add_paragraph()
    add_center(doc, "TP. HỒ CHÍ MINH, THÁNG 06 NĂM 2026", 13, True)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "LỜI CẢM ƠN", 1)
    add_para(doc, "Nhóm thực hiện xin gửi lời cảm ơn đến quý Thầy Cô phụ trách học phần Ngôn ngữ lập trình Java SE330.Q21 đã định hướng kiến thức và tạo điều kiện để nhóm hoàn thành đồ án. Thông qua quá trình xây dựng DevOrbit, nhóm có cơ hội vận dụng các kiến thức về Java, Spring Boot, thiết kế API, cơ sở dữ liệu và kiểm thử phần mềm vào một sản phẩm có tính ứng dụng thực tế.")
    add_para(doc, "Do thời gian thực hiện có hạn, sản phẩm chắc chắn vẫn còn điểm cần hoàn thiện. Nhóm mong nhận được góp ý từ quý Thầy Cô để tiếp tục cải thiện hệ thống và rút kinh nghiệm cho các dự án sau.")

    add_heading(doc, "MỤC LỤC", 1)
    for line in [
        "Chương 1. Giới thiệu đề tài",
        "Chương 2. Yêu cầu chức năng và phân tích thiết kế",
        "Chương 3. Kiến trúc và xây dựng ứng dụng",
        "Chương 4. Kết quả sản phẩm",
        "Chương 5. Đánh giá kết quả thực hiện và phân công công việc",
        "Chương 6. Thuận lợi, khó khăn và bài học kinh nghiệm",
        "Tài liệu tham khảo",
    ]:
        add_para(doc, line, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "DANH MỤC HÌNH", 1)
    add_para(doc, "Không sử dụng hình minh họa riêng trong bản báo cáo này.", first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "DANH MỤC BẢNG", 1)
    for line in [
        "Bảng 0.1: Danh mục từ viết tắt",
        "Bảng 1.1: Công nghệ sử dụng trong đồ án",
        "Bảng 2.1: Yêu cầu chức năng chính",
        "Bảng 3.1: Vai trò các lớp trong backend",
        "Bảng 4.1: Kết quả sản phẩm đạt được",
        "Bảng 5.1: Phân công công việc và đánh giá kết quả",
    ]:
        add_para(doc, line, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "DANH MỤC TỪ VIẾT TẮT", 1)
    add_caption(doc, "Bảng 0.1: Danh mục từ viết tắt")
    add_table(doc, ["Từ viết tắt", "Ý nghĩa", "Ghi chú"], [
        ["API", "Application Programming Interface", "Giao diện lập trình ứng dụng để client gọi backend."],
        ["DTO", "Data Transfer Object", "Đối tượng dùng để trao đổi dữ liệu giữa client và backend."],
        ["JWT", "JSON Web Token", "Token dùng trong xác thực và phân quyền."],
        ["JPA", "Java Persistence API", "Cơ chế ánh xạ object với dữ liệu quan hệ."],
        ["RAG", "Retrieval Augmented Generation", "Cách bổ sung ngữ cảnh dữ liệu trước khi AI trả lời."],
        ["CRUD", "Create, Read, Update, Delete", "Nhóm thao tác cơ bản với dữ liệu."],
    ])

    add_heading(doc, "Chương 1. GIỚI THIỆU ĐỀ TÀI", 1)
    add_heading(doc, "1.1. Giới thiệu chung", 2)
    add_para(doc, "DevOrbit là hệ thống hỗ trợ sinh viên quản lý, tìm kiếm và khai thác mã nguồn học thuật trong quá trình học tập. Thay vì chỉ lưu trữ đường dẫn repository rời rạc, hệ thống gom thông tin môn học, repository, đánh giá, trao đổi cộng đồng và trợ lý học tập vào cùng một nền tảng.")
    add_para(doc, "Trong phạm vi đồ án môn Java, nhóm tập trung nhiều nhất vào phần backend DevOrbit API. Đây là nơi xử lý nghiệp vụ chính, cung cấp REST API, quản lý dữ liệu, xác thực người dùng, quét GitHub repository, xử lý chat cộng đồng và cung cấp các chức năng AI hỗ trợ học tập.")
    add_heading(doc, "1.2. Mục tiêu đề tài", 2)
    add_para(doc, "Mục tiêu chính của đồ án là xây dựng một backend Java có cấu trúc rõ ràng, có thể phục vụ nhiều chức năng thực tế của một hệ thống học thuật. Sản phẩm cần có API ổn định, mô hình dữ liệu hợp lý, kiểm soát quyền truy cập, xử lý lỗi tốt và có khả năng mở rộng thêm tính năng sau này.")
    add_heading(doc, "1.3. Phạm vi thực hiện", 2)
    add_para(doc, "Phạm vi đồ án gồm các module chính: quản lý người dùng và xác thực, quản lý môn học, quản lý repository, quét dữ liệu GitHub, đánh giá repository, community chat, AI tutor, thông báo và các API phục vụ giao diện minh họa. Phần giao diện chỉ được xem là lớp client để kiểm chứng API, không phải trọng tâm chính của báo cáo.")
    add_heading(doc, "1.4. Công nghệ sử dụng", 2)
    add_caption(doc, "Bảng 1.1: Công nghệ sử dụng trong đồ án")
    add_table(doc, ["Nhóm công nghệ", "Công nghệ", "Vai trò trong đồ án"], [
        ["Backend", "Java 21, Spring Boot", "Xây dựng REST API, service nghiệp vụ, bảo mật và xử lý dữ liệu."],
        ["Cơ sở dữ liệu", "PostgreSQL, Flyway", "Lưu trữ dữ liệu chính và quản lý migration."],
        ["Bảo mật", "Spring Security, JWT", "Đăng nhập, phân quyền và bảo vệ endpoint."],
        ["Realtime", "WebSocket", "Hỗ trợ community chat và cập nhật theo thời gian thực."],
        ["AI", "RAG, service AI tutor", "Hỗ trợ hỏi đáp học tập dựa trên dữ liệu hệ thống."],
        ["Client minh họa", "React", "Kiểm chứng API và mô phỏng trải nghiệm người dùng."],
    ])
    add_heading(doc, "1.5. Đối tượng sử dụng", 2)
    add_para(doc, "Đối tượng sử dụng chính của DevOrbit là sinh viên đang học các môn lập trình, đặc biệt là các môn có bài tập lớn, đồ án hoặc yêu cầu tham khảo mã nguồn. Sinh viên có thể dùng hệ thống để tìm repository theo môn học, xem thông tin dự án, đọc đánh giá, trao đổi với cộng đồng và nhận hỗ trợ từ AI tutor.")
    add_para(doc, "Giảng viên hoặc trợ giảng có thể dùng hệ thống như một kênh tham khảo để quan sát cách sinh viên tổ chức mã nguồn, cách repository được phân loại theo môn học và mức độ hoàn thiện của từng dự án. Trong phạm vi đồ án, nhóm chưa triển khai đầy đủ toàn bộ nghiệp vụ quản trị học vụ, nhưng backend đã được thiết kế để có thể mở rộng theo hướng đó.")
    add_heading(doc, "1.6. Ý nghĩa thực tiễn của đề tài", 2)
    add_para(doc, "Trong quá trình học tập, sinh viên thường lưu mã nguồn ở nhiều nơi khác nhau, chủ yếu là GitHub cá nhân hoặc các nhóm nhỏ. Khi cần tìm lại bài làm cũ, tham khảo cách cài đặt một môn học hoặc xem các repository có chất lượng, việc tìm kiếm thường mất thời gian và thiếu ngữ cảnh. DevOrbit được xây dựng để giải quyết vấn đề này ở mức nền tảng, giúp repository không chỉ là đường dẫn mà còn gắn với môn học, đánh giá, trao đổi và hỗ trợ học tập.")
    add_para(doc, "Đề tài cũng giúp nhóm thực hành nhiều nội dung quan trọng của lập trình Java hiện đại: tổ chức source code theo lớp, xây dựng REST API, cấu hình bảo mật, quản lý database bằng migration, xử lý realtime, viết test và tích hợp nhiều module trong cùng một hệ thống. Đây là các kỹ năng cần thiết khi phát triển ứng dụng backend thực tế.")

    add_heading(doc, "Chương 2. YÊU CẦU CHỨC NĂNG VÀ PHÂN TÍCH THIẾT KẾ", 1)
    add_heading(doc, "2.1. Cách nhóm tìm hiểu hệ thống", 2)
    add_para(doc, "Nhóm bắt đầu bằng việc đọc cấu trúc dự án, xác định các package chính trong backend và đối chiếu với các chức năng đang có. Các class được phân loại theo vai trò quen thuộc trong Spring Boot như controller, service, repository, entity, config và DTO.")
    add_para(doc, "Sau đó nhóm kiểm tra từng controller để nắm danh sách API, đọc service để hiểu nghiệp vụ xử lý, xem entity và migration để xác định dữ liệu được lưu như thế nào. Cách làm này phù hợp với báo cáo đồ án vì giúp nhóm mô tả đúng sản phẩm đã xây dựng, không chỉ viết theo ý tưởng ban đầu.")
    add_para(doc, "Khi đọc mã nguồn, nhóm không xem từng file một cách rời rạc mà đi theo luồng chức năng. Ví dụ, với chức năng hỏi đáp AI tutor, nhóm bắt đầu từ controller nhận request, đi tiếp sang service xử lý câu hỏi, sau đó kiểm tra các lớp cấu hình AI, DTO trả về và test tương ứng. Với chức năng repository, nhóm đi từ API repository, sang service quét GitHub, repository lưu dữ liệu và các phần xử lý đánh giá.")
    add_para(doc, "Cách trình bày trong báo cáo vì vậy ưu tiên trả lời ba câu hỏi: chức năng đó dùng để làm gì, backend xử lý như thế nào và kết quả người dùng nhận được là gì. Những chi tiết quá sâu về thuật toán hoặc công cụ chỉ được nhắc khi thật sự liên quan đến việc xây dựng sản phẩm.")
    add_heading(doc, "2.2. Yêu cầu chức năng", 2)
    add_caption(doc, "Bảng 2.1: Yêu cầu chức năng chính")
    add_table(doc, ["Mã", "Chức năng", "Mô tả"], [
        ["F01", "Đăng nhập và phân quyền", "Cho phép người dùng đăng nhập, nhận JWT và truy cập chức năng phù hợp với vai trò."],
        ["F02", "Quản lý môn học", "Lưu thông tin môn học, học kỳ, tài nguyên học tập và dữ liệu liên quan."],
        ["F03", "Quản lý repository", "Lưu và hiển thị repository học thuật gắn với môn học hoặc sinh viên."],
        ["F04", "Quét GitHub repository", "Lấy thông tin repository, phân tích metadata và hỗ trợ đánh giá mã nguồn."],
        ["F05", "Community chat", "Cho phép sinh viên trao đổi theo kênh, hỗ trợ WebSocket và trạng thái online."],
        ["F06", "AI tutor", "Hỗ trợ hỏi đáp học tập dựa trên dữ liệu hệ thống và nội dung đã nạp."],
        ["F07", "Thông báo", "Gửi thông báo cho người dùng về các hoạt động liên quan."],
    ])
    add_para(doc, "Đối với chức năng đăng nhập và phân quyền, yêu cầu quan trọng nhất là hệ thống phải nhận diện đúng người dùng và không cho phép truy cập trái quyền. Sau khi đăng nhập thành công, client nhận token và dùng token đó cho các request tiếp theo. Backend chịu trách nhiệm kiểm tra token, xác định vai trò và quyết định endpoint nào được phép truy cập.")
    add_para(doc, "Đối với chức năng quản lý môn học, hệ thống cần lưu được thông tin cơ bản của môn, mã môn, tên môn và các dữ liệu liên quan. Đây là phần nền để liên kết repository, tài nguyên học tập và các luồng hỏi đáp theo ngữ cảnh môn học. Nếu phần dữ liệu môn học không ổn định, các chức năng phía trên như repository discovery hoặc AI tutor sẽ khó hoạt động chính xác.")
    add_para(doc, "Đối với chức năng repository, người dùng cần xem được các repository gắn với môn học, tìm kiếm repository phù hợp và xem các thông tin mô tả cơ bản. Backend cần lưu metadata, liên kết repository với môn học, hỗ trợ cập nhật thông tin và tạo nền cho các chức năng đánh giá repository.")
    add_para(doc, "Đối với chức năng quét GitHub, hệ thống cần lấy dữ liệu từ repository thật, phân tích các thông tin cần thiết và lưu lại để dùng trong các màn hình đánh giá. Đây là phần giúp DevOrbit khác với một danh sách link thông thường, vì hệ thống có thể tự động đọc một phần thông tin kỹ thuật từ repository.")
    add_para(doc, "Đối với community chat, yêu cầu chính là người dùng có thể trao đổi theo kênh, nhận tin nhắn mới và thấy trạng thái hoạt động của các thành viên. Backend cần hỗ trợ cả lưu trữ tin nhắn lẫn cập nhật realtime, đồng thời đảm bảo người dùng chỉ tương tác với các kênh phù hợp.")
    add_para(doc, "Đối với AI tutor, hệ thống cần nhận câu hỏi học tập, tìm ngữ cảnh liên quan và trả lời theo hướng hỗ trợ sinh viên hiểu bài. Nhóm đặt mục tiêu AI tutor đóng vai trò trợ lý học tập trong hệ thống, không thay thế hoàn toàn tài liệu chính thức hoặc giảng viên.")
    add_heading(doc, "2.3. Yêu cầu phi chức năng", 2)
    add_para(doc, "Hệ thống cần có cấu trúc dễ bảo trì, API trả lỗi rõ ràng, có kiểm thử cho các service quan trọng, cấu hình tách biệt theo môi trường và không để lộ thông tin nhạy cảm. Các chức năng chính cần được thiết kế để có thể mở rộng thêm client giao diện hoặc tích hợp dịch vụ ngoài.")
    add_para(doc, "Về khả năng bảo trì, source code cần được chia thành các package có vai trò rõ ràng. Controller không nên chứa quá nhiều nghiệp vụ, service cần thể hiện luồng xử lý chính, repository tập trung vào truy vấn dữ liệu và DTO giúp kiểm soát dữ liệu trao đổi với client. Cách chia này giúp thành viên khác trong nhóm có thể đọc code nhanh hơn.")
    add_para(doc, "Về độ ổn định, backend cần xử lý được các lỗi thường gặp như dữ liệu không tồn tại, request thiếu thông tin, token không hợp lệ hoặc lỗi khi gọi dịch vụ ngoài. Thay vì để lỗi hệ thống trả về lộn xộn, backend cần chuyển lỗi thành response dễ hiểu để client có thể hiển thị phù hợp.")
    add_para(doc, "Về bảo mật, hệ thống không chỉ cần đăng nhập mà còn cần kiểm soát quyền truy cập theo vai trò. Các thông tin cấu hình nhạy cảm cần tách ra khỏi source code, hạn chế hard-code secret và bảo đảm endpoint quan trọng không bị gọi tự do.")
    add_heading(doc, "2.4. Thiết kế dữ liệu tổng quát", 2)
    add_para(doc, "Dữ liệu của hệ thống xoay quanh các nhóm chính: người dùng, môn học, repository, đánh giá, nội dung chat, thông báo và dữ liệu phục vụ AI tutor. Các bảng được quản lý bằng migration để việc khởi tạo và cập nhật database có thể lặp lại được trong nhiều môi trường.")
    add_para(doc, "Nhóm người dùng lưu thông tin tài khoản, vai trò và các dữ liệu cần thiết cho quá trình xác thực. Nhóm môn học lưu thông tin học phần và làm điểm nối cho các repository hoặc tài nguyên liên quan. Nhóm repository lưu thông tin dự án mã nguồn, đường dẫn GitHub, mô tả, dữ liệu quét và các chỉ số phục vụ đánh giá.")
    add_para(doc, "Nhóm community chat lưu kênh trao đổi, tin nhắn và trạng thái người dùng trong kênh. Nhóm thông báo lưu các sự kiện cần gửi đến người dùng. Riêng nhóm dữ liệu phục vụ AI tutor có vai trò cung cấp ngữ cảnh cho câu hỏi, giúp câu trả lời không quá chung chung mà bám vào dữ liệu học tập trong hệ thống.")
    add_para(doc, "Việc dùng Flyway migration giúp nhóm quản lý thay đổi database có thứ tự. Khi cần thêm bảng hoặc sửa cấu trúc bảng, thay đổi được ghi thành migration riêng thay vì sửa thủ công trên một database cụ thể. Điều này giúp quá trình chạy lại backend trên máy khác dễ kiểm soát hơn.")

    add_heading(doc, "Chương 3. KIẾN TRÚC VÀ XÂY DỰNG ỨNG DỤNG", 1)
    add_heading(doc, "3.1. Kiến trúc backend", 2)
    add_para(doc, "Backend DevOrbit API được tổ chức theo kiến trúc nhiều lớp quen thuộc trong Spring Boot. Controller nhận request và trả response, service xử lý nghiệp vụ, repository giao tiếp với database, entity mô tả dữ liệu, DTO dùng để trao đổi dữ liệu với client.")
    add_table(doc, ["Lớp", "Vai trò", "Ví dụ trong hệ thống"], [
        ["Controller", "Nhận request, kiểm tra đầu vào cơ bản và gọi service.", "SubjectQaController, CommunityController, GithubRepoController."],
        ["Service", "Xử lý nghiệp vụ chính của hệ thống.", "SubjectQaService, GithubScanService, CommunityChatService."],
        ["Repository", "Truy vấn và lưu dữ liệu.", "CourseRepository, GithubRepoRepository, UserRepository."],
        ["Entity", "Biểu diễn bảng dữ liệu.", "User, Course, GithubRepo, CommunityMessage."],
        ["Config", "Cấu hình bảo mật, CORS, AI, cache và runtime.", "SecurityConfig, AiConfig, WebSocketConfig."],
    ])
    add_para(doc, "Luồng xử lý phổ biến của backend bắt đầu từ request gửi vào controller. Controller nhận dữ liệu, chuyển sang DTO phù hợp và gọi service. Service kiểm tra điều kiện nghiệp vụ, lấy dữ liệu thông qua repository, xử lý kết quả và trả lại DTO phản hồi. Cách tổ chức này giúp controller ngắn gọn hơn và các nghiệp vụ quan trọng có thể được kiểm thử độc lập.")
    add_para(doc, "Trong quá trình xây dựng, nhóm ưu tiên giữ các service có trách nhiệm rõ ràng. Ví dụ, service liên quan đến repository tập trung vào xử lý repository, service liên quan đến chat tập trung vào tin nhắn và kênh, service liên quan đến AI tutor tập trung vào câu hỏi, ngữ cảnh và phản hồi. Điều này giúp hệ thống không bị dồn quá nhiều logic vào một lớp duy nhất.")
    add_para(doc, "Các lớp cấu hình đóng vai trò quan trọng trong việc đưa ứng dụng từ code sang trạng thái chạy được. Những cấu hình như bảo mật, CORS, WebSocket, AI, datasource và cache giúp backend có thể kết nối với client, database và các dịch vụ ngoài một cách ổn định hơn.")
    add_heading(doc, "3.2. Các module chính", 2)
    add_para(doc, "Module xác thực chịu trách nhiệm đăng nhập, cấp JWT và bảo vệ API. Module môn học và repository quản lý dữ liệu học thuật. Module GitHub scan hỗ trợ thu thập và phân tích thông tin repository. Module community chat cung cấp kênh trao đổi realtime. Module AI tutor xử lý hỏi đáp học tập, kết hợp dữ liệu có sẵn trong hệ thống.")
    add_para(doc, "Module xác thực là cửa vào của hệ thống. Khi người dùng đăng nhập, backend kiểm tra thông tin, tạo token và yêu cầu client gửi token này trong các request tiếp theo. Nhờ đó, các API quan trọng như quản lý dữ liệu hoặc gửi tin nhắn không bị truy cập tùy ý.")
    add_para(doc, "Module môn học là nền dữ liệu học thuật của DevOrbit. Các chức năng repository, tài nguyên học tập và AI tutor đều cần ngữ cảnh môn học để hoạt động có ý nghĩa. Vì vậy nhóm xem module này như phần dữ liệu lõi, dù giao diện hiển thị có thể thay đổi theo từng client.")
    add_para(doc, "Module repository và GitHub scan là phần thể hiện rõ đặc trưng của đề tài. Backend không chỉ lưu đường dẫn repository mà còn có khả năng lấy thêm thông tin, phân loại và hỗ trợ đánh giá repository. Chức năng này giúp sinh viên xem repository trong bối cảnh học tập thay vì chỉ xem một trang GitHub riêng lẻ.")
    add_para(doc, "Module community chat giúp hệ thống có yếu tố tương tác giữa người học. Tin nhắn và kênh trao đổi được lưu ở backend, trong khi WebSocket giúp client nhận cập nhật nhanh hơn. Đây là phần cho thấy backend không chỉ phục vụ request một lần mà còn hỗ trợ luồng realtime.")
    add_para(doc, "Module AI tutor là phần hỗ trợ sinh viên đặt câu hỏi trong quá trình học. Backend chịu trách nhiệm nhận câu hỏi, chuẩn bị ngữ cảnh, gọi phần xử lý AI và trả kết quả về client. Trong phạm vi đồ án, nhóm tập trung vào việc làm cho luồng này có thể chạy được, có cấu trúc service rõ ràng và có kiểm thử cho những tình huống quan trọng.")
    add_heading(doc, "3.3. Xây dựng API", 2)
    add_para(doc, "Các API được xây dựng theo hướng REST, sử dụng request và response DTO để tách dữ liệu trao đổi với mô hình lưu trữ nội bộ. Các endpoint quan trọng đều đi qua service, giúp code dễ kiểm thử và giảm việc đưa nghiệp vụ trực tiếp vào controller.")
    add_para(doc, "Nhóm ưu tiên đặt tên endpoint theo tài nguyên hoặc nhóm chức năng, ví dụ nhóm API cho môn học, repository, chat, thông báo và AI tutor. Mỗi API cần có vai trò rõ ràng: lấy danh sách, xem chi tiết, tạo mới, cập nhật, xóa hoặc thực hiện một hành động nghiệp vụ cụ thể.")
    add_para(doc, "Với các API nhận dữ liệu từ client, DTO giúp backend kiểm soát dữ liệu đầu vào tốt hơn. Entity không được đưa trực tiếp ra ngoài khi không cần thiết, vì entity thường gắn với cấu trúc database và có thể chứa những trường không phù hợp để client nhìn thấy.")
    add_para(doc, "Với các API trả dữ liệu cho client, response DTO giúp nhóm định hình chính xác dữ liệu hiển thị. Điều này cũng giúp lớp giao diện không phụ thuộc quá sâu vào cấu trúc bảng trong database. Khi database thay đổi, backend vẫn có thể giữ response ổn định nếu thiết kế DTO tốt.")
    add_heading(doc, "3.4. Bảo mật và xử lý lỗi", 2)
    add_para(doc, "Hệ thống sử dụng JWT để xác thực request và Spring Security để cấu hình quyền truy cập. Các lỗi phổ biến được xử lý tập trung để client nhận được thông báo rõ ràng hơn, đồng thời backend dễ theo dõi nguyên nhân khi có sự cố.")
    add_para(doc, "Trong đồ án, bảo mật không chỉ là phần đăng nhập mà còn là cách bảo vệ các luồng nghiệp vụ. Một số endpoint chỉ nên dành cho người dùng đã đăng nhập, một số endpoint cần vai trò quản trị hoặc cần kiểm tra quyền sở hữu dữ liệu trước khi cho phép thao tác.")
    add_para(doc, "Xử lý lỗi tập trung giúp giảm tình trạng mỗi controller trả lỗi theo một kiểu khác nhau. Khi hệ thống có cùng một quy ước về lỗi, client dễ hiển thị thông báo hơn và nhóm phát triển cũng dễ debug hơn. Đây là điểm quan trọng khi hệ thống có nhiều module và nhiều người cùng phát triển.")
    add_heading(doc, "3.5. Kiểm thử và vận hành", 2)
    add_para(doc, "Nhóm bổ sung kiểm thử cho nhiều service quan trọng, đồng thời dùng file cấu hình và script chạy để thuận tiện hơn khi khởi động backend. Đây là phần cần thiết để sản phẩm không chỉ chạy được trên máy một thành viên mà có thể được tái lập ở môi trường khác.")
    add_para(doc, "Kiểm thử trong đồ án tập trung vào các phần có nghiệp vụ rõ hoặc dễ phát sinh lỗi, chẳng hạn xác thực, AI tutor, community chat, hardening database và các service xử lý repository. Các test giúp nhóm phát hiện lỗi sớm khi sửa code, đặc biệt trong giai đoạn nhiều thành viên cùng thay đổi source.")
    add_para(doc, "Về vận hành, backend cần đọc cấu hình từ file cấu hình hoặc biến môi trường để phù hợp với nhiều máy khác nhau. Các thông tin như database URL, token, khóa dịch vụ ngoài hoặc cấu hình AI không nên viết cố định trong code. Cách làm này giúp sản phẩm dễ chạy lại khi chuyển sang máy khác hoặc môi trường triển khai khác.")
    add_para(doc, "Nhóm cũng chú ý đến log và thông báo lỗi khi khởi động. Một backend đồ án tốt không chỉ cần chạy được khi mọi thứ đúng, mà còn cần báo lỗi đủ rõ khi thiếu cấu hình, sai database hoặc dịch vụ ngoài chưa sẵn sàng. Điều này giúp quá trình demo và bảo trì bớt phụ thuộc vào một thành viên duy nhất.")

    add_heading(doc, "Chương 4. KẾT QUẢ SẢN PHẨM", 1)
    add_heading(doc, "4.1. Kết quả đạt được", 2)
    add_para(doc, "Sau quá trình thực hiện, nhóm đã xây dựng được backend DevOrbit API với nhiều chức năng liên quan đến quản lý học thuật, repository, community chat và AI tutor. Hệ thống có cấu trúc source code rõ ràng, tách lớp theo Spring Boot, có migration database và có các test cho những phần quan trọng.")
    add_caption(doc, "Bảng 4.1: Kết quả sản phẩm đạt được")
    add_table(doc, ["Nhóm kết quả", "Mô tả"], [
        ["Backend Java API", "Cung cấp API cho đăng nhập, môn học, repository, chat, thông báo và AI tutor."],
        ["Database", "Có entity và migration để quản lý dữ liệu PostgreSQL."],
        ["Bảo mật", "Có JWT, phân quyền và cấu hình bảo vệ endpoint."],
        ["Realtime", "Có WebSocket phục vụ community chat."],
        ["AI tutor", "Có luồng hỏi đáp học tập và xử lý dữ liệu liên quan."],
        ["Client minh họa", "Giao diện web hỗ trợ kiểm chứng các API chính."],
    ])
    add_para(doc, "Kết quả quan trọng nhất là backend đã hình thành được bộ khung tương đối đầy đủ cho một hệ thống học thuật. Các chức năng không đứng riêng lẻ mà có liên kết với nhau: người dùng đăng nhập, xem môn học, tìm repository, tham gia trao đổi, nhận thông báo và dùng AI tutor trong cùng một hệ thống.")
    add_para(doc, "Phần database đã có migration để mô tả các thay đổi cấu trúc theo thời gian. Điều này giúp quá trình phát triển rõ ràng hơn so với việc tạo bảng thủ công. Khi cần kiểm tra hoặc chạy lại hệ thống, nhóm có thể dựa vào migration để tái tạo cấu trúc database.")
    add_para(doc, "Phần bảo mật đã có nền tảng đăng nhập và phân quyền. Đây là yêu cầu bắt buộc với một hệ thống có dữ liệu người dùng và chức năng cộng đồng. Nếu không có lớp bảo mật, các API như gửi tin nhắn, xem dữ liệu cá nhân hoặc thao tác quản trị sẽ không đủ an toàn để mở rộng.")
    add_para(doc, "Phần AI tutor là điểm mở rộng đáng chú ý của đồ án. Dù chưa thể xem đây là một hệ thống AI hoàn chỉnh ở mức sản phẩm thương mại, nhóm đã xây dựng được luồng backend để nhận câu hỏi, xử lý ngữ cảnh và trả phản hồi. Điều này cho thấy backend có khả năng tích hợp với các chức năng thông minh hơn ngoài CRUD thông thường.")
    add_heading(doc, "4.2. Một số màn hình và luồng sử dụng", 2)
    add_para(doc, "Các màn hình giao diện trong đồ án chủ yếu dùng để minh họa cách client tiêu thụ API. Người dùng có thể đăng nhập, xem môn học, tìm repository, trao đổi trong community, nhận thông báo và dùng AI tutor để hỏi đáp học tập.")
    add_image(doc, str(ROOT / "docs/report/screenshots/01-homepage.png"), "Hình 4.1: Trang chủ DevOrbit")
    add_para(doc, "Luồng đăng nhập bắt đầu khi người dùng nhập thông tin tài khoản ở client. Client gửi request đến backend, backend kiểm tra thông tin và trả token nếu hợp lệ. Sau đó client dùng token để gọi các API cần xác thực. Luồng này giúp minh họa cách Spring Security và JWT được sử dụng trong hệ thống.")
    add_image(doc, str(ROOT / "docs/report/screenshots/05-login.png"), "Hình 4.2: Giao diện đăng nhập")
    add_para(doc, "Luồng repository bắt đầu từ việc người dùng xem hoặc tìm repository theo môn học. Khi cần cập nhật thông tin từ GitHub, backend có thể gọi service quét repository, lấy metadata và lưu lại kết quả. Client chỉ cần gọi API của DevOrbit thay vì tự xử lý trực tiếp toàn bộ logic GitHub.")
    add_image(doc, str(ROOT / "docs/report/screenshots/02-courses.png"), "Hình 4.3: Trang danh sách môn học")
    add_image(doc, str(ROOT / "docs/report/screenshots/03-courses-list.png"), "Hình 4.4: Kết quả tìm kiếm môn học")
    add_para(doc, "Luồng community chat bắt đầu khi người dùng chọn kênh trao đổi. Backend trả danh sách tin nhắn cũ qua API và gửi cập nhật mới thông qua WebSocket. Cách kết hợp REST API và WebSocket giúp hệ thống vừa có dữ liệu lịch sử, vừa có trải nghiệm realtime.")
    add_image(doc, str(ROOT / "docs/report/screenshots/04-community.png"), "Hình 4.5: Trang Cộng đồng")
    add_para(doc, "Ngoài ra, hệ thống còn cung cấp các tính năng bổ sung như Photobooth trực tuyến và Lộ trình học tập giúp sinh viên lên kế hoạch học tập 4 năm tại UIT.")
    add_image(doc, str(ROOT / "docs/report/screenshots/06-photobooth.png"), "Hình 4.6: Trang Photobooth")
    add_image(doc, str(ROOT / "docs/report/screenshots/07-knowledge-graph.png"), "Hình 4.7: Lộ trình học tập")
    add_para(doc, "Luồng AI tutor bắt đầu khi người dùng đặt câu hỏi. Backend nhận câu hỏi, chuẩn bị dữ liệu liên quan, xử lý qua service AI và trả về câu trả lời. Luồng này minh họa cách backend đóng vai trò điều phối giữa client, dữ liệu hệ thống và thành phần AI.")
    add_heading(doc, "4.3. Mức độ hoàn thiện", 2)
    add_para(doc, "Sản phẩm đã đáp ứng được các chức năng chính trong phạm vi đồ án. Một số phần vẫn còn có thể mở rộng thêm như tối ưu hiệu năng khi dữ liệu lớn, bổ sung kiểm thử tích hợp và hoàn thiện trải nghiệm người dùng ở client.")
    add_para(doc, "Ở mức đồ án môn học, hệ thống đã có đủ các phần chính để trình bày: backend Java, database, bảo mật, API, realtime, AI tutor và client minh họa. Các phần này cho thấy nhóm không chỉ làm một chức năng đơn lẻ mà đã xây dựng một hệ thống có nhiều module phối hợp với nhau.")
    add_para(doc, "Tuy nhiên, sản phẩm vẫn còn giới hạn. Một số chức năng cần thêm kiểm thử tích hợp để chắc chắn hoạt động tốt khi kết nối database thật và dịch vụ ngoài. Một số luồng client vẫn có thể được tinh chỉnh thêm để trải nghiệm mượt hơn. Ngoài ra, phần AI tutor cần thêm dữ liệu học tập phong phú hơn để câu trả lời có chiều sâu hơn.")
    add_para(doc, "Nhóm đánh giá mức độ hoàn thiện hiện tại là phù hợp với phạm vi đồ án cuối kỳ. Sản phẩm đủ để demo các chức năng chính, đủ để trình bày kiến trúc Java backend và đủ cơ sở để tiếp tục phát triển nếu có thêm thời gian.")

    add_heading(doc, "Chương 5. ĐÁNH GIÁ KẾT QUẢ THỰC HIỆN VÀ PHÂN CÔNG CÔNG VIỆC", 1)
    add_heading(doc, "5.1. Đánh giá chung", 2)
    add_para(doc, "Nhìn chung, nhóm đã hoàn thành được mục tiêu xây dựng một hệ thống backend Java có nhiều chức năng thực tế, có tổ chức source code rõ ràng và có khả năng mở rộng. Phần backend là trọng tâm của đồ án, trong khi giao diện đóng vai trò hỗ trợ trình bày và kiểm chứng chức năng.")
    add_para(doc, "Về tiến độ, nhóm có khối lượng commit lớn trong giai đoạn từ 14/03/2026 đến 14/06/2026 trên nhánh master. Các commit được tổng hợp theo GitHub Contributors, loại trừ merge commit, để phản ánh tương đối quá trình đóng góp của từng thành viên.")
    add_para(doc, "Về mặt sản phẩm, nhóm đã xây dựng được nhiều chức năng có liên hệ với nhau thay vì chỉ làm các màn hình rời rạc. Backend có thể phục vụ các luồng chính như đăng nhập, xem dữ liệu học thuật, quản lý repository, chat cộng đồng và AI tutor. Đây là điểm quan trọng vì đồ án Java cần thể hiện được cách tổ chức backend và xử lý nghiệp vụ.")
    add_para(doc, "Về mặt kỹ thuật, hệ thống sử dụng các thành phần quen thuộc của Spring Boot như controller, service, repository, entity, DTO, configuration và test. Nhóm có cố gắng tách trách nhiệm giữa các lớp, dùng migration cho database và giữ cấu hình ở mức có thể tái sử dụng. Những điểm này giúp sản phẩm không chỉ chạy được mà còn dễ đọc hơn khi cần bảo trì.")
    add_para(doc, "Về mặt làm việc nhóm, lịch sử commit cho thấy các thành viên có phạm vi đóng góp khác nhau. Có thành viên tham gia nhiều vào backend lõi, có thành viên tập trung vào community và trải nghiệm sử dụng, có thành viên xử lý cấu hình và hardening, có thành viên tập trung vào repository discovery. Cách phân công này giúp sản phẩm bao phủ được nhiều mảng, dù khối lượng giữa các thành viên chưa hoàn toàn cân bằng.")
    add_heading(doc, "5.2. Phân công công việc", 2)
    add_caption(doc, "Bảng 5.1: Phân công công việc và đánh giá kết quả")
    add_table(doc, ["TV", "Số commit", "Phân công chính", "Đánh giá kết quả"], [
        ["A", str(counts["A"]), "Backend lõi (entities, API, JWT auth), AI/RAG (subject QA, streaming chat, embeddings), knowledge graph (impact scoring, simulation mode), frontend utilities, CI/CD.", "Đóng góp lớn nhất về số lượng, đảm nhận phần kỹ thuật phức tạp nhất của hệ thống."],
        ["B", str(counts["B"]), "Community chat (REST/WebSocket, channel presence, subscriptions), GPA calculator (goal planner, presets, autosave, what-if projections), mobile features (bookmarks, search filters, tech stack).", "Xây dựng tính năng community từ đầu đến GPA calculator đầy đủ chức năng."],
        ["C", str(counts["C"]), "Admin panel rewrite, auth system (OTP, forgot password), photobooth backend API, knowledge graph (elective selection, SE2025 curriculum), hardening (CI/CD, Dependabot, dead file cleanup).", "Bao phủ nhiều mảng từ admin đến auth, photobooth đến infrastructure."],
        ["D", str(counts["D"]), "Repository discovery (search matching, alias map, query intent classifier), repo evaluation (X-ray summaries, type classification, dynamic rating), community features (vote/review, WebSocket chat, guest gating).", "Phát triển chức năng phân tích repository, đặc trưng nhất của đề tài."],
    ])
    add_para(doc, "Bảng phân công trên được tổng hợp từ commit trên nhánh master, chỉ tính giai đoạn từ 14/03/2026 đến hết 14/06/2026 và loại trừ merge commit. Nhóm không liệt kê từng commit trong báo cáo để tránh làm phần này quá dài, nhưng khi tổng hợp đã đọc commit message, author và phạm vi file thay đổi để gom theo nội dung công việc.")
    add_para(doc, "Số commit được dùng như một dấu vết tham gia chứ không phải thước đo duy nhất để xếp hạng đóng góp. Có những phần việc cần nhiều commit nhỏ để chỉnh dần, cũng có những phần việc ít commit hơn nhưng đòi hỏi đọc hiểu kỹ, sửa đúng điểm và giữ cho hệ thống không bị lệch hướng. Vì vậy phần đánh giá cá nhân bên dưới tập trung vào vai trò, tinh thần hoàn thành và giá trị mà từng thành viên mang lại cho sản phẩm chung.")
    add_heading(doc, "5.3. Đánh giá cá nhân", 2)
    add_para(doc, "Thành viên A là người đóng góp lớn nhất về số lượng commit và cũng là người đảm nhận nhiều phần kỹ thuật phức tạp nhất. A xây dựng backend lõi từ đầu: entity, API, JWT auth, admin course management, GitHub candidate scanning. Sau đó A phát triển toàn bộ hệ thống AI/RAG bao gồm subject QA streaming chat, Fireworks embedding provider, RAG semantic retrieval với hybrid retrieval, query expansion và reranking. A cũng triển khai knowledge graph với impact scoring, simulation mode và cascade animations.")
    add_para(doc, "Bên cạnh backend, A còn đóng góp nhiều vào frontend: cosmic redesign, shared components (Avatar, LoadingSkeleton, Toast notification), Vietnamese localization, performance monitoring utilities và hàng chục unit test. Về infrastructure, A enforce CI/CD pre-push gate, thực hiện security hardening và tối ưu database initialization với batch inserts. A giữ vai trò như trục kỹ thuật chính của nhóm, đảm bảo phần lõi hệ thống hoạt động ổn định.")
    add_para(doc, "Thành viên B xây dựng tính năng community chat từ đầu, bao gồm REST/WebSocket endpoints, channel presence tracking, online members UI và subscription management. B cũng phát triển đầy đủ GPA calculator với nhiều tính năng: goal what-if projections, semester presets từ roadmap, autosave draft, cumulative GPA estimate mode và background customization. Trên mobile, B đóng góp course bookmarks, search filters, repo tech stack filter và AI roadmap QA viewmodels.")
    add_para(doc, "Điểm đáng chú ý trong phần đóng góp của B là tính hệ thống: mỗi tính năng đều được phát triển từ backend đến frontend, từ web đến mobile. Community chat không chỉ có WebSocket mà còn có entity, migration, REST endpoints và test contract. GPA calculator không chỉ có giao diện mà còn có goal planner, draft persistence và integration với roadmap. Cách làm này giúp sản phẩm có độ hoàn thiện cao ở các tính năng B phụ trách.")
    add_para(doc, "Thành viên C tham gia sâu vào nhiều mảng quan trọng của hệ thống. Về quản trị, C viết lại admin panel với layout mới, bổ sung hệ thống thông báo và chỉnh UI dashboard sinh viên. Về xác thực, C triển khai đầy đủ luồng đăng nhập, đăng ký, quên mật khẩu và xác thực OTP qua email — phần này ảnh hưởng trực tiếp đến trải nghiệm người dùng đầu tiên khi tiếp cận hệ thống.")
    add_para(doc, "C cũng là người xây dựng backend API cho Photobooth thay vì dùng Supabase trực tiếp, xử lý frame management, per-slot filters và crop/zoom. Về knowledge graph, C cải thiện card chọn môn tự chọn và tích hợp chương trình SE2025 với AI roadmap generator. Ngoài ra, C khởi tạo dự án DevOrbit Mobile và triển khai Room migration. Về vận hành, C bổ sung CI/CD pipelines, Dependabot, ESLint, các file bảo vệ repo và thực hiện refactor loại bỏ hàng chục file chết trên cả mobile, web và API. Kiểu đóng góp của C vừa rộng vừa chạm đến nhiều layer — từ giao diện admin đến backend auth, từ mobile đến infrastructure — giúp hệ thống chắc hơn ở nhiều điểm cùng lúc.")
    add_para(doc, "Thành viên D tập trung vào repository discovery và evaluation — phần đặc trưng nhất của đề tài DevOrbit. D xây dựng hệ thống search matching với alias map, repo-based course matching và query intent classifier giúp sinh viên tìm repository chính xác hơn. Về repository evaluation, D phát triển X-ray repo summaries, repo type classification, signal detection, contextual analysis và dynamic rating — những tính năng giúp DevOrbit khác biệt với một danh sách link GitHub thông thường.")
    add_para(doc, "D cũng đóng góp vào community features: repo vote và review system, channel subscribe/unsubscribe, guest gating và real-time WebSocket chat với STOMP và JWT auth. Trên frontend, D cải thiện course search relevance, pagination, review component với star rating và share dialog cho repo. Dù số commit ít hơn một số thành viên khác, phần việc của D có định hướng rõ ràng và tạo giá trị riêng cho sản phẩm: DevOrbit không chỉ quản lý môn học mà còn phân tích và đánh giá mã nguồn học thuật một cách có chiều sâu.")
    add_para(doc, "Đóng góp của D cho thấy giá trị của sự tập trung. Một nhóm đồ án không chỉ cần người làm nhiều phần rộng, mà cũng cần người giữ một mảng chức năng đủ rõ để sản phẩm có chiều sâu. Vì vậy, phần đánh giá cá nhân không xem số commit thấp hơn là đóng góp thấp hơn, mà nhìn vào vai trò thực tế của module đối với mục tiêu chung.")
    add_heading(doc, "5.4. Nhận xét về kết quả", 2)
    add_para(doc, "Điểm mạnh của nhóm là đã chia được các mảng tương đối rõ và có sản phẩm chạy được với nhiều chức năng. Điểm cần cải thiện là khối lượng công việc giữa các thành viên chưa thật đồng đều, một số chức năng còn cần thêm kiểm thử tích hợp và tài liệu hướng dẫn sử dụng chi tiết hơn.")
    add_para(doc, "Nếu xét theo mục tiêu môn học, đồ án đã thể hiện được nhiều kiến thức Java backend: tổ chức package, xây dựng REST API, xử lý service, kết nối database, migration, bảo mật, realtime và kiểm thử. Đây là các nội dung có thể liên hệ trực tiếp với yêu cầu học phần, không chỉ là phần trình bày giao diện.")
    add_para(doc, "Nếu xét theo sản phẩm, DevOrbit đã có hướng đi rõ và có thể phát triển tiếp. Các module hiện tại đủ để tạo nền cho một hệ thống hỗ trợ sinh viên quản lý và khám phá mã nguồn học thuật. Tuy nhiên, để trở thành sản phẩm hoàn chỉnh, hệ thống cần thêm kiểm thử tích hợp, dữ liệu thật phong phú hơn và cơ chế quản trị chi tiết hơn.")
    add_para(doc, "Nhóm cũng nhận thấy việc làm một hệ thống nhiều module đòi hỏi thống nhất quy ước code từ sớm. Nếu không có quy ước rõ về tên API, DTO, cách xử lý lỗi và cách viết test, quá trình ghép module sẽ dễ phát sinh lỗi. Đây là kinh nghiệm quan trọng cho những dự án tiếp theo.")

    add_heading(doc, "Chương 6. THUẬN LỢI, KHÓ KHĂN VÀ BÀI HỌC KINH NGHIỆM", 1)
    add_heading(doc, "6.1. Thuận lợi", 2)
    add_para(doc, "Nhóm có sẵn nền tảng kiến thức Java và Spring Boot từ học phần, đồng thời đề tài có tính gần với nhu cầu học tập thực tế của sinh viên. Việc dùng GitHub làm nguồn dữ liệu cũng giúp nhóm dễ kiểm chứng và trình bày kết quả sản phẩm.")
    add_para(doc, "Một thuận lợi khác là hệ sinh thái Spring Boot có nhiều thư viện hỗ trợ sẵn cho các nhu cầu thường gặp như xây dựng REST API, bảo mật, WebSocket, cache và kiểm thử. Nhờ đó nhóm có thể tập trung nhiều hơn vào cách tổ chức hệ thống và nghiệp vụ DevOrbit thay vì phải tự xây toàn bộ hạ tầng từ đầu.")
    add_para(doc, "Dự án cũng có dữ liệu và bối cảnh khá gần với sinh viên UIT, nên nhóm dễ hình dung người dùng thật sẽ cần gì. Các chức năng như tìm repository theo môn học, trao đổi cộng đồng hoặc hỏi AI tutor đều xuất phát từ nhu cầu tương đối quen thuộc trong quá trình học lập trình.")
    add_heading(doc, "6.2. Khó khăn", 2)
    add_para(doc, "Khó khăn chính nằm ở việc hệ thống có nhiều module liên quan với nhau, từ xác thực, database, repository, realtime chat đến AI tutor. Khi một module thay đổi, các phần còn lại cũng cần được kiểm tra để tránh lỗi dây chuyền. Ngoài ra, việc phối hợp giữa backend và lớp giao diện minh họa cũng làm tăng khối lượng kiểm thử.")
    add_para(doc, "Khó khăn tiếp theo là việc tích hợp với dịch vụ ngoài như GitHub hoặc AI tutor có nhiều tình huống không ổn định. Dịch vụ ngoài có thể thay đổi dữ liệu, giới hạn request hoặc trả lỗi ngoài dự đoán. Backend vì vậy cần xử lý lỗi tốt hơn so với các chức năng chỉ đọc ghi database nội bộ.")
    add_para(doc, "Việc quản lý database cũng là một phần cần cẩn thận. Khi nhiều thành viên cùng thêm entity hoặc migration, nếu không thống nhất quy ước đặt tên và thứ tự migration, hệ thống có thể gặp lỗi khi chạy lại từ đầu. Nhóm đã dùng Flyway để giảm rủi ro này, nhưng vẫn cần kiểm tra kỹ mỗi khi thay đổi cấu trúc dữ liệu.")
    add_para(doc, "Ngoài ra, do phạm vi đồ án khá rộng, nhóm phải cân bằng giữa việc thêm chức năng mới và việc làm chắc các chức năng đã có. Nếu chỉ thêm chức năng mà không kiểm thử, sản phẩm dễ có nhiều lỗi nhỏ. Nếu chỉ tập trung kiểm thử, tiến độ demo tính năng lại bị chậm. Đây là bài toán nhóm phải điều chỉnh trong quá trình thực hiện.")
    add_heading(doc, "6.3. Bài học kinh nghiệm", 2)
    add_para(doc, "Qua đồ án, nhóm nhận ra rằng xây dựng một backend không chỉ là viết cho đủ endpoint. Một hệ thống muốn đứng vững cần có cách chia lớp rõ ràng, dữ liệu được quản lý có trật tự, lỗi được xử lý có trách nhiệm và các thành viên hiểu được phần việc của nhau. Khi những điều này chưa rõ, code vẫn có thể chạy trong một thời điểm, nhưng rất khó sửa và khó mở rộng khi đồ án lớn dần.")
    add_para(doc, "Bài học đầu tiên là cần giữ nghiệp vụ ở đúng vị trí. Controller nên là nơi tiếp nhận request và trả response, còn service mới là nơi thể hiện suy nghĩ chính của hệ thống. Khi nhóm để service gánh phần xử lý nghiệp vụ, việc đọc code trở nên mạch lạc hơn, test cũng dễ viết hơn và mỗi thay đổi nhỏ ít gây ảnh hưởng dây chuyền hơn.")
    add_para(doc, "Bài học thứ hai là dữ liệu cần được tôn trọng ngay từ đầu. Entity, DTO và migration không chỉ là phần phụ trợ, mà là cách hệ thống tự mô tả chính mình. Nếu dữ liệu được thiết kế vội, các chức năng phía trên sẽ phải sửa đi sửa lại. Nếu migration không rõ, mỗi máy chạy một kiểu. Vì vậy, nhóm học được rằng database cũng là một phần quan trọng của source code.")
    add_para(doc, "Bài học thứ ba là kiểm thử không nên để đến cuối. Khi dự án còn nhỏ, chạy thử thủ công có vẻ nhanh hơn. Nhưng khi số module tăng lên, mỗi lần sửa một service đều có thể ảnh hưởng đến controller, repository hoặc client. Test giúp nhóm giữ lại sự tự tin sau mỗi lần thay đổi, nhất là ở các phần như xác thực, chat, AI tutor và xử lý repository.")
    add_para(doc, "Bài học thứ tư là giao tiếp trong nhóm quan trọng không kém kỹ thuật. Một thành viên có thể làm phần backend lõi, một thành viên chăm chút luồng sử dụng, một thành viên giữ cấu hình và bảo mật, một thành viên đào sâu repository discovery. Những phần việc này khác nhau về hình thức, nhưng đều cần được nối lại thành một sản phẩm chung. Khi nhóm hiểu vai trò của nhau, việc đánh giá đóng góp cũng công bằng và nhẹ nhàng hơn.")
    add_para(doc, "Bài học cuối cùng là một đồ án tốt không nhất thiết phải hoàn hảo, nhưng cần trung thực với những gì đã làm được và những gì còn thiếu. DevOrbit vẫn còn nhiều hướng để phát triển, nhưng quá trình thực hiện đã giúp nhóm hiểu rõ hơn cách biến kiến thức Java trong lớp học thành một hệ thống có cấu trúc, có dữ liệu, có người dùng và có khả năng tiếp tục đi xa hơn sau môn học.")
    add_heading(doc, "6.4. Hướng phát triển", 2)
    add_para(doc, "Trong tương lai, nhóm có thể tiếp tục hoàn thiện kiểm thử tích hợp, tối ưu hiệu năng với dữ liệu lớn, bổ sung dashboard quản trị đầy đủ hơn và mở rộng AI tutor để hỗ trợ nhiều loại tài liệu học tập hơn.")
    add_para(doc, "Hướng phát triển gần nhất là bổ sung integration test cho các luồng quan trọng như đăng nhập, repository, community chat và AI tutor. Khi có integration test, nhóm có thể tự tin hơn mỗi khi thay đổi code backend hoặc cập nhật database.")
    add_para(doc, "Tiếp theo, hệ thống có thể mở rộng phần repository analysis để đọc thêm cấu trúc thư mục, ngôn ngữ sử dụng, file README và thông tin commit. Những dữ liệu này giúp việc đánh giá repository có cơ sở hơn, đồng thời hỗ trợ sinh viên tìm được dự án phù hợp với nhu cầu học tập.")
    add_para(doc, "AI tutor cũng có thể được cải thiện bằng cách nạp thêm tài liệu học phần, slide, bài tập mẫu hoặc hướng dẫn thực hành. Khi nguồn dữ liệu phong phú hơn, câu trả lời có thể bám sát môn học hơn và hữu ích hơn cho sinh viên.")
    add_para(doc, "Về vận hành, nhóm có thể bổ sung logging, monitoring và tài liệu triển khai chi tiết. Đây là các phần thường chưa được ưu tiên trong đồ án đầu tiên, nhưng rất quan trọng nếu muốn đưa hệ thống đến gần môi trường sử dụng thật.")

    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        ("Spring Boot Documentation", "https://docs.spring.io/spring-boot/"),
        ("Spring Security Reference", "https://docs.spring.io/spring-security/reference/"),
        ("PostgreSQL Documentation", "https://www.postgresql.org/docs/"),
        ("Flyway Documentation", "https://documentation.red-gate.com/fd"),
        ("GitHub REST API Documentation", "https://docs.github.com/en/rest"),
        ("DevOrbit GitHub Contributors", "https://github.com/huyhoang171106/DevOrbit/graphs/contributors?from=3%2F14%2F2026"),
    ]
    for idx, (label, url) in enumerate(refs, start=1):
        p = add_para(doc, f"[{idx}] {label}. ", first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        add_hyperlink(p, url, url)

    doc.save(OUTPUT)
    print(f"created={OUTPUT}")
    print(f"commits={dict(counts)} total={total}")


if __name__ == "__main__":
    build_report()
