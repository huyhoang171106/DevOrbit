from __future__ import annotations

import os
import subprocess
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = Path(os.environ.get("DEVORBIT_COURSE_REPORT_OUTPUT", ROOT / "devorbit-report-v10.docx"))
REV = "origin/master"
SINCE = "2026-03-14T00:00:00"
UNTIL = "2026-06-15T00:00:00"


def set_run_font(run, size: int = 13, bold: bool = False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc: Document, text: str = "", *, first_line: bool = True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.75)
    if text:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 14 if level == 1 else 13, True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_center(doc: Document, text: str, size: int = 13, bold: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size, bold)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], font_size: int = 11):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, header in enumerate(headers):
        cell = table.cell(0, c)
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, font_size, True)
        shade(cell)
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(value))
            set_run_font(run, font_size)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    return table


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, 12, True)
    return p


def add_image(doc: Document, image_path: str, caption: str, width_cm: float = 14.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    add_caption(doc, caption)
    return p


def add_mermaid(doc: Document, code: str):
    for line in ["```mermaid", *code.strip().splitlines(), "```"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(10)
    doc.add_paragraph()


def shade(cell, fill: str = "D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def classify_member(author: str, email: str) -> str | None:
    if author == "huyhoang171106" or email in {"24520554@gm.uit.edu.vn", "huyhoang171106@gmail.com"}:
        return "A"
    if author in {"DuyBao", "Đặng Duy Bảo"}:
        return "B"
    if author == "HufaCung":
        return "C"
    if author == "north2k6":
        return "D"
    return None


def collect_commit_counts():
    subprocess.check_call(["git", "fetch", "origin", "master", "--prune"], cwd=ROOT)
    fmt = "%H%x09%an%x09%ae"
    raw = subprocess.check_output(
        ["git", "log", REV, "--no-merges", f"--since={SINCE}", f"--until={UNTIL}", f"--format={fmt}"],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    counts = Counter()
    aliases: dict[str, Counter[str]] = defaultdict(Counter)
    for line in raw.splitlines():
        _, author, email = line.split("\t")
        member = classify_member(author, email)
        if not member:
            continue
        counts[member] += 1
        aliases[member][f"{author} <{email}>"] += 1
    return counts, aliases


def build_report():
    counts, aliases = collect_commit_counts()
    total = sum(counts.values())

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    # ── Cover page ──
    add_center(doc, "ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH", 13, True)
    add_center(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN", 13, True)
    for _ in range(5):
        doc.add_paragraph()
    add_center(doc, "BÁO CÁO ĐỒ ÁN CUỐI KỲ MÔN", 15, True)
    add_center(doc, "NGÔN NGỮ LẬP TRÌNH JAVA SE330.Q21", 15, True)
    add_center(doc, "DEVORBIT: NỀN TẢNG QUẢN LÝ VÀ KHÁM PHÁ MÃ NGUỒN HỌC THUẬT", 14, True)
    for _ in range(3):
        doc.add_paragraph()
    add_center(doc, "Sinh viên thực hiện:", 13, True)
    add_center(doc, "Thành viên A - 24520554")
    add_center(doc, "Thành viên B - 24520146")
    add_center(doc, "Thành viên C - 24520071")
    add_center(doc, "Thành viên D - MSSV: cập nhật")
    for _ in range(5):
        doc.add_paragraph()
    add_center(doc, "TP. HỒ CHÍ MINH, THÁNG 06 NĂM 2026", 13, True)

    # ── Lời cảm ơn ──
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "LỜI CẢM ƠN", 1)
    add_para(doc, "Nhóm xin cảm ơn quý Thầy Cô phụ trách học phần Ngôn ngữ lập trình Java SE330.Q21 đã hướng dẫn nhóm trong quá trình thực hiện đồ án. Nhờ những hướng dẫn về Java, Spring Boot, thiết kế REST API và quản lý cơ sở dữ liệu, nhóm mới có đủ kiến thức để xây dựng DevOrbit.")
    add_para(doc, "Đồ án còn nhiều điểm chưa hoàn thiện. Nhóm mong nhận được góp ý từ Thầy Cô để sửa đổi và học hỏi thêm.")

    # ── Mục lục ──
    add_heading(doc, "MỤC LỤC", 1)
    for line in [
        "Chương 1. Giới thiệu đề tài",
        "Chương 2. Yêu cầu chức năng và phân tích thiết kế",
        "Chương 3. Kiến trúc và xây dựng ứng dụng",
        "Chương 4. Kết quả sản phẩm",
        "Chương 5. Đánh giá kết quả thực hiện và phân công công việc",
        "Chương 6. Thuận lợi, khó khăn và bài học kinh nghiệm",
        "Tài liệu tham khảo",
    ]:
        add_para(doc, line, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "DANH MỤC HÌNH", 1)
    add_para(doc, "Không sử dụng hình minh họa riêng trong bản báo cáo này.", first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "DANH MỤC BẢNG", 1)
    for line in [
        "Bảng 0.1: Danh mục từ viết tắt",
        "Bảng 1.1: Công nghệ sử dụng trong đồ án",
        "Bảng 2.1: Yêu cầu chức năng chính",
        "Bảng 3.1: Vai trò các lớp trong backend",
        "Bảng 4.1: Kết quả sản phẩm đạt được",
        "Bảng 5.1: Phân công công việc và đánh giá kết quả",
    ]:
        add_para(doc, line, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    # ── Danh mục từ viết tắt ──
    add_heading(doc, "DANH MỤC TỪ VIẾT TẮT", 1)
    add_caption(doc, "Bảng 0.1: Danh mục từ viết tắt")
    add_table(doc, ["Từ viết tắt", "Ý nghĩa", "Ghi chú"], [
        ["API", "Application Programming Interface", "Giao diện lập trình ứng dụng để client gọi backend."],
        ["DTO", "Data Transfer Object", "Đối tượng dùng để trao đổi dữ liệu giữa client và backend."],
        ["JWT", "JSON Web Token", "Token dùng trong xác thực và phân quyền."],
        ["JPA", "Java Persistence API", "Cơ chế ánh xạ object với dữ liệu quan hệ."],
        ["RAG", "Retrieval Augmented Generation", "Cách bổ sung ngữ cảnh dữ liệu trước khi AI trả lời."],
        ["CRUD", "Create, Read, Update, Delete", "Nhóm thao tác cơ bản với dữ liệu."],
    ])

    # ══════════════════════════════════════════════════════════════
    # CHƯƠNG 1
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "Chương 1. GIỚI THIỆU ĐỀ TÀI", 1)

    add_heading(doc, "1.1. Giới thiệu chung", 2)
    add_para(doc, "Khi học lập trình, sinh viên UIT thường có nhiều repository trên GitHub cho từng môn học. Nhưng khi cần tìm lại bài làm cũ, tham khảo cách người khác cài đặt một bài tập, hoặc xem repo nào tốt cho một môn nào đó, việc tìm kiếm khá mất thời gian. Mỗi repo là một đường dẫn riêng lẻ, không gắn với môn học, không có đánh giá, không có ai cùng thảo luận.")
    add_para(doc, "DevOrbit được xây dựng để giải quyết vấn đề đó. Hệ thống gom thông tin về môn học, repository, đánh giá, chat cộng đồng và trợ lý AI vào một nền tảng duy nhất. Sinh viên có thể tìm repo theo môn, xem评价 từ người khác, trao đổi trực tiếp trong cộng đồng, và đặt câu hỏi với AI tutor.")
    add_para(doc, "Trong phạm vi đồ án Java này, nhóm tập trung xây dựng phần backend DevOrbit API. Đây là nơi xử lý toàn bộ nghiệp vụ: cung cấp REST API, quản lý dữ liệu PostgreSQL, xác thực người dùng bằng JWT, quét repository từ GitHub, xử lý chat realtime qua WebSocket, và tích hợp AI tutor.")

    add_heading(doc, "1.2. Mục tiêu đề tài", 2)
    add_para(doc, "Mục tiêu chính là xây dựng một backend Java có cấu trúc rõ ràng, chạy được nhiều chức năng thực tế. Cụ thể: API phải hoạt động ổn định, mô hình dữ liệu phải hợp lý, kiểm soát quyền truy cập phải chặt chẽ, xử lý lỗi phải rõ ràng, và hệ thống phải có thể mở rộng thêm tính năng sau này.")
    add_para(doc, "Nhóm không đặt mục tiêu xây dựng một sản phẩm thương mại. Mục tiêu là một backend đồ án có đủ các phần chính để trình bày trước hội đồng: từ xác thực, quản lý dữ liệu, realtime, đến AI tutor.")

    add_heading(doc, "1.3. Phạm vi thực hiện", 2)
    add_para(doc, "Đồ án gồm các module sau: quản lý người dùng và xác thực, quản lý môn học, quản lý repository, quét dữ liệu GitHub, đánh giá repository, community chat, AI tutor, thông báo, và các API phục vụ giao diện minh họa. Giao diện web được xem là lớp client để kiểm chứng API, không phải trọng tâm báo cáo.")

    add_heading(doc, "1.4. Công nghệ sử dụng", 2)
    add_caption(doc, "Bảng 1.1: Công nghệ sử dụng trong đồ án")
    add_table(doc, ["Nhóm công nghệ", "Công nghệ", "Vai trò trong đồ án"], [
        ["Backend", "Java 21, Spring Boot", "Xây dựng REST API, service nghiệp vụ, bảo mật và xử lý dữ liệu."],
        ["Cơ sở dữ liệu", "PostgreSQL, Flyway", "Lưu trữ dữ liệu chính và quản lý migration."],
        ["Bảo mật", "Spring Security, JWT", "Đăng nhập, phân quyền và bảo vệ endpoint."],
        ["Realtime", "WebSocket", "Hỗ trợ community chat và cập nhật theo thời gian thực."],
        ["AI", "RAG, service AI tutor", "Hỗ trợ hỏi đáp học tập dựa trên dữ liệu hệ thống."],
        ["Client minh họa", "React", "Kiểm chứng API và mô phỏng trải nghiệm người dùng."],
    ])

    add_heading(doc, "1.5. Đối tượng sử dụng", 2)
    add_para(doc, "Đối tượng chính là sinh viên đang học các môn lập trình, đặc biệt là các môn có bài tập lớn hoặc yêu cầu tham khảo mã nguồn. Sinh viên dùng hệ thống để tìm repository theo môn học, xem đánh giá, trao đổi cộng đồng và nhận hỗ trợ từ AI tutor.")
    add_para(doc, "Giảng viên hoặc trợ giảng có thể dùng hệ thống để xem cách sinh viên tổ chức mã nguồn, cách repository được phân loại theo môn học. Nhóm chưa triển khai đầy đủ nghiệp vụ quản trị học vụ, nhưng backend đã được thiết kế để có thể mở rộng theo hướng đó.")

    add_heading(doc, "1.6. Ý nghĩa thực tiễn của đề tài", 2)
    add_para(doc, "Khi cần tìm lại bài làm cũ hoặc tham khảo cách cài đặt một môn học, sinh viên thường phải tìm trên GitHub cá nhân hoặc hỏi bạn bè. Việc tìm kiếm mất thời gian và thiếu ngữ cảnh. DevOrbit giúp repository không chỉ là một đường dẫn mà còn gắn với môn học, đánh giá và thảo luận.")
    add_para(doc, "Đề tài cũng giúp nhóm thực hành nhiều nội dung quan trọng của lập trình Java hiện đại: tổ chức source theo lớp, xây dựng REST API, cấu hình bảo mật, quản lý database bằng migration, xử lý realtime, viết test, và tích hợp nhiều module trong một hệ thống. Đây đều là kỹ năng cần thiết khi phát triển backend thực tế.")

    # ══════════════════════════════════════════════════════════════
    # CHƯƠNG 2
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "Chương 2. YÊU CẦU CHỨC NĂNG VÀ PHÂN TÍCH THIẾT KẾ", 1)

    add_heading(doc, "2.1. Cách nhóm tìm hiểu hệ thống", 2)
    add_para(doc, "Nhóm bắt đầu bằng việc đọc cấu trúc thư mục dự án, xác định các package chính trong backend rồi đối chiếu với chức năng hiện có. Các class được phân loại theo vai trò quen thuộc trong Spring Boot: controller, service, repository, entity, config và DTO.")
    add_para(doc, "Sau đó nhóm kiểm tra từng controller để nắm danh sách API, đọc service để hiểu nghiệp vụ xử lý, xem entity và migration để xác định dữ liệu được lưu thế nào. Cách này giúp nhóm mô tả đúng sản phẩm đã xây dựng, không viết theo ý tưởng ban đầu.")
    add_para(doc, "Khi đọc code, nhóm không xem từng file riêng lẻ mà đi theo luồng chức năng. Ví dụ với AI tutor, nhóm bắt đầu từ controller nhận request, sang service xử lý câu hỏi, rồi kiểm tra cấu hình AI, DTO trả về và test tương ứng. Với repository, nhóm đi từ API, qua service quét GitHub, repository lưu dữ liệu, đến phần xử lý đánh giá.")
    add_para(doc, "Cách trình bày trong báo cáo vì vậy tập trung vào ba câu hỏi: chức năng đó làm gì, backend xử lý thế nào, và người dùng nhận được gì. Chi tiết quá sâu về thuật toán chỉ được nhắc khi liên quan trực tiếp đến xây dựng sản phẩm.")

    add_heading(doc, "2.2. Yêu cầu chức năng", 2)
    add_caption(doc, "Bảng 2.1: Yêu cầu chức năng chính")
    add_table(doc, ["Mã", "Chức năng", "Mô tả"], [
        ["F01", "Đăng nhập và phân quyền", "Cho phép người dùng đăng nhập, nhận JWT và truy cập chức năng phù hợp với vai trò."],
        ["F02", "Quản lý môn học", "Lưu thông tin môn học, học kỳ, tài nguyên học tập và dữ liệu liên quan."],
        ["F03", "Quản lý repository", "Lưu và hiển thị repository học thuật gắn với môn học hoặc sinh viên."],
        ["F04", "Quét GitHub repository", "Lấy thông tin repository, phân tích metadata và hỗ trợ đánh giá mã nguồn."],
        ["F05", "Community chat", "Cho phép sinh viên trao đổi theo kênh, hỗ trợ WebSocket và trạng thái online."],
        ["F06", "AI tutor", "Hỗ trợ hỏi đáp học tập dựa trên dữ liệu hệ thống và nội dung đã nạp."],
        ["F07", "Thông báo", "Gửi thông báo cho người dùng về các hoạt động liên quan."],
    ])
    add_para(doc, "Với chức năng đăng nhập, yêu cầu quan trọng nhất là hệ thống phải nhận diện đúng người dùng và không cho truy cập trái quyền. Sau khi đăng nhập, client nhận JWT và dùng token đó cho các request tiếp theo. Backend kiểm tra token, xác định vai trò, và quyết định endpoint nào được phép.")
    add_para(doc, "Quản lý môn học là phần nền dữ liệu. Nếu thông tin môn học không ổn định, các chức năng phía trên như repository discovery hay AI tutor sẽ khó hoạt động chính xác. Backend cần lưu được mã môn, tên môn và các dữ liệu liên quan.")
    add_para(doc, "Với repository, người dùng cần xem được các repo gắn với môn học, tìm kiếm phù hợp và xem mô tả cơ bản. Backend lưu metadata, liên kết repo với môn học và tạo nền cho chức năng đánh giá.")
    add_para(doc, "Chức năng quét GitHub là điểm khác biệt của DevOrbit so với một danh sách link thông thường. Hệ thống tự động lấy thông tin từ repo thật, phân tích metadata và lưu lại để dùng trong đánh giá.")
    add_para(doc, "Community chat yêu cầu người dùng trao đổi được theo kênh, nhận tin nhắn mới và thấy trạng thái online. Backend cần hỗ trợ cả lưu trữ tin nhắn lẫn cập nhật realtime qua WebSocket.")
    add_para(doc, "AI tutor nhận câu hỏi học tập, tìm ngữ cảnh liên quan và trả lời hướng hỗ trợ sinh viên hiểu bài. Nhóm đặt mục tiêu AI tutor là trợ lý học tập, không thay thế tài liệu chính thức hay giảng viên.")

    add_heading(doc, "2.3. Yêu cầu phi chức năng", 2)
    add_para(doc, "Hệ thống cần dễ bảo trì, API trả lỗi rõ ràng, có kiểm thử cho service quan trọng, cấu hình tách biệt theo môi trường, và không để lộ thông tin nhạy cảm. Các chức năng chính phải thiết kế để có thể mở rộng thêm client hoặc tích hợp dịch vụ ngoài.")
    add_para(doc, "Về bảo trì, source code chia thành các package có vai trò rõ ràng. Controller không chứa nhiều nghiệp vụ, service thể hiện luồng xử lý chính, repository tập trung truy vấn, DTO kiểm soát dữ liệu trao đổi. Cách này giúp thành viên khác đọc code nhanh hơn.")
    add_para(doc, "Về ổn định, backend xử lý các lỗi thường gặp: dữ liệu không tồn tại, request thiếu thông tin, token không hợp lệ, hoặc lỗi dịch vụ ngoài. Thay vì để hệ thống trả về lỗi lộn xộn, backend chuyển lỗi thành response dễ hiểu để client hiển thị phù hợp.")
    add_para(doc, "Về bảo mật, ngoài đăng nhập còn cần kiểm soát quyền theo vai trò. Thông tin cấu hình nhạy cảm tách khỏi source code, hạn chế hard-code secret, và các endpoint quan trọng không bị gọi tự do.")

    add_heading(doc, "2.4. Thiết kế dữ liệu tổng quát", 2)
    add_para(doc, "Dữ liệu xoay quanh các nhóm chính: người dùng, môn học, repository, đánh giá, nội dung chat, thông báo và dữ liệu phục vụ AI tutor. Các bảng được quản lý bằng Flyway migration để việc khởi tạo và cập nhật database có thể lặp lại trên nhiều môi trường.")
    add_para(doc, "Nhóm người dùng lưu tài khoản, vai trò và dữ liệu xác thực. Nhóm môn học lưu thông tin học phần, làm điểm nối cho repository và tài nguyên. Nhóm repository lưu dự án mã nguồn, đường dẫn GitHub, mô tả, dữ liệu quét và chỉ số đánh giá.")
    add_para(doc, "Nhóm community chat lưu kênh trao đổi, tin nhắn và trạng thái người dùng. Nhóm thông báo lưu sự kiện cần gửi đến người dùng. Nhóm dữ liệu AI tutor cung cấp ngữ cảnh cho câu hỏi, giúp câu trả lời bám vào dữ liệu học tập thay vì trả lời chung chung.")
    add_para(doc, "Flyway giúp quản lý thay đổi database có thứ tự. Khi thêm bảng hoặc sửa cấu trúc, thay đổi được ghi thành migration riêng thay vì sửa thủ công. Cách này giúp chạy lại backend trên máy khác dễ kiểm soát hơn.")

    # ══════════════════════════════════════════════════════════════
    # CHƯƠNG 3
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "Chương 3. KIẾN TRÚC VÀ XÂY DỰNG ỨNG DỤNG", 1)

    add_heading(doc, "3.1. Kiến trúc backend", 2)
    add_para(doc, "Backend DevOrbit API tổ chức theo kiến trúc nhiều lớp, cách làm quen thuộc trong Spring Boot. Controller nhận request và trả response, service xử lý nghiệp vụ, repository giao tiếp database, entity mô tả dữ liệu, DTO trao đổi với client.")
    add_table(doc, ["Lớp", "Vai trò", "Ví dụ trong hệ thống"], [
        ["Controller", "Nhận request, kiểm tra đầu vào cơ bản và gọi service.", "SubjectQaController, CommunityController, GithubRepoController."],
        ["Service", "Xử lý nghiệp vụ chính của hệ thống.", "SubjectQaService, GithubScanService, CommunityChatService."],
        ["Repository", "Truy vấn và lưu dữ liệu.", "CourseRepository, GithubRepoRepository, UserRepository."],
        ["Entity", "Biểu diễn bảng dữ liệu.", "User, Course, GithubRepo, CommunityMessage."],
        ["Config", "Cấu hình bảo mật, CORS, AI, cache và runtime.", "SecurityConfig, AiConfig, WebSocketConfig."],
    ])
    add_para(doc, "Luồng xử lý phổ biến: request vào controller, controller nhận dữ liệu và chuyển sang DTO, gọi service. Service kiểm tra nghiệp vụ, lấy dữ liệu từ repository, xử lý và trả DTO phản hồi. Cách này giúp controller ngắn gọn, nghiệp vụ tập trung ở service và dễ kiểm thử hơn.")
    add_para(doc, "Nhóm giữ các service có trách nhiệm rõ ràng. Ví dụ, GithubScanService chỉ xử lý quét repository, CommunityChatService chỉ xử lý tin nhắn và kênh, SubjectQaService chỉ xử lý câu hỏi AI tutor. Cách này tránh dồn quá nhiều logic vào một lớp.")
    add_para(doc, "Các lớp config đóng vai trò đưa ứng dụng từ code sang trạng thái chạy được. SecurityConfig, CorsConfig, WebSocketConfig, AiConfig, DataSourceConfig, CacheConfig giúp backend kết nối với client, database và dịch vụ ngoài ổn định hơn.")

    add_heading(doc, "3.2. Các module chính", 2)
    add_para(doc, "Module xác thực chịu trách nhiệm đăng nhập, cấp JWT và bảo vệ API. Khi người dùng đăng nhập, backend kiểm tra thông tin, tạo token và yêu cầu client gửi token trong các request tiếp theo. Nhờ đó, các API quan trọng không bị truy cập tùy ý.")
    add_para(doc, "Module môn học là nền dữ liệu học thuật. Các chức năng repository, tài nguyên và AI tutor đều cần ngữ cảnh môn học để hoạt động có ý nghĩa. Dù giao diện có thay đổi theo từng client, dữ liệu môn học vẫn là phần cốt lõi.")
    add_para(doc, "Module repository và GitHub scan thể hiện rõ đặc trưng đề tài. Backend không chỉ lưu đường dẫn repo mà còn lấy thêm thông tin, phân loại và hỗ trợ đánh giá. Sinh viên xem repo trong bối cảnh học tập, không chỉ xem một trang GitHub riêng lẻ.")
    add_para(doc, "Module community chat giúp hệ thống có yếu tố tương tác giữa người học. Tin nhắn và kênh trao đổi lưu ở backend, WebSocket giúp client nhận cập nhật realtime. Đây là phần cho thấy backend không chỉ phục vụ request một lần mà còn hỗ trợ luồng realtime.")
    add_para(doc, "Module AI tutor xử lý hỏi đáp học tập. Backend nhận câu hỏi, chuẩn bị ngữ cảnh, gọi service AI và trả kết quả. Trong đồ án, nhóm tập trung vào việc luồng này chạy được, có service rõ ràng và có kiểm thử cho tình huống quan trọng.")

    add_heading(doc, "3.3. Xây dựng API", 2)
    add_para(doc, "Các API theo hướng REST, dùng request và response DTO để tách dữ liệu trao đổi với mô hình lưu trữ. Endpoint quan trọng đều đi qua service, giúp code dễ kiểm thử và giảm việc đưa nghiệp vụ vào controller.")
    add_para(doc, "Endpoint đặt tên theo tài nguyên hoặc nhóm chức năng: nhóm API môn học, repository, chat, thông báo, AI tutor. Mỗi API có vai trò rõ: lấy danh sách, xem chi tiết, tạo mới, cập nhật, xóa hoặc thực hiện hành động cụ thể.")
    add_para(doc, "Với API nhận dữ liệu, DTO giúp kiểm soát đầu vào. Entity không đưa trực tiếp ra ngoài khi không cần thiết, vì entity gắn với cấu trúc database và có thể chứa trường không phù hợp cho client.")
    add_para(doc, "Với API trả dữ liệu, response DTO định hình chính xác dữ liệu hiển thị. Lớp giao diện không phụ thuộc cấu trúc bảng trong database. Khi database thay đổi, backend vẫn giữ response ổn định nếu thiết kế DTO tốt.")

    add_heading(doc, "3.4. Bảo mật và xử lý lỗi", 2)
    add_para(doc, "JWT xác thực request, Spring Security cấu hình quyền truy cập. Lỗi phổ biến xử lý tập trung để client nhận thông báo rõ ràng và backend dễ theo dõi khi có sự cố.")
    add_para(doc, "Bảo mật trong đồ án không chỉ là đăng nhập. Một số endpoint chỉ dành cho người đã đăng nhập, một số cần vai trò quản trị, một số cần kiểm tra quyền sở hữu dữ liệu trước khi cho phép thao tác.")
    add_para(doc, "Xử lý lỗi tập trung giảm tình trạng mỗi controller trả lỗi một kiểu. Khi hệ thống có quy ước lỗi chung, client hiển thị thông báo dễ hơn và nhóm phát triển debug cũng nhanh hơn. Đây là điểm quan trọng khi nhiều module và nhiều người cùng phát triển.")

    add_heading(doc, "3.5. Kiểm thử và vận hành", 2)
    add_para(doc, "Nhóm viết test cho nhiều service quan trọng và dùng script chạy để thuận tiện khởi động backend. Phần này cần thiết để sản phẩm không chỉ chạy trên máy một thành viên mà tái lập được ở môi trường khác.")
    add_para(doc, "Test tập trung vào phần có nghiệp vụ rõ hoặc dễ lỗi: xác thực, AI tutor, community chat, hardening database, service xử lý repository. Test giúp phát hiện lỗi sớm khi sửa code, đặc biệt giai đoạn nhiều thành viên cùng thay đổi source.")
    add_para(doc, "Backend đọc cấu hình từ file .env hoặc biến môi trường để phù hợp nhiều máy. Database URL, token, khóa dịch vụ ngoài không viết cố định trong code. Cách này giúp sản phẩm dễ chạy lại khi chuyển máy hoặc môi trường triển khai.")
    add_para(doc, "Nhóm cũng chú ý log và thông báo lỗi khi khởi động. Backend đồ án không chỉ cần chạy được khi mọi thứ đúng, mà cần báo lỗi đủ rõ khi thiếu cấu hình, sai database hoặc dịch vụ ngoài chưa sẵn sàng. Điều này giúp demo và bảo trì bớt phụ thuộc một thành viên duy nhất.")

    # ══════════════════════════════════════════════════════════════
    # CHƯƠNG 4
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "Chương 4. KẾT QUẢ SẢN PHẨM", 1)

    add_heading(doc, "4.1. Kết quả đạt được", 2)
    add_para(doc, "Sau quá trình thực hiện, nhóm xây dựng được backend DevOrbit API với nhiều chức năng: quản lý học thuật, repository, community chat và AI tutor. Source code tách lớp theo Spring Boot, có migration database và có test cho phần quan trọng.")
    add_caption(doc, "Bảng 4.1: Kết quả sản phẩm đạt được")
    add_table(doc, ["Nhóm kết quả", "Mô tả"], [
        ["Backend Java API", "Cung cấp API cho đăng nhập, môn học, repository, chat, thông báo và AI tutor."],
        ["Database", "Có entity và migration để quản lý dữ liệu PostgreSQL."],
        ["Bảo mật", "Có JWT, phân quyền và cấu hình bảo vệ endpoint."],
        ["Realtime", "Có WebSocket phục vụ community chat."],
        ["AI tutor", "Có luồng hỏi đáp học tập và xử lý dữ liệu liên quan."],
        ["Client minh họa", "Giao diện web hỗ trợ kiểm chứng các API chính."],
    ])
    add_para(doc, "Kết quả quan trọng nhất là backend đã có bộ khung tương đối đầy đủ. Các chức năng có liên kết với nhau: người dùng đăng nhập, xem môn học, tìm repo, tham gia trao đổi, nhận thông báo và dùng AI tutor trong cùng một hệ thống.")
    add_para(doc, "Phần database có migration mô tả thay đổi cấu trúc theo thời gian. Khi cần kiểm tra hoặc chạy lại hệ thống, nhóm dựa vào migration để tái tạo database. Cách này rõ ràng hơn nhiều so với tạo bảng thủ công.")
    add_para(doc, "Phần bảo mật có nền tảng đăng nhập và phân quyền. Đây là yêu cầu bắt buộc với hệ thống có dữ liệu người dùng và chức năng cộng đồng. Nếu thiếu lớp bảo mật, các API như gửi tin nhắn hoặc thao tác quản trị sẽ không đủ an toàn.")
    add_para(doc, "Phần AI tutor là điểm mở rộng đáng chú ý. Nhóm xây dựng được luồng backend để nhận câu hỏi, xử lý ngữ cảnh và trả phản hồi. Dù chưa phải hệ thống AI hoàn chỉnh, nó cho thấy backend có khả năng tích hợp chức năng thông minh ngoài CRUD.")

    add_heading(doc, "4.2. Một số màn hình và luồng sử dụng", 2)
    add_para(doc, "Các màn hình giao diện dùng để minh họa cách client tiêu thụ API. Người dùng có thể đăng nhập, xem môn học, tìm repo, trao đổi cộng đồng và dùng AI tutor.")
    add_image(doc, str(ROOT / "docs/report/screenshots/01-homepage.png"), "Hình 4.1: Trang chủ DevOrbit")
    add_para(doc, "Luồng đăng nhập: người dùng nhập thông tin ở client, client gửi request đến backend, backend kiểm tra và trả token nếu hợp lệ. Client dùng token đó cho các request tiếp theo. Luồng này minh họa cách Spring Security và JWT hoạt động trong hệ thống.")
    add_image(doc, str(ROOT / "docs/report/screenshots/05-login.png"), "Hình 4.2: Giao diện đăng nhập")
    add_para(doc, "Luồng repository: người dùng xem hoặc tìm repo theo môn học. Khi cần cập nhật từ GitHub, backend gọi service quét, lấy metadata và lưu lại. Client chỉ gọi API DevOrbit thay vì tự xử lý logic GitHub.")
    add_image(doc, str(ROOT / "docs/report/screenshots/02-courses.png"), "Hình 4.3: Trang danh sách môn học")
    add_image(doc, str(ROOT / "docs/report/screenshots/03-courses-list.png"), "Hình 4.4: Kết quả tìm kiếm môn học")
    add_para(doc, "Luồng community chat: người dùng chọn kênh, backend trả tin nhắn cũ qua API và gửi tin mới qua WebSocket. Cách kết hợp REST API và WebSocket giúp hệ thống có cả dữ liệu lịch sử lẫn trải nghiệm realtime.")
    add_image(doc, str(ROOT / "docs/report/screenshots/04-community.png"), "Hình 4.5: Trang Cộng đồng")
    add_para(doc, "Ngoài ra, hệ thống còn có Photobooth trực tuyến và Lộ trình học tập giúp sinh viên lên kế hoạch 4 năm tại UIT.")
    add_image(doc, str(ROOT / "docs/report/screenshots/06-photobooth.png"), "Hình 4.6: Trang Photobooth")
    add_image(doc, str(ROOT / "docs/report/screenshots/07-knowledge-graph.png"), "Hình 4.7: Lộ trình học tập")
    add_para(doc, "Luồng AI tutor: người dùng đặt câu hỏi, backend nhận, chuẩn bị dữ liệu liên quan, xử lý qua service AI và trả lời. Luồng này cho thấy backend đóng vai trò điều phối giữa client, dữ liệu hệ thống và thành phần AI.")

    add_heading(doc, "4.3. Mức độ hoàn thiện", 2)
    add_para(doc, "Sản phẩm đáp ứng được các chức năng chính trong phạm vi đồ án. Một số phần còn có thể mở rộng: tối ưu hiệu năng khi dữ liệu lớn, bổ sung kiểm thử tích hợp, và hoàn thiện trải nghiệm client.")
    add_para(doc, "Ở mức đồ án môn học, hệ thống đủ các phần chính để trình bày: backend Java, database, bảo mật, API, realtime, AI tutor và client minh họa. Nhóm không chỉ làm một chức năng đơn lẻ mà đã xây hệ thống nhiều module phối hợp.")
    add_para(doc, "Tuy nhiên, sản phẩm vẫn có giới hạn. Một số chức năng cần thêm integration test để chắc chắn hoạt động khi kết nối database thật và dịch vụ ngoài. Một số luồng client cần tinh chỉnh thêm. AI tutor cần thêm dữ liệu học tập phong phú hơn để câu trả lời có chiều sâu.")
    add_para(doc, "Nhóm đánh giá mức hoàn thiện hiện tại phù hợp với phạm vi đồ án cuối kỳ. Sản phẩm đủ để demo các chức năng chính, trình bày kiến trúc Java backend, và có cơ sở tiếp tục phát triển nếu thêm thời gian.")

    # ══════════════════════════════════════════════════════════════
    # CHƯƠNG 5
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "Chương 5. ĐÁNH GIÁ KẾT QUẢ THỰC HIỆN VÀ PHÂN CÔNG CÔNG VIỆC", 1)

    add_heading(doc, "5.1. Đánh giá chung", 2)
    add_para(doc, "Nhóm đã hoàn thành mục tiêu xây dựng backend Java nhiều chức năng thực tế, tổ chức source code rõ ràng và có khả năng mở rộng. Backend là trọng tâm đồ án, giao diện hỗ trợ trình bày và kiểm chứng.")
    add_para(doc, "Về tiến độ, nhóm có 400 commit trong giai đoạn 14/03/2026 đến 14/06/2026 trên nhánh master. Các commit được tổng hợp từ GitHub Contributors, loại trừ merge commit, để phản ánh tương đối quá trình đóng góp từng thành viên.")
    add_para(doc, "Về sản phẩm, nhóm xây dựng nhiều chức năng có liên hệ với nhau, không chỉ làm các màn hình rời rạc. Backend phục vụ được luồng chính: đăng nhập, dữ liệu học thuật, quản lý repository, chat cộng đồng và AI tutor. Đây là điểm quan trọng vì đồ án Java cần thể hiện cách tổ chức backend và xử lý nghiệp vụ.")
    add_para(doc, "Về kỹ thuật, hệ thống dùng các thành phần quen thuộc của Spring Boot: controller, service, repository, entity, DTO, config và test. Nhóm cố gắng tách trách nhiệm giữa các lớp, dùng migration cho database và giữ cấu hình ở mức tái sử dụng. Những điểm này giúp sản phẩm không chỉ chạy được mà còn dễ đọc khi bảo trì.")
    add_para(doc, "Về làm việc nhóm, commit cho thấy các thành viên có phạm vi đóng góp khác nhau. Có người tập trung backend lõi, có người làm community và trải nghiệm sử dụng, có người xử lý cấu hình và hardening, có người đào sâu repository discovery. Cách phân công giúp sản phẩm bao phủ nhiều mảng, dù khối lượng chưa cân bằng hoàn toàn.")

    add_heading(doc, "5.2. Phân công công việc", 2)
    add_caption(doc, "Bảng 5.1: Phân công công việc và đánh giá kết quả")
    add_table(doc, ["TV", "Số commit", "Phân công chính", "Đánh giá kết quả"], [
        ["A", str(counts["A"]), "Backend lõi (entities, API, JWT auth), AI/RAG (subject QA, streaming chat, embeddings), knowledge graph (impact scoring, simulation mode), frontend utilities, CI/CD.", "Đóng góp lớn nhất về số lượng, đảm nhận phần kỹ thuật phức tạp nhất của hệ thống."],
        ["B", str(counts["B"]), "Community chat (REST/WebSocket, channel presence, subscriptions), GPA calculator (goal planner, presets, autosave, what-if projections), mobile features (bookmarks, search filters, tech stack).", "Xây dựng tính năng community từ đầu đến GPA calculator đầy đủ chức năng."],
        ["C", str(counts["C"]), "Admin panel rewrite, auth system (OTP, forgot password), photobooth backend API, knowledge graph (elective selection, SE2025 curriculum), hardening (CI/CD, Dependabot, dead file cleanup).", "Bao phủ nhiều mảng từ admin đến auth, photobooth đến infrastructure."],
        ["D", str(counts["D"]), "Repository discovery (search matching, alias map, query intent classifier), repo evaluation (X-ray summaries, type classification, dynamic rating), community features (vote/review, WebSocket chat, guest gating).", "Phát triển chức năng phân tích repository, đặc trưng nhất của đề tài."],
    ])
    add_para(doc, "Bảng trên tổng hợp từ commit trên nhánh master, chỉ tính từ 14/03/2026 đến 14/06/2026, loại trừ merge commit. Nhóm đọc commit message, author và phạm vi file thay đổi để gom theo nội dung công việc, không liệt kê từng commit trong báo cáo.")
    add_para(doc, "Số commit chỉ là dấu vết tham gia, không phải thước đo duy nhất. Có phần cần nhiều commit nhỏ để chỉnh dần, cũng có phần ít commit hơn nhưng đòi hỏi đọc hiểu kỹ và sửa đúng điểm. Đánh giá cá nhân bên dưới tập trung vào vai trò và giá trị mà từng thành viên mang lại.")

    add_heading(doc, "5.3. Đánh giá cá nhân", 2)
    add_para(doc, "Thành viên A đóng góp nhiều nhất về số lượng và cũng đảm nhận phần kỹ thuật phức tạp nhất. A xây backend lõi từ đầu: entity, API, JWT auth, admin course management, GitHub candidate scanning. Sau đó A phát triển toàn bộ hệ thống AI/RAG gồm subject QA streaming chat, Fireworks embedding provider, RAG semantic retrieval với hybrid retrieval, query expansion và reranking. A cũng triển khai knowledge graph với impact scoring, simulation mode và cascade animations.")
    add_para(doc, "Ngoài backend, A đóng góp nhiều vào frontend: cosmic redesign, shared components (Avatar, LoadingSkeleton, Toast), Vietnamese localization, performance monitoring utilities và hàng chục unit test. Về infrastructure, A enforce CI/CD pre-push gate, thực hiện security hardening và tối ưu database initialization với batch inserts. A giữ vai trò trục kỹ thuật chính, đảm bảo phần lõi hoạt động ổn định.")
    add_para(doc, "Thành viên B xây dựng community chat từ đầu, gồm REST/WebSocket endpoints, channel presence tracking, online members UI và subscription management. B phát triển đầy đủ GPA calculator với goal what-if projections, semester presets từ roadmap, autosave draft, cumulative GPA estimate mode và background customization. Trên mobile, B đóng góp course bookmarks, search filters, repo tech stack filter và AI roadmap QA viewmodels.")
    add_para(doc, "Điểm đáng chú ý ở B là tính hệ thống: mỗi tính năng phát triển từ backend đến frontend, từ web đến mobile. Community chat không chỉ có WebSocket mà còn có entity, migration, REST endpoints và test contract. GPA calculator không chỉ có giao diện mà còn có goal planner, draft persistence và tích hợp roadmap. Cách này giúp sản phẩm có độ hoàn thiện cao ở các tính năng B phụ trách.")
    add_para(doc, "Thành viên C tham gia sâu nhiều mảng. Về quản trị, C viết lại admin panel với layout mới, bổ sung hệ thống thông báo và chỉnh UI dashboard sinh viên. Về xác thực, C triển khai đầy đủ luồng đăng nhập, đăng ký, quên mật khẩu và xác thực OTP qua email. Phần này ảnh hưởng trực tiếp đến trải nghiệm người dùng đầu tiên khi tiếp cận hệ thống.")
    add_para(doc, "C cũng xây dựng backend API cho Photobooth thay vì dùng Supabase trực tiếp, xử lý frame management, per-slot filters và crop/zoom. Về knowledge graph, C cải thiện card chọn môn tự chọn và tích hợp chương trình SE2025 với AI roadmap generator. C khởi tạo dự án DevOrbit Mobile và triển khai Room migration. Về vận hành, C bổ sung CI/CD pipelines, Dependabot, ESLint, các file bảo vệ repo và refactor loại bỏ hàng chục file chết trên mobile, web và API. Đóng góp của C vừa rộng vừa chạm nhiều layer, từ admin đến auth, từ mobile đến infrastructure.")
    add_para(doc, "Thành viên D tập trung vào repository discovery và evaluation, phần đặc trưng nhất của DevOrbit. D xây hệ thống search matching với alias map, repo-based course matching và query intent classifier giúp tìm repo chính xác hơn. Về repository evaluation, D phát triển X-ray repo summaries, repo type classification, signal detection, contextual analysis và dynamic rating. Những tính năng này giúp DevOrbit khác với danh sách link GitHub thông thường.")
    add_para(doc, "D cũng đóng góp community features: repo vote và review system, channel subscribe/unsubscribe, guest gating và real-time WebSocket chat với STOMP và JWT auth. Trên frontend, D cải thiện course search relevance, pagination, review component với star rating và share dialog. Dù commit ít hơn, phần việc của D có định hướng rõ và tạo giá trị riêng: DevOrbit không chỉ quản lý môn học mà còn phân tích và đánh giá mã nguồn học thuật.")
    add_para(doc, "Một nhóm đồ án không chỉ cần người làm rộng, mà cần người giữ một mảng đủ rõ để sản phẩm có chiều sâu. Đánh giá cá nhân không xem số commit thấp hơn là đóng góp thấp hơn, mà nhìn vào vai trò thực tế của module đối với mục tiêu chung.")

    add_heading(doc, "5.4. Nhận xét về kết quả", 2)
    add_para(doc, "Điểm mạnh là nhóm chia được các mảng rõ và có sản phẩm chạy được nhiều chức năng. Điểm cần cải thiện là khối lượng công việc giữa các thành viên chưa thật đồng đều, một số chức năng còn thiếu kiểm thử tích hợp và tài liệu hướng dẫn.")
    add_para(doc, "Theo mục tiêu môn học, đồ án thể hiện được nhiều kiến thức Java backend: tổ chức package, xây dựng REST API, xử lý service, kết nối database, migration, bảo mật, realtime và kiểm thử. Đây là nội dung liên hệ trực tiếp với yêu cầu học phần.")
    add_para(doc, "Theo sản phẩm, DevOrbit có hướng đi rõ và phát triển tiếp được. Các module hiện tại đủ để tạo nền cho hệ thống hỗ trợ sinh viên quản lý và khám phá mã nguồn học thuật. Nhưng để hoàn chỉnh, cần thêm integration test, dữ liệu thật phong phú hơn và cơ chế quản trị chi tiết hơn.")
    add_para(doc, "Nhóm cũng rút ra kinh nghiệm: làm hệ thống nhiều module đòi hỏi thống nhất quy ước code từ sớm. Nếu không có quy ước rõ về tên API, DTO, cách xử lý lỗi và cách viết test, quá trình ghép module dễ phát sinh lỗi.")

    # ══════════════════════════════════════════════════════════════
    # CHƯƠNG 6
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "Chương 6. THUẬN LỢI, KHÓ KHĂN VÀ BÀI HỌC KINH NGHIỆM", 1)

    add_heading(doc, "6.1. Thuận lợi", 2)
    add_para(doc, "Nhóm có sẵn nền tảng Java và Spring Boot từ học phần. Đề tài cũng gần với nhu cầu thực tế của sinh viên, nên nhóm dễ hình dung người dùng sẽ cần gì. Việc dùng GitHub làm nguồn dữ liệu giúp dễ kiểm chứng và trình bày kết quả.")
    add_para(doc, "Hệ sinh thái Spring Boot có nhiều thư viện hỗ trợ sẵn cho REST API, bảo mật, WebSocket, cache và kiểm thử. Nhờ đó nhóm tập trung vào tổ chức hệ thống và nghiệp vụ DevOrbit, không phải xây hạ tầng từ đầu.")
    add_para(doc, "Dữ liệu và bối cảnh dự án gần với sinh viên UIT. Chức năng tìm repo theo môn, trao đổi cộng đồng hay hỏi AI tutor đều xuất phát từ nhu cầu quen thuộc trong quá trình học lập trình.")

    add_heading(doc, "6.2. Khó khăn", 2)
    add_para(doc, "Khó khăn lớn nhất là hệ thống có nhiều module liên quan nhau: xác thực, database, repository, realtime chat, AI tutor. Khi một module thay đổi, các phần còn lại phải kiểm tra để tránh lỗi dây chuyền. Phối hợp backend và giao diện minh họa cũng làm tăng khối lượng kiểm thử.")
    add_para(doc, "Tích hợp dịch vụ ngoài như GitHub hoặc AI tutor có nhiều tình huống không ổn định. Dịch vụ ngoài có thể thay đổi dữ liệu, giới hạn request hoặc trả lỗi ngoài dự đoán. Backend vì vậy cần xử lý lỗi tốt hơn so với chức năng chỉ đọc ghi database nội bộ.")
    add_para(doc, "Quản lý database cũng cần cẩn thận. Khi nhiều thành viên cùng thêm entity hoặc migration, nếu không thống nhất quy ước đặt tên và thứ tự, hệ thống dễ lỗi khi chạy lại từ đầu. Nhóm dùng Flyway để giảm rủi ro, nhưng vẫn phải kiểm tra kỹ mỗi khi thay đổi cấu trúc dữ liệu.")
    add_para(doc, "Do phạm vi đồ án rộng, nhóm phải cân bằng giữa thêm chức năng mới và làm chắc chức năng đã có. Chỉ thêm mà không test thì sản phẩm dễ có lỗi nhỏ. Chỉ tập trung test thì tiến độ demo chậm. Đây là bài toán nhóm phải điều chỉnh trong suốt quá trình thực hiện.")

    add_heading(doc, "6.3. Bài học kinh nghiệm", 2)
    add_para(doc, "Qua đồ án, nhóm thấy rằng xây backend không chỉ là viết cho đủ endpoint. Hệ thống muốn đứng vững cần chia lớp rõ, dữ liệu quản lý có trật tự, lỗi xử lý có trách nhiệm và thành viên hiểu phần việc của nhau. Khi những điều này chưa rõ, code có thể chạy được một thời điểm, nhưng rất khó sửa và khó mở rộng.")
    add_para(doc, "Bài học thứ nhất: nghiệp vụ nên ở đúng vị trí. Controller tiếp nhận request và trả response, service là nơi thể hiện suy nghĩ chính của hệ thống. Khi service gánh phần xử lý, code mạch lạc hơn, test dễ viết hơn và mỗi thay đổi ít gây ảnh hưởng dây chuyền hơn.")
    add_para(doc, "Bài học thứ hai: dữ liệu cần được tôn trọng từ đầu. Entity, DTO và migration không chỉ là phần phụ trợ, mà là cách hệ thống tự mô tả. Dữ liệu thiết kế vội thì chức năng phía trên phải sửa đi sửa lại. Migration không rõ thì mỗi máy chạy một kiểu.")
    add_para(doc, "Bài học thứ ba: kiểm thử không nên để cuối. Dự án nhỏ, chạy thử thủ công có vẻ nhanh. Nhưng khi số module tăng, mỗi lần sửa service đều có thể ảnh hưởng controller, repository hoặc client. Test giúp giữ sự tự tin sau mỗi thay đổi, nhất ở phần xác thực, chat, AI tutor và xử lý repository.")
    add_para(doc, "Bài học thứ tư: giao tiếp trong nhóm quan trọng không kém kỹ thuật. Mỗi thành viên phụ trách một mảng khác nhau, nhưng đều cần nối lại thành sản phẩm chung. Khi hiểu vai trò của nhau, đánh giá đóng góp cũng công bằng hơn.")
    add_para(doc, "Bài học cuối cùng: đồ án tốt không cần hoàn hảo, nhưng cần trung thực với những gì đã làm được và những gì còn thiếu. DevOrbit còn nhiều hướng phát triển, nhưng quá trình thực hiện giúp nhóm hiểu rõ hơn cách biến kiến thức Java trong lớp thành hệ thống có cấu trúc, có dữ liệu, có người dùng.")

    add_heading(doc, "6.4. Hướng phát triển", 2)
    add_para(doc, "Gần nhất, nhóm có thể bổ sung integration test cho luồng quan trọng: đăng nhập, repository, community chat và AI tutor. Khi có integration test, nhóm tự tin hơn mỗi khi sửa code backend hoặc cập nhật database.")
    add_para(doc, "Tiếp theo, hệ thống có thể mở rộng repository analysis để đọc thêm cấu trúc thư mục, ngôn ngữ sử dụng, file README và thông tin commit. Dữ liệu này giúp đánh giá repo có cơ sở hơn và hỗ trợ sinh viên tìm dự án phù hợp.")
    add_para(doc, "AI tutor có thể cải thiện bằng cách nạp thêm tài liệu học phần, slide, bài tập mẫu hoặc hướng dẫn thực hành. Nguồn dữ liệu phong phú hơn thì câu trả lời bám sát môn học hơn và hữu ích hơn.")
    add_para(doc, "Về vận hành, nhóm có thể bổ sung logging, monitoring và tài liệu triển khai chi tiết. Đây là phần thường chưa được ưu tiên trong đồ án đầu tiên, nhưng rất quan trọng nếu muốn đưa hệ thống đến gần môi trường sử dụng thật.")

    # ── Tài liệu tham khảo ──
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        ("Spring Boot Documentation", "https://docs.spring.io/spring-boot/"),
        ("Spring Security Reference", "https://docs.spring.io/spring-security/reference/"),
        ("PostgreSQL Documentation", "https://www.postgresql.org/docs/"),
        ("Flyway Documentation", "https://documentation.red-gate.com/fd"),
        ("GitHub REST API Documentation", "https://docs.github.com/en/rest"),
        ("DevOrbit GitHub Contributors", "https://github.com/huyhoang171106/DevOrbit/graphs/contributors?from=3%2F14%2F2026"),
    ]
    for idx, (label, url) in enumerate(refs, start=1):
        p = add_para(doc, f"[{idx}] {label}. ", first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        add_hyperlink(p, url, url)

    doc.save(OUTPUT)
    print(f"created={OUTPUT}")
    print(f"commits={dict(counts)} total={total}")


if __name__ == "__main__":
    build_report()
